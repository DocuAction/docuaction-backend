"""Shared builder for AGT compliance documents.

One builder, not 43 hand-written scripts. Every document gets the same cover
page, heading styles, control-mapping table format, version history and approval
block, so a reviewer moving between them is reading one document set rather than
43 differently-shaped files. Divergent formatting across a policy suite reads as
carelessness to an assessor, and it is the first thing they notice.

The content lives in per-document spec modules; this file owns only presentation.
"""

from __future__ import annotations

import datetime as _dt
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

COMPANY = "Alliance Global Tech Inc."
COMPANY_SHORT = "AGT"
PLATFORM = "DocuAction TEFCA ARC"
CEO = "Imran Siddiqui"
NAVY = RGBColor(0x0F, 0x2B, 0x4C)
ACCENT = RGBColor(0x0F, 0x6C, 0xBD)
GREY = RGBColor(0x5A, 0x5A, 0x5A)

# Every document states the same platform facts. Stated once here so a stack
# change is one edit, not 43 - and so no document can quietly contradict another.
PLATFORM_FACTS = [
    ("Frontend", "Next.js 14 (React), static export, Azure Static Web Apps"),
    ("Backend", "FastAPI on Python 3.12, Azure App Service (Linux)"),
    ("Database", "Azure Database for PostgreSQL - Flexible Server, geo-redundant"),
    ("Identity", "Microsoft Entra ID; OAuth 2.0, OIDC, JWT bearer, RBAC, MFA"),
    ("Secrets", "Azure Key Vault, resolved at runtime by system-assigned managed identity"),
    ("Monitoring", "Azure Monitor and Application Insights"),
    ("Healthcare standards", "FHIR R4, HL7, TEFCA Common Agreement / RCE IG v1.14.0"),
    ("Certifications", "CMMI Level 3, ISO 27001, ISO 9001, ISO 20000-1, DoD Final Contractor"),
]


def _shade(cell, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


class AGTDoc:
    def __init__(self, doc_id: str, title: str, version: str = "1.0",
                 date: str = "2026-07-28", classification: str = "Internal"):
        self.doc_id, self.title = doc_id, title
        self.version, self.date, self.classification = version, date, classification
        self.d = Document()
        self._styles()
        self._cover()
        self._page_footer()

    # ---------- presentation ----------
    def _styles(self) -> None:
        n = self.d.styles["Normal"]
        n.font.name, n.font.size = "Calibri", Pt(10.5)
        n.paragraph_format.space_after = Pt(6)
        n.paragraph_format.line_spacing = 1.15
        for lvl, size in ((1, 16), (2, 13), (3, 11.5)):
            st = self.d.styles[f"Heading {lvl}"]
            st.font.name, st.font.size = "Calibri", Pt(size)
            st.font.color.rgb = NAVY if lvl < 3 else ACCENT
            st.font.bold = True
        for s in self.d.sections:
            s.left_margin = s.right_margin = Inches(1.0)
            s.top_margin = s.bottom_margin = Inches(0.9)

    def _cover(self) -> None:
        for _ in range(4):
            self.d.add_paragraph()
        p = self.d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(COMPANY_SHORT); r.font.size, r.font.bold, r.font.color.rgb = Pt(44), True, NAVY
        p = self.d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(COMPANY); r.font.size, r.font.color.rgb = Pt(13), GREY
        self.d.add_paragraph()
        p = self.d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(self.title); r.font.size, r.font.bold, r.font.color.rgb = Pt(24), True, NAVY
        p = self.d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(PLATFORM); r.font.size, r.font.italic, r.font.color.rgb = Pt(13), True, ACCENT
        for _ in range(3):
            self.d.add_paragraph()
        t = self.d.add_table(rows=0, cols=2); t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for k, v in (("Document ID", self.doc_id), ("Version", self.version),
                     ("Effective date", self.date), ("Classification", self.classification),
                     ("Owner", f"Chief Executive Officer, {COMPANY_SHORT}"),
                     ("Approved by", CEO), ("Review cycle", "Annual, or on material change")):
            row = t.add_row().cells
            row[0].width, row[1].width = Inches(1.9), Inches(3.6)
            rr = row[0].paragraphs[0].add_run(k); rr.font.bold = True; rr.font.size = Pt(10)
            _shade(row[0], "EEF3F8")
            rv = row[1].paragraphs[0].add_run(v); rv.font.size = Pt(10)
        self.d.add_page_break()

    def _page_footer(self) -> None:
        for s in self.d.sections:
            p = s.footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{self.doc_id} v{self.version}  |  {COMPANY_SHORT} {self.classification}  |  "
                          f"Uncontrolled when printed")
            r.font.size, r.font.color.rgb = Pt(8), GREY

    # ---------- content ----------
    def h1(self, text): self.d.add_heading(text, 1); return self
    def h2(self, text): self.d.add_heading(text, 2); return self
    def h3(self, text): self.d.add_heading(text, 3); return self

    def p(self, text):
        self.d.add_paragraph(text); return self

    def bullets(self, items):
        for i in items:
            self.d.add_paragraph(str(i), style="List Bullet")
        return self

    def numbered(self, items):
        for i in items:
            self.d.add_paragraph(str(i), style="List Number")
        return self

    def table(self, headers, rows, widths=None):
        t = self.d.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            r = c.paragraphs[0].add_run(str(h)); r.font.bold = True; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade(c, "0F2B4C")
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row[:len(headers)]):
                rr = cells[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(9)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(w)
        self.d.add_paragraph()
        return self

    def note(self, text):
        p = self.d.add_paragraph()
        r = p.add_run(text); r.font.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
        return self

    def page_break(self):
        self.d.add_page_break(); return self

    # ---------- standard sections ----------
    def platform_context(self):
        self.h2("System context")
        self.p(f"This document applies to {PLATFORM}, operated by {COMPANY} on Microsoft "
               f"Azure. The control descriptions in this document are written against the "
               f"following implementation, not against a generic reference architecture.")
        self.table(["Layer", "Implementation"], PLATFORM_FACTS, widths=[1.8, 4.4])
        return self

    def compliance_mapping(self, rows):
        self.h2("Compliance mapping")
        self.p("Each framework reference below identifies the control this document is "
               "written to satisfy. Where a control is only partially satisfied, the gap is "
               "stated in the row rather than omitted - a mapping table that shows only "
               "green is not evidence, it is decoration.")
        self.table(["Framework", "Control reference", "How this document satisfies it", "Status"],
                   rows, widths=[1.1, 1.5, 3.0, 0.8])
        return self

    def related(self, rows):
        self.h2("Related documents")
        self.table(["Document ID", "Title", "Relationship"], rows, widths=[1.3, 2.9, 2.2])
        return self

    def roles(self, rows):
        self.h2("Roles and responsibilities")
        self.table(["Role", "Responsibility"], rows, widths=[1.8, 4.4])
        return self

    def definitions(self, rows):
        self.h2("Definitions")
        self.table(["Term", "Definition"], rows, widths=[1.5, 4.7])
        return self

    def closing(self, history=None):
        self.page_break()
        self.h1("Version history")
        self.table(["Version", "Date", "Author", "Summary of change"],
                   history or [["1.0", self.date, CEO, "Initial issue."]],
                   widths=[0.8, 1.1, 1.4, 2.9])
        self.h1("Approval")
        self.p("This document is approved for use and is binding on all personnel and "
               "contractors with access to the systems in scope.")
        self.table(["Role", "Name", "Signature", "Date"],
                   [["Chief Executive Officer", CEO, "", ""],
                    ["Security Officer", "", "", ""],
                    ["Privacy Officer", "", "", ""]],
                   widths=[1.8, 1.6, 1.6, 1.2])
        self.note("Approval is recorded in the document management system; a wet signature "
                  "is not required where an auditable electronic approval exists.")
        return self

    def save(self, directory: Path) -> Path:
        directory.mkdir(parents=True, exist_ok=True)
        safe = self.title.replace(" ", "_").replace("/", "-").replace(":", "")
        out = directory / f"{self.doc_id}_{safe}.docx"
        self.d.save(out)
        return out
