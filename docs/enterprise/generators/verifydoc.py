"""Structural verification of a generated .docx.

LibreOffice and pdftoppm are not available on this machine, so a PDF render and
visual page review cannot be performed. This checks the things a render would
have shown — branding, page setup, header/footer, heading hierarchy, table
integrity, empty cells — directly against the document XML.
"""
import pathlib
import sys

from docx import Document
from docx.shared import RGBColor

NAVY = RGBColor(0x0B, 0x3C, 0x5D)


def verify(path: pathlib.Path):
    d = Document(str(path))
    print(f"\n=== {path.name} ({path.stat().st_size:,} bytes) ===")

    s = d.sections[0]
    print(f"  page      : {s.page_width.inches:.1f}\" x {s.page_height.inches:.1f}\" "
          f"(expect 8.5 x 11.0)")
    print(f"  margins   : L{s.left_margin.inches:.2f} R{s.right_margin.inches:.2f} "
          f"T{s.top_margin.inches:.2f} B{s.bottom_margin.inches:.2f}")
    print(f"  font      : {d.styles['Normal'].font.name} "
          f"{d.styles['Normal'].font.size.pt:.0f}pt (expect Arial 10)")

    hdr = " | ".join(p.text for p in s.header.paragraphs if p.text.strip())
    ftr = " | ".join(p.text for p in s.footer.paragraphs if p.text.strip())
    print(f"  header    : {hdr.strip() or '(EMPTY)'}")
    print(f"  footer    : {ftr.strip() or '(EMPTY)'}")

    heads = {}
    for p in d.paragraphs:
        st = p.style.name
        if st.startswith("Heading"):
            heads[st] = heads.get(st, 0) + 1
    print(f"  headings  : " + ", ".join(f"{k}={v}" for k, v in sorted(heads.items())))

    h1s = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    print(f"  H1 count  : {len(h1s)}")

    # table integrity
    empty_rows, ragged, total_cells = 0, 0, 0
    for t in d.tables:
        ncol = len(t.columns)
        for r in t.rows:
            cells = r.cells
            total_cells += len(cells)
            if len(cells) != ncol:
                ragged += 1
            if all(not c.text.strip() for c in cells):
                empty_rows += 1
    print(f"  tables    : {len(d.tables)}  rows={sum(len(t.rows) for t in d.tables)}  "
          f"cells={total_cells}")
    print(f"  integrity : ragged_rows={ragged}  fully_empty_rows={empty_rows} "
          f"(both should be 0)")

    # Branding lives on the STYLE, not on individual runs — add_heading()
    # produces runs with no explicit colour and inherits from the style.
    style_navy = [n for n in ("Heading 1", "Heading 2", "Heading 3")
                  if d.styles[n].font.color and d.styles[n].font.color.rgb == NAVY]
    print(f"  navy styles: {', '.join(style_navy) or 'NONE'}")

    # Field codes. Header/footer live in SEPARATE XML parts, so the body xml
    # alone will always report PAGE missing.
    body_xml = d.element.xml
    ftr_xml = "".join(s.footer._element.xml for s in d.sections)
    hdr_xml = "".join(s.header._element.xml for s in d.sections)
    all_xml = body_xml + ftr_xml + hdr_xml
    print(f"  TOC field : {'present' if 'TOC ' in body_xml else 'MISSING'}")
    print(f"  PAGE field: {'present' if 'PAGE' in ftr_xml else 'MISSING'}")

    words = sum(len(p.text.split()) for p in d.paragraphs)
    tblwords = sum(len(c.text.split()) for t in d.tables for r in t.rows
                   for c in r.cells)
    print(f"  word count: {words:,} body + {tblwords:,} in tables "
          f"= {words + tblwords:,}")

    issues = []
    if abs(s.page_width.inches - 8.5) > 0.01: issues.append("page width")
    if d.styles["Normal"].font.name != "Arial": issues.append("body font")
    if ragged: issues.append("ragged table rows")
    if empty_rows: issues.append("empty table rows")
    if "TOC " not in body_xml: issues.append("no TOC field")
    if "PAGE" not in ftr_xml: issues.append("no PAGE field in footer")
    if len(style_navy) < 3: issues.append("heading styles not navy")
    if not hdr.strip(): issues.append("empty header")
    if not ftr.strip(): issues.append("empty footer")
    print(f"  RESULT    : {'PASS' if not issues else 'ISSUES — ' + ', '.join(issues)}")
    return not issues


if __name__ == "__main__":
    ok = True
    for arg in sys.argv[1:]:
        ok &= verify(pathlib.Path(arg))
    print(f"\noverall: {'PASS' if ok else 'ISSUES FOUND'}")
