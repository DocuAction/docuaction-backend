"""AGT branded .docx builder — shared by all enterprise documentation volumes.

Branding is fixed here rather than per-document so the four volumes cannot drift
apart: navy #0B3C5D, blue #0078D4, gold #D4A843, Arial, 10pt body, US Letter.
"""
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x0B, 0x3C, 0x5D)
BLUE = RGBColor(0x00, 0x78, 0xD4)
GOLD = RGBColor(0xD4, 0xA8, 0x43)
NAVY_HEX = "0B3C5D"
BLUE_HEX = "0078D4"
GOLD_HEX = "D4A843"
GREY_HEX = "F2F4F7"

ORG = "Alliance Global Tech, Inc."
CONTRACT = "7571MN26F80064"
CAGE = "8ERE8"
UEI = "MP2FLV1MAW93"


def _shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def _field(paragraph, instr):
    """Insert a Word field (PAGE, NUMPAGES, TOC) that Word evaluates on open."""
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._r.append(fld)
    r2 = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r2._r.append(it)
    r3 = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)
    r4 = paragraph.add_run("1")
    r5 = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)
    return r4


class AGTDoc:
    def __init__(self, doc_id, title, subtitle, version="1.0",
                 date="August 2026", author="Imran Siddiqui"):
        self.doc_id = doc_id
        self.title = title
        self.subtitle = subtitle
        self.version = version
        self.date = date
        self.author = author
        self.d = Document()
        self._page_setup()
        self._styles()

    # ── setup ────────────────────────────────────────────────────────────────
    def _page_setup(self):
        for s in self.d.sections:
            s.page_width = Inches(8.5)
            s.page_height = Inches(11)
            s.left_margin = Inches(1)
            s.right_margin = Inches(1)
            s.top_margin = Inches(0.9)
            s.bottom_margin = Inches(0.9)

    def _styles(self):
        n = self.d.styles["Normal"]
        n.font.name = "Arial"
        n.font.size = Pt(10)
        rpr = n.element.get_or_add_rPr()
        rf = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), "Arial")
        rpr.append(rf)
        n.paragraph_format.space_after = Pt(6)
        n.paragraph_format.line_spacing = 1.12

        sizes = {"Heading 1": 14, "Heading 2": 12, "Heading 3": 11}
        for name, pt in sizes.items():
            st = self.d.styles[name]
            st.font.name = "Arial"
            st.font.size = Pt(pt)
            st.font.bold = True
            st.font.color.rgb = NAVY
            st.paragraph_format.space_before = Pt(12)
            st.paragraph_format.space_after = Pt(4)

    # ── header / footer ──────────────────────────────────────────────────────
    def _header_footer(self):
        for s in self.d.sections:
            s.different_first_page_header_footer = True

            h = s.header.paragraphs[0]
            h.text = f"{self.doc_id}\tCONFIDENTIAL"
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in h.runs:
                r.font.size = Pt(8)
                r.font.name = "Arial"
                r.font.color.rgb = NAVY
                r.bold = True

            f = s.footer.paragraphs[0]
            f.text = ""
            r = f.add_run(f"{ORG}\t")
            r.font.size = Pt(8)
            r.font.name = "Arial"
            r.font.color.rgb = NAVY
            r2 = f.add_run("Page ")
            r2.font.size = Pt(8)
            r2.font.name = "Arial"
            pr = _field(f, "PAGE")
            pr.font.size = Pt(8)
            pr.font.name = "Arial"
            f.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── building blocks ──────────────────────────────────────────────────────
    def cover(self, volume_line):
        d = self.d
        for _ in range(4):
            d.add_paragraph()
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(ORG.upper())
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = GOLD
        r.font.name = "Arial"

        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(self.title)
        r.font.size = Pt(26)
        r.font.bold = True
        r.font.color.rgb = NAVY
        r.font.name = "Arial"

        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(volume_line)
        r.font.size = Pt(16)
        r.font.color.rgb = BLUE
        r.font.name = "Arial"

        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("_" * 58)
        r.font.color.rgb = GOLD

        for line, sz, bold in ((self.subtitle, 11, False),
                               (f"Contract {CONTRACT}", 11, True),
                               (f"CAGE {CAGE}  |  UEI {UEI}", 10, False),
                               ("", 10, False),
                               (f"Document ID: {self.doc_id}", 10, False),
                               (f"Version {self.version}  |  {self.date}", 10, False)):
            p = d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            r.font.size = Pt(sz)
            r.font.bold = bold
            r.font.name = "Arial"
            r.font.color.rgb = NAVY if bold else RGBColor(0x33, 0x33, 0x33)

        for _ in range(6):
            d.add_paragraph()
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("CONFIDENTIAL — Contains Controlled Unclassified "
                      "Information (CUI)")
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = NAVY
        r.font.name = "Arial"
        self.page_break()

    def doc_control(self, rows):
        self.h1("Document Control")
        self.table(["Field", "Value"], rows, widths=(2.2, 4.3))

    def revision_history(self, rows):
        self.h2("Revision History")
        self.table(["Version", "Date", "Author", "Description"], rows,
                   widths=(0.9, 1.3, 1.7, 2.6))

    def toc(self):
        self.h1("Table of Contents")
        p = self.d.add_paragraph()
        _field(p, r'TOC \o "1-3" \h \z \u')
        note = self.d.add_paragraph()
        r = note.add_run("If the contents above appear as a single line, open in "
                         "Microsoft Word and press Ctrl+A then F9 to populate the "
                         "table of contents.")
        r.font.size = Pt(8)
        r.italic = True
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        self.page_break()

    def page_break(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def h1(self, text):
        self.d.add_heading(text, 1)

    def h2(self, text):
        self.d.add_heading(text, 2)

    def h3(self, text):
        self.d.add_heading(text, 3)

    def p(self, text, bold=False, italic=False, size=10):
        par = self.d.add_paragraph()
        r = par.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Arial"
        return par

    def bullets(self, items):
        for it in items:
            par = self.d.add_paragraph(style="List Bullet")
            r = par.add_run(it)
            r.font.size = Pt(10)
            r.font.name = "Arial"

    def numbered(self, items):
        for it in items:
            par = self.d.add_paragraph(style="List Number")
            r = par.add_run(it)
            r.font.size = Pt(10)
            r.font.name = "Arial"

    def table(self, headers, rows, widths=None, font_size=8.5):
        t = self.d.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, htext in enumerate(headers):
            hdr[i].text = ""
            par = hdr[i].paragraphs[0]
            r = par.add_run(str(htext))
            r.bold = True
            r.font.size = Pt(font_size)
            r.font.name = "Arial"
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade(hdr[i], NAVY_HEX)
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for i in range(len(headers)):
                val = str(row[i]) if i < len(row) else ""
                cells[i].text = ""
                par = cells[i].paragraphs[0]
                r = par.add_run(val)
                r.font.size = Pt(font_size)
                r.font.name = "Arial"
                if ri % 2 == 1:
                    _shade(cells[i], GREY_HEX)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    if i < len(row.cells):
                        row.cells[i].width = Inches(w)
        self.d.add_paragraph()
        return t

    def diagram(self, lines, caption=""):
        """Render a text-based diagram inside a shaded, bordered single cell.

        Monospace (Consolas) is used here and ONLY here. Arial is proportional,
        so box-drawing and arrow alignment collapses under it — the diagram
        would render as ragged text rather than a figure. Body copy stays Arial
        per the branding standard.
        """
        t = self.d.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        c = t.rows[0].cells[0]
        c.text = ""
        par = c.paragraphs[0]
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.line_spacing = 1.0
        for i, line in enumerate(lines):
            r = par.add_run()
            if i:
                r.add_break()          # soft line break inside one paragraph
            r.add_text(line)
            r.font.name = "Consolas"
            r.font.size = Pt(7.5)
            rpr = r._element.get_or_add_rPr()
            rf = OxmlElement("w:rFonts")
            for a in ("w:ascii", "w:hAnsi", "w:cs"):
                rf.set(qn(a), "Consolas")
            rpr.append(rf)
        _shade(c, "F7F9FC")
        if caption:
            cp = self.d.add_paragraph()
            cr = cp.add_run(caption)
            cr.font.size = Pt(8)
            cr.italic = True
            cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            cr.font.name = "Arial"
        self.d.add_paragraph()

    def callout(self, text, label="NOTE"):
        t = self.d.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        c = t.rows[0].cells[0]
        c.text = ""
        par = c.paragraphs[0]
        r = par.add_run(f"{label}: ")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = NAVY
        r.font.name = "Arial"
        r2 = par.add_run(text)
        r2.font.size = Pt(9)
        r2.font.name = "Arial"
        _shade(c, "FFF8E7")
        self.d.add_paragraph()

    def save(self, path):
        self._header_footer()
        self.d.save(str(path))
        return path
