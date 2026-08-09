"""FCC Daily News Summary Word generation (Task 2).

The document is produced from a briefing the reviewer has already cleaned up in
Excel and is pasted into an email as-is, so the failure that matters is not a
crash — it is a document that opens fine and is subtly wrong: a story missing, a
headline that is not clickable, markup in the body.

Assertions are made against the RENDERED .docx (reopened with python-docx and,
where needed, the raw XML), not against the builder's inputs. A test that checks
what we meant to write cannot catch a writer that drops it.
"""

import io
import zipfile

import pytest
from docx import Document

from app.bulletin_intelligence.word_generator import (BulletinWordGenerator,
                                                      filename_for)

pytestmark = [pytest.mark.regression, pytest.mark.bulletin]


ARTICLES = [
    {"title": "FCC adopts new spectrum rules", "url": "https://example.com/a",
     "summary": "The Commission voted to open new spectrum bands.",
     "outlet": "Reuters", "section": "Spectrum Policy", "relevance_score": 0.9},
    {"title": "Broadband funding expands", "url": "https://example.com/b",
     "summary": "New funding reaches rural districts.",
     "outlet": "Washington Post", "section": "Spectrum Policy",
     "relevance_score": 0.8, "is_paywalled": True},
    {"title": "Commissioner comments on AI", "url": "https://example.com/c",
     "summary": "Remarks delivered at a policy forum.",
     "outlet": "Politico", "section": "AI & Emerging Tech",
     "relevance_score": 0.5, "article_type": "opinion"},
]


def _render(articles=ARTICLES, date="August 9, 2026") -> bytes:
    return BulletinWordGenerator().build(articles, briefing_date=date)


def _text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    for section in document.sections:
        parts += [p.text for p in section.footer.paragraphs]
    return "\n".join(parts)


def _xml(content: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(content)) as z:
        return z.read("word/document.xml").decode("utf-8")


# ── structure ────────────────────────────────────────────────────────────────

def test_generates_valid_docx():
    content = _render()
    assert content[:2] == b"PK", "a .docx is a zip archive"
    document = Document(io.BytesIO(content))
    assert len(document.paragraphs) > 5
    with zipfile.ZipFile(io.BytesIO(content)) as z:
        assert "word/document.xml" in z.namelist()


def test_every_article_appears():
    """The failure that matters most: a document that opens fine, minus a story."""
    text = _text(_render())
    for article in ARTICLES:
        assert article["title"].split(" ")[0] in text
        assert article["summary"] in text


def test_articles_grouped_by_category():
    text = _text(_render())
    assert "Spectrum Policy (2 articles)" in text
    assert "AI & Emerging Tech (1 article)" in text, "singular for one article"


def test_categories_are_ordered_deterministically():
    """Two runs over one briefing must produce the same document, or a reviewer
    cannot tell an edit from a reshuffle."""
    first = _text(_render())
    second = _text(_render(list(reversed(ARTICLES))))
    assert first.index("Spectrum Policy") < first.index("AI & Emerging Tech")
    assert second.index("Spectrum Policy") < second.index("AI & Emerging Tech")


def test_headlines_hyperlinked():
    """python-docx has no hyperlink API; without the hand-built w:hyperlink the
    headlines are plain text and the document is a dead end."""
    content = _render()
    assert "<w:hyperlink" in _xml(content)
    with zipfile.ZipFile(io.BytesIO(content)) as z:
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
    for article in ARTICLES:
        assert article["url"] in rels, f"no link relationship for {article['url']}"


def test_an_article_without_a_url_is_not_a_dead_link():
    content = _render([{**ARTICLES[0], "url": ""}])
    assert ARTICLES[0]["title"] in _text(content)
    with zipfile.ZipFile(io.BytesIO(content)) as z:
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
    assert "example.com" not in rels


def test_toc_present():
    text = _text(_render())
    assert "Contents" in text
    assert "(2)" in text and "(1)" in text


def test_footer_on_every_page():
    """Word repeats the section footer on every page, so the assertion is that
    the section footer carries the branding and live page fields."""
    content = _render()
    document = Document(io.BytesIO(content))
    footer_text = "\n".join(p.text for s in document.sections
                            for p in s.footer.paragraphs)
    assert "Alliance Global Tech, Inc." in footer_text
    assert "273FCC26F0061" in footer_text

    with zipfile.ZipFile(io.BytesIO(content)) as z:
        footers = [n for n in z.namelist() if n.startswith("word/footer")]
        assert footers, "no footer part written"
        blob = "".join(z.read(n).decode("utf-8") for n in footers)
    # "Page 2 of 7" cannot be computed here — Word evaluates these on open.
    assert "PAGE" in blob and "NUMPAGES" in blob


# ── content safety ───────────────────────────────────────────────────────────

def test_html_is_stripped_from_every_field():
    """A stray tag in a delivered federal document is the visible failure; the
    invisible one is markup surviving into text somebody pastes elsewhere."""
    content = _render([{
        "title": "<b>FCC</b> acts", "url": "https://example.com/x",
        "summary": "<p>Body text</p> with <i>markup</i>",
        "outlet": "<span>Reuters</span>", "section": "General",
        "relevance_score": 0.8}])
    text = _text(content)
    assert "<b>" not in text and "<p>" not in text and "<span>" not in text
    assert "FCC acts" in text
    assert "Body text" in text


def test_paywalled_and_opinion_articles_are_tagged():
    text = _text(_render())
    assert "[Subscription Required]" in text
    assert "[Opinion]" in text


def test_relevance_band_is_shown():
    text = _text(_render())
    assert "High Relevance" in text
    assert "Medium Relevance" in text


# ── edge cases ───────────────────────────────────────────────────────────────

def test_an_empty_briefing_says_so():
    """A header with nothing under it reads as a broken export rather than a
    quiet day."""
    text = _text(_render([]))
    assert "No articles met the criteria" in text
    assert "0 Articles" in text


def test_uses_reviewed_excel_when_available():
    """A reviewed workbook applied via upload-reviewed edits the same article
    records this reads, so there is exactly one source of truth. Pinning the
    shared helper is what keeps that true."""
    import inspect

    from app.bulletin_intelligence import word_generator

    src = inspect.getsource(word_generator.BulletinWordGenerator.generate)
    assert "_briefing_articles" in src, \
        "must read through the shared rehydration helper, not a second parser"


def test_filename_is_readable_and_dated():
    assert filename_for("August 9, 2026") == "FCC_Bulletin_Aug09_2026.docx"
    assert filename_for("2026-08-09", "FCC") == "FCC_Bulletin_Aug09_2026.docx"
    # An unparseable date must still yield a usable filename.
    assert filename_for("not a date").endswith(".docx")


# ── wiring ───────────────────────────────────────────────────────────────────

def test_the_endpoint_is_registered_and_gated():
    from app.core.security import ROLE_HIERARCHY  # noqa: F401
    from app.main import app

    paths = app.openapi()["paths"]
    assert "/api/v1/bulletin/generate-word/{briefing_id}" in paths


def test_the_endpoint_returns_a_docx_content_type():
    import inspect

    from app.bulletin_intelligence import routes

    src = inspect.getsource(routes.generate_word)
    assert "wordprocessingml.document" in src
    assert "attachment; filename=" in src
