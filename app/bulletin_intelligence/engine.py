"""
DocuAction Bulletin Intelligence — Multi-Source Intelligence Engine v2
News Sources:
  - GDELT Project API    FREE  — global news every 15 min, no key needed
  - Tavily AI Search     FREE  — 1,000/mo free, AI-optimized (TAVILY_API_KEY)
  - NewsAPI              FREE  — dev tier, 80K+ sources (NEWSAPI_KEY)
  - Claude web_search    —     — fallback, uses existing ANTHROPIC_API_KEY
  - Federal Register     FREE  — no key needed
  - Congress.gov         FREE  — CONGRESS_API_KEY (optional)

LLM Visibility Tracker (unique — no competitor has this):
  - Claude (Anthropic)   existing ANTHROPIC_API_KEY
  - ChatGPT (OpenAI)     OPENAI_API_KEY
  - Perplexity           PERPLEXITY_API_KEY (sonar has real-time web)
  - Gemini (Google)      GEMINI_API_KEY

Classification:  Claude Haiku
Briefing:        Claude Sonnet → FCC Daily News Summary format

Required env vars:
  ANTHROPIC_API_KEY      already set ✓
  SENDGRID_API_KEY       set in Railway
  TAVILY_API_KEY         free at tavily.com
  NEWSAPI_KEY            free at newsapi.org
  OPENAI_API_KEY         for ChatGPT LLM visibility
  PERPLEXITY_API_KEY     for Perplexity LLM visibility
  GEMINI_API_KEY         for Gemini LLM visibility
  CONGRESS_API_KEY       free at api.congress.gov
"""

import os, json, logging, asyncio, hashlib, re
from datetime import datetime, timedelta, timezone
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, asdict, field

import httpx
from anthropic import AsyncAnthropic

# Extended feed + keyword lists (ADD-only companion modules; optional at runtime).
try:
    from app.bulletin_intelligence.fcc_feeds_extended import EXTENDED_FCC_OFFICIAL, EXTENDED_OUTLET_FEEDS
    from app.bulletin_intelligence.fcc_keywords_extended import EXTENDED_FCC_KEYWORDS, EXTENDED_FCC_OFFICIALS, EXTENDED_FCC_CORE_PHRASES
except ImportError:
    EXTENDED_FCC_OFFICIAL = []
    EXTENDED_OUTLET_FEEDS = []
    EXTENDED_FCC_KEYWORDS = []
    EXTENDED_FCC_OFFICIALS = []
    EXTENDED_FCC_CORE_PHRASES = []

logger = logging.getLogger(__name__)

ANTHROPIC_KEY   = os.getenv("ANTHROPIC_API_KEY", "")
SENDGRID_KEY    = os.getenv("SENDGRID_API_KEY", "")
CONGRESS_KEY    = os.getenv("CONGRESS_API_KEY", "")
TAVILY_KEY      = os.getenv("TAVILY_API_KEY", "")
NEWSAPI_KEY     = os.getenv("NEWSAPI_KEY", "")
# NewsAPI.ai (Event Registry) — additive collector. Auto-detected: present = on,
# absent = skipped gracefully (no crash). Distinct from NEWSAPI_KEY (newsapi.org).
NEWSAPI_AI_KEY  = os.getenv("NEWSAPI_AI_KEY", "")
OPENAI_KEY      = os.getenv("OPENAI_API_KEY", "")
PERPLEXITY_KEY  = os.getenv("PERPLEXITY_API_KEY", "")
GEMINI_KEY      = os.getenv("GEMINI_API_KEY", "")

TIMEOUT = httpx.Timeout(30.0)
HTTP_HEADERS = {"User-Agent": "DocuAction-BulletinIntelligence/1.0 (Alliance Global Tech)"}
# Fallback UA for hosts (notably fcc.gov) that block non-browser/cloud clients with
# a 403 or a hung connection. Only used as a retry when the default UA fails.
_BROWSER_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )
}

# ── Topic taxonomy ─────────────────────────────────────────────────────────────
FCC_TOPICS = [
    "fcc_news_events", "consumers_advocacy", "media_broadcasting",
    "public_safety_emergency", "wireless_mobile", "ai_emerging_tech",
    "business_industry", "international_affairs", "space_communications", "spectrum_policy",
]

TOPIC_LABELS = {
    "fcc_news_events":         "FCC News & Events",
    "consumers_advocacy":      "Consumers & Advocacy",
    "media_broadcasting":      "Media & Broadcasting",
    "public_safety_emergency": "Public Safety",
    "wireless_mobile":         "Wireless & Mobile",
    "ai_emerging_tech":        "AI & Emerging Tech",
    "business_industry":       "Business & Industry",
    "international_affairs":   "International Affairs",
    "space_communications":    "Space Communications",
    "spectrum_policy":         "Spectrum & Policy",
    "other":                   "Other",
}
FCC_TOPIC_LABELS = TOPIC_LABELS  # alias for backward compatibility

# ── Client display sections (the 6 buckets in the AGT FCC Daily News email) ─────
# These are what the CLIENT sees, grouped from the finer internal topics above.
# Order here is the order they appear in the briefing.
AGT_SECTIONS = [
    "General",            # client's Appendix A label for FCC_NEWS; always first
    "Consumers",
    "Media & Broadcasting",
    "Public Safety / Cybersecurity / Privacy",
    "Wireless & Spectrum",
    "Broadband & Infrastructure",   # added 2026-07-02 (FCC-org structure)
    "Space Policy",
    "Business & Tech",
    "Enforcement & Consumer",        # added 2026-07-02 (FCC-org structure)
    "AI / Machine Learning",
    "International",
]

# Fallback topic -> section map (used if the classifier doesn't set a section).
TOPIC_TO_SECTION = {
    "fcc_news_events":          "General",
    "consumers_advocacy":       "Consumers",
    "media_broadcasting":       "Media & Broadcasting",
    "public_safety_emergency":  "Public Safety / Cybersecurity / Privacy",
    "wireless_mobile":          "Wireless & Spectrum",
    "ai_emerging_tech":         "AI / Machine Learning",
    "business_industry":        "Business & Tech",
    "international_affairs":     "International",
    "space_communications":     "Space Policy",
    "spectrum_policy":          "Wireless & Spectrum",
    # extra topic names the classifier sometimes emits — map them too so nothing drops
    "broadband_infrastructure": "Broadband & Infrastructure",
    "enforcement":              "Enforcement & Consumer",
    "equipment_authorization":  "General",
    "cybersecurity_privacy":    "Public Safety / Cybersecurity / Privacy",
    "other":                    "General",
}


# ── Data classes ───────────────────────────────────────────────────────────────
@dataclass
class AgencyConfig:
    agency_id: str
    name: str
    short_name: str
    primary_color: str
    search_queries: List[str]
    topics: List[str]
    distribution_email: str
    distribution_list: List[str]
    delivery_time_et: str = "07:30"
    include_broadcast: bool = True
    include_social: bool = True
    include_regulatory: bool = True
    archive_months: int = 12
    logo_url: str = ""         # agency's own logo (e.g. FCC seal) shown in the header


@dataclass
class Article:
    article_id: str
    agency_id: str
    source: str
    source_type: str
    title: str
    url: str
    published_at: str
    summary: str
    full_text: str
    author: str
    outlet: str
    topic: str = "other"
    section: str = ""          # client display section (one of AGT_SECTIONS)
    article_type: str = "news"
    relevance_score: float = 0.5
    sentiment: str = "neutral"
    is_paywalled: bool = False
    broadcast_clip_url: str = ""
    ingested_at: str = ""
    dedup_hash: str = ""
    # ── Provider tracking (additive; defaults keep older stored rows valid) ──────
    # Which collector surfaced this article, so per-provider analytics + the
    # provider column in exports can be computed from real data (never guessed).
    provider: str = ""            # e.g. "NewsAPI.ai", "GDELT", "RSS", "Federal Register"
    provider_url: str = ""        # provider/API canonical URL
    source_name: str = ""         # human outlet name (mirrors `outlet` when unset)
    collection_method: str = ""   # rss | news_api | search_api | news_index | gov_api
    collection_time: str = ""     # ISO-8601 UTC time this article was collected


@dataclass
class Briefing:
    briefing_id: str
    agency_id: str
    briefing_date: str
    status: str
    html_content: str
    article_count: int
    topic_counts: Dict[str, int] = field(default_factory=dict)
    generated_at: str = ""
    approved_at: str = ""
    delivered_at: str = ""
    delivery_recipients: int = 0
    docx_b64: str = ""         # editable Word version, base64 (for the download endpoint)


# ── In-memory store ────────────────────────────────────────────────────────────
_articles:  Dict[str, Article]  = {}
_briefings: Dict[str, Briefing] = {}
_agencies:  Dict[str, AgencyConfig] = {}

# Cost guard: prevents overlapping/duplicate ingest cycles (each cycle makes
# ~20+ Claude calls). Maps agency_id -> start timestamp; a cycle is considered
# in-progress while a recent timestamp is present. Auto-expires after the TTL so
# a crashed cycle can never permanently block future runs.
_running_cycles: Dict[str, float] = {}
CYCLE_LOCK_TTL = 180  # seconds (covers a typical 1-2 min cycle)


def _hash(url: str, title: str) -> str:
    return hashlib.md5(f"{url}{title}".encode()).hexdigest()


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _us_date(dt) -> str:
    """US long date with NO leading zero on the day: 'July 7, 2026' (never
    'July 07, 2026' or '7 July 2026'). Portable across OSes (avoids the
    non-portable %-d / %#d strftime flags)."""
    try:
        return f"{dt.strftime('%B')} {dt.day}, {dt.year}"
    except Exception:
        return str(dt)


def _us_date_short(dt) -> str:
    """US month+day, no year, no leading zero: 'July 7' (for range starts)."""
    try:
        return f"{dt.strftime('%B')} {dt.day}"
    except Exception:
        return str(dt)


def _parse_pub_dt(s: str) -> Optional[datetime]:
    """Parse a feed publish date to an aware UTC datetime, or None if it cannot be
    parsed. Handles ISO-8601 (Atom), GDELT (YYYYMMDDTHHMMSSZ), Unix epoch, and
    RFC-2822 (RSS).

    Returns None (NOT now()) on failure so callers can treat an unverifiable date
    as STALE rather than silently fresh. This is the crux of the stale-article fix:
    Atom/ISO dates used to fail the RFC-2822-only parser and get stamped now(),
    letting weeks-old items pass every freshness cutoff."""
    s = (s or "").strip()
    if not s:
        return None

    def _aware(dt: datetime) -> datetime:
        return dt if dt.tzinfo else dt.replace(tzinfo=timezone.utc)

    # ISO-8601 (Atom feeds, and our own stored dates). Handles date-only too.
    try:
        return _aware(datetime.fromisoformat(s.replace("Z", "+00:00")))
    except Exception:
        pass
    # GDELT compact form.
    try:
        return datetime.strptime(s, "%Y%m%dT%H%M%SZ").replace(tzinfo=timezone.utc)
    except Exception:
        pass
    # Unix epoch seconds (Reddit created_utc), guarded to year 2001+ to avoid
    # misreading short numeric strings as dates.
    try:
        v = float(s)
        if v >= 1_000_000_000:
            return datetime.fromtimestamp(v, tz=timezone.utc)
    except Exception:
        pass
    # RFC-2822 (RSS pubDate).
    try:
        from email.utils import parsedate_to_datetime
        return _aware(parsedate_to_datetime(s))
    except Exception:
        return None


def _normalize_pub(s: str) -> str:
    """Normalize a publish date to ISO 8601 for storage/display. Falls back to
    now() when unparseable (fine for display); callers that must gate on FRESHNESS
    should use _parse_pub_dt and treat None as stale."""
    dt = _parse_pub_dt(s)
    return dt.isoformat() if dt else _now()


# FCC operates on US Eastern time; EST/EDT is handled automatically by the tz db
# (never hardcode an offset). Same mechanism the scheduler uses.
try:
    from zoneinfo import ZoneInfo   # Python 3.9+ stdlib
    _ET = ZoneInfo("America/New_York")
except Exception:                    # pragma: no cover
    import pytz
    _ET = pytz.timezone("America/New_York")


def get_briefing_window():
    """Freshness window for a briefing: the PREVIOUS BUSINESS DAY(S) in US Eastern
    time, not a rolling 72h. Returns (start, end) as aware ET datetimes; an article
    is in-window iff start <= published < end (end = last midnight ET, so today's
    items are excluded).

      Tue-Sat  -> previous day only            (1 day)
      Monday   -> Fri 00:00 .. Mon 00:00 ET     (Fri+Sat+Sun)
      Sunday   -> Fri 00:00 .. Sat 00:00 ET     (shouldn't run)

    Comparisons are between timezone-AWARE datetimes, so a UTC article date is
    compared as the same instant — no manual offset conversion needed."""
    now = datetime.now(_ET)
    today = now.date()
    weekday = today.weekday()  # 0=Monday
    if weekday == 0:            # Monday — cover Fri+Sat+Sun
        start = today - timedelta(days=3)
    elif weekday == 6:         # Sunday — shouldn't run
        start = today - timedelta(days=2)
    else:                      # Tue-Sat — previous day only
        start = today - timedelta(days=1)
    end = today                # up to midnight last night ET
    return (
        datetime.combine(start, datetime.min.time(), tzinfo=_ET),
        datetime.combine(end, datetime.min.time(), tzinfo=_ET),
    )


# Last briefing window + in/out-of-window counts per agency, surfaced in the
# run_daily_cycle result so a run can report exactly what the window admitted.
_last_window_stats: Dict[str, Dict[str, Any]] = {}


def _dict_to_article(d: dict, agency_id: str, default_source_type: str = "news") -> Optional["Article"]:
    """Convert a raw ingester dict (GDELT, Reddit, etc.) into an Article. Returns
    None for unusable rows (no url/title)."""
    url = (d.get("url") or "").strip()
    title = (d.get("title") or "").strip()
    if not url or not title:
        return None
    summary = (d.get("summary") or "").strip() or title
    outlet = d.get("outlet") or d.get("source") or ""
    return Article(
        article_id=f"{agency_id}_{_hash(url, title)}",
        agency_id=agency_id,
        source=d.get("source", "") or default_source_type,
        source_type=d.get("source_type", default_source_type),
        title=title,
        url=url,
        published_at=_normalize_pub(d.get("published_at", "")),
        summary=summary,
        full_text=d.get("full_text", "") or summary,
        author=d.get("author", ""),
        outlet=outlet,
        ingested_at=_now(),
        dedup_hash=_hash(url, title),
    )


async def hydrate_from_store() -> Dict[str, int]:
    """Restore persisted articles/briefings into the in-memory cache on startup.

    Best-effort: if the store is empty or unavailable, the cache simply stays
    empty and the app runs as before.
    """
    try:
        from dataclasses import fields as _fields
        from app.bulletin_intelligence import bulletin_store
        articles, briefings = await bulletin_store.load_all()
        art_fields = {f.name for f in _fields(Article)}
        brief_fields = {f.name for f in _fields(Briefing)}
        for d in articles:
            try:
                _articles[d["article_id"]] = Article(**{k: v for k, v in d.items() if k in art_fields})
            except Exception:
                pass
        for d in briefings:
            try:
                _briefings[d["briefing_id"]] = Briefing(**{k: v for k, v in d.items() if k in brief_fields})
            except Exception:
                pass
        # ── One-time migration to the LIVE-FEED model ──────────────────────────
        # The old flow parked briefings in "pending_approval" behind an approval
        # gate that we've removed. Any such briefing is really just a live briefing
        # that never got clicked through — mark it "delivered" (i.e. live/viewable)
        # so it shows up in latest/today/history immediately. Idempotent: only
        # touches pending_approval rows, and re-persists just those.
        migrated = 0
        for b in _briefings.values():
            if b.status == "pending_approval":
                b.status = "delivered"
                migrated += 1
                try:
                    await bulletin_store.save_briefing(asdict(b))
                except Exception as e:
                    logger.warning(f"Migration persist failed for {b.briefing_id}: {e}")
        if migrated:
            logger.info(f"Bulletin hydrate: migrated {migrated} pending_approval -> delivered (live-feed model)")
        logger.info(f"Bulletin hydrate: {len(_articles)} articles, {len(_briefings)} briefings restored")
        return {"articles": len(_articles), "briefings": len(_briefings), "migrated_to_live": migrated}
    except Exception as e:
        logger.warning(f"Bulletin hydrate failed: {e}")
        return {"articles": 0, "briefings": 0}


# ── Agency management ──────────────────────────────────────────────────────────
def register_agency(config: AgencyConfig) -> None:
    _agencies[config.agency_id] = config
    logger.info(f"Bulletin: registered agency {config.name}")


def get_agency(agency_id: str) -> Optional[AgencyConfig]:
    return _agencies.get(agency_id)


def list_agencies() -> List[AgencyConfig]:
    return list(_agencies.values())


# ── Claude API helper ──────────────────────────────────────────────────────────
def _get_client() -> AsyncAnthropic:
    return AsyncAnthropic(api_key=ANTHROPIC_KEY)


async def _record_llm_cost(resp, *, operation: str, model: str) -> None:
    """Phase 1 cost tracking shim — records tokens/cost for one Claude response.

    Imported lazily and fully swallowed so that neither a missing costs package nor
    a DB problem can affect a bulletin run. No-op unless
    BULLETIN_COST_TRACKING_ENABLED=true. Records only; changes no behaviour.
    """
    try:
        from app.bulletin_intelligence.costs.cost_tracker import record_usage
        await record_usage(resp, operation=operation, model=model)
    except Exception:
        pass


def _extract_text(content) -> str:
    """Extract plain text from Anthropic message content blocks."""
    parts = []
    for block in (content if isinstance(content, list) else [content]):
        if hasattr(block, 'type'):
            if block.type == 'text':
                parts.append(block.text)
        elif isinstance(block, dict) and block.get('type') == 'text':
            parts.append(block.get('text', ''))
    return '\n'.join(parts)


def _parse_json_safe(text: str) -> Any:
    """Parse JSON from Claude response — handles markdown code fences."""
    text = text.strip()
    text = re.sub(r'^```json\s*', '', text)
    text = re.sub(r'^```\s*', '', text)
    text = re.sub(r'\s*```$', '', text)
    text = text.strip()

    # Find first [ or { and last ] or }
    start = min(
        (text.find('[') if '[' in text else len(text)),
        (text.find('{') if '{' in text else len(text))
    )
    end_bracket = text.rfind(']')
    end_brace = text.rfind('}')
    end = max(end_bracket, end_brace)

    if start <= end:
        text = text[start:end + 1]

    return json.loads(text)


# ── FCC Relevance Pre-Filter ──────────────────────────────────────────────────
FCC_KEYWORDS = {
    "fcc", "federal communications commission", "fcc chairman", "fcc commissioner",
    "brendan carr", "olivia trusty", "anna gomez",
    "spectrum auction", "spectrum license", "aws-3", "spectrum policy",
    "robocall", "tcpa", "stir-shaken", "robocall mitigation",
    "net neutrality", "open internet", "e-rate", "lifeline program",
    "media ownership", "broadcast license", "radio license", "tv license",
    "submarine cable", "undersea cable", "subsea cable",
    "911 fcc", "e911", "psap", "emergency alert system",
    "fcc enforcement", "fcc fine", "fcc ruling", "fcc vote", "fcc proposes",
    "fcc approves", "fcc order", "fcc notice", "fcc meeting",
    "telecom policy", "telecommunications policy", "fcc regulation",
}

def _is_fcc_relevant(title: str, summary: str) -> bool:
    """Return True only if article is genuinely FCC-relevant."""
    text = (title + " " + summary).lower()
    return any(kw in text for kw in FCC_KEYWORDS)




# ── INGESTION: RSS Feeds (Appendix B Sources — FREE, always FCC-relevant) ─────
# Outlets the client has asked us NOT to include (not relevant to the FCC
# mission). Matched as a substring of the article URL, so "techdirt.com" also
# covers www.techdirt.com. Applied both at ingestion (new articles never enter)
# and at download/render time (so already-stored articles are hidden too).
EXCLUDED_DOMAINS = {
    "techdirt.com",
}


def _is_excluded_domain(url: str) -> bool:
    u = (url or "").lower()
    return any(dom in u for dom in EXCLUDED_DOMAINS)


# Terms that mark genuine FCC relevance — used to filter the broad major-outlet
# feeds down to FCC stories only, so we add their coverage without flooding the
# briefing with unrelated national news.
_FCC_RELEVANCE_TERMS = (
    "fcc", "federal communications commission", "f.c.c.",
    "brendan carr", "anna gomez", "olivia trusty", "geoffrey starks", "nathan simington",
    # Expanded 2026-06-29 for broader FCC coverage from Google News / broad outlet
    # feeds. Substring-matched (safe because each is specific enough on its own).
    "lifeline", "affordable connectivity", "universal service fund",
    "rip and replace", "huawei", "spectrum auction", "space bureau",
    "ngso", "earth station", "data breach", "spectrum pipeline", "digital equity",
    "e-rate", "robocall", "spoofing", "media ownership", "broadcast",
    "tower siting", "small cell", "atsc", "nextgen tv", "emergency alert",
    "tcpa", "pole attachment", "cable franchise",
    # Expanded 2026-06-30 (URGENT volume push) — broaden the gate so more FCC-adjacent
    # stories from the broad outlet / Google News feeds pass. ADD-only; existing kept.
    "spectrum", "broadband", "telecom", "telecommunications", "wireless", "5g", "6g",
    "net neutrality", "open internet", "universal service", "pirate radio",
    "enforcement action", "section 230", "content moderation", "ng911", "satellite",
    "starlink", "spacex", "commissioner", "gomez", "c-band", "cbrs",
)

# Short abbreviations matched as WHOLE WORDS only — naive substring matching would
# false-positive (e.g. "eas" in "release"/"season", "acp" in "backpack", "bead" in
# "beads", "usf" in "useful", "carr" in "carrier"/"carry") and effectively disable
# the gate. Same terms the client asked for (EAS, ACP, USF, BEAD, ZTE, Carr),
# just boundary-safe.
_FCC_RELEVANCE_WORD_TERMS = ("eas", "acp", "usf", "bead", "zte", "carr")
_FCC_WORD_RE = re.compile(r"\b(?:" + "|".join(_FCC_RELEVANCE_WORD_TERMS) + r")\b")

# Merge the client's extended keyword/official lists into the substring gate — ADD
# only, existing kept. Skip short/ambiguous tokens (<4 chars, e.g. "usf", "fcc")
# that are already handled safely as whole-word terms above; substring-matching them
# would false-positive and effectively disable the gate.
_EXTENDED_GATE_TERMS = tuple(sorted({
    k.lower().strip()
    for k in (list(EXTENDED_FCC_KEYWORDS) + list(EXTENDED_FCC_OFFICIALS))
    if len(k.strip()) >= 4 and k.lower().strip() not in _FCC_RELEVANCE_TERMS
}))
if _EXTENDED_GATE_TERMS:
    _FCC_RELEVANCE_TERMS = _FCC_RELEVANCE_TERMS + _EXTENDED_GATE_TERMS


def _mentions_fcc_legacy(text: str) -> bool:
    """Flat substring/word gate. Superseded by _is_fcc_relevant_v2; kept for
    rollback (flip the call sites back to this name)."""
    t = (text or "").lower()
    if any(term in t for term in _FCC_RELEVANCE_TERMS):
        return True
    return bool(_FCC_WORD_RE.search(t))


# ── 3-tier FCC relevance filter (2026-07-02) ─────────────────────────────────
# Replaces the flat gate above. Generic tech words (broadband, 5g, wireless,
# satellite) no longer pass on their own — they need an explicit FCC signal —
# which cuts the general-tech noise that made the briefing read like an RSS dump.
_FCC_TIER1_EXPLICIT = (
    "federal communications commission",
    "fcc chairman", "fcc commissioner", "fcc chairwoman",
    "fcc ruling", "fcc order", "fcc vote", "fcc fine", "fcc enforcement",
    "fcc license", "fcc approval", "fcc regulation", "fcc docket", "fcc action",
    "fcc proposes", "fcc approves", "fcc adopts", "fcc notice", "fcc proposal",
)
# Bare 'fcc' / 'f.c.c.' as a whole word (avoids matching inside other tokens).
_FCC_TIER1_RE = re.compile(r"\bf\.?c\.?c\.?\b")

_FCC_TIER2_OFFICIALS = (
    "brendan carr", "chairman carr", "olivia trusty", "anna gomez",
    "commissioner gomez", "commissioner trusty", "commissioner carr",
    "wireline competition bureau", "wireless telecommunications bureau",
    "media bureau", "enforcement bureau", "public safety bureau",
    "space bureau", "office of economics and analytics",
    "consumer and governmental affairs bureau", "international bureau",
)

# Tier 3 must appear in the TITLE (not just the summary) to pass — these are
# FCC-specific proceedings, so a title match is a strong on-topic signal.
_FCC_TIER3_PROGRAMS_TITLE = (
    "e-rate", "lifeline program", "connect america fund",
    "robocall mitigation database", "rip and replace", "covered list",
    "stir/shaken", "stir-shaken", "universal service fund", "net neutrality",
    "spectrum auction", "retransmission consent", "next generation 911",
    "ng911", "pirate radio", "pole attachment", "huawei ban", "zte ban",
)


def _is_fcc_relevant_v2(title: str, summary: str = "") -> bool:
    """3-tier FCC relevance gate.
      Tier 1  explicit FCC mention (title OR summary) -> pass
      Tier 2  named FCC official / bureau (title OR summary) -> pass
      Tier 3  FCC-specific program/proceeding, TITLE ONLY -> pass
    Generic tech terms (broadband/5g/wireless/satellite) never pass alone."""
    t = (title or "").lower()
    blob = f"{title} {summary}".lower()
    if _FCC_TIER1_RE.search(blob) or any(term in blob for term in _FCC_TIER1_EXPLICIT):
        return True
    if any(term in blob for term in _FCC_TIER2_OFFICIALS):
        return True
    if any(term in t for term in _FCC_TIER3_PROGRAMS_TITLE):
        return True
    return False


# Major national outlets the client wants represented on big FCC stories (e.g. the
# WaPo/NYT/WSJ "ABC vs. FCC / The View" coverage). Their RSS section feeds carry
# mostly non-FCC news, so items are kept ONLY when they mention the FCC — see the
# relevance_required path in ingest_rss. Format: (feed_url, outlet, paywalled).
MAJOR_OUTLET_FEEDS = [
    ("https://rss.nytimes.com/services/xml/rss/nyt/MediaandAdvertising.xml", "The New York Times", True),
    ("https://rss.nytimes.com/services/xml/rss/nyt/Technology.xml",          "The New York Times", True),
    ("https://rss.nytimes.com/services/xml/rss/nyt/Business.xml",            "The New York Times", True),
    ("https://rss.nytimes.com/services/xml/rss/nyt/Politics.xml",            "The New York Times", True),
    ("https://feeds.washingtonpost.com/rss/business/technology",             "The Washington Post", True),
    ("https://feeds.washingtonpost.com/rss/business",                        "The Washington Post", True),
    ("https://feeds.washingtonpost.com/rss/politics",                        "The Washington Post", True),
    ("https://feeds.a.dj.com/rss/RSSWSJD.xml",                               "The Wall Street Journal", True),
    ("https://feeds.a.dj.com/rss/WSJcomUSBusiness.xml",                      "The Wall Street Journal", True),
    ("https://thehill.com/homenews/media/feed/",                            "The Hill", False),
    # National tech press (Tier 2) — broad, so FCC-relevance filtered.
    ("https://api.axios.com/feed/",                                          "Axios", False),
    ("https://www.theverge.com/rss/index.xml",                              "The Verge", False),
    ("https://feeds.arstechnica.com/arstechnica/index",                     "Ars Technica", False),
    ("https://search.cnbc.com/rs/search/combinedcms/view.xml?partnerId=wrss01&id=19854910", "CNBC", False),
    # Telecom/broadcast trade press (Tier 1).
    ("https://www.lightreading.com/rss.xml",                                "Light Reading", True),
    ("https://www.cablefax.com/feed",                                       "Cablefax", False),
    ("https://wirelessestimator.com/feed/",                                 "Wireless Estimator", False),
    ("https://tvnewscheck.com/feed/",                                       "TVNewsCheck", False),
    # Regulatory + industry associations (Tier 3/4) — FCC-relevance filtered.
    ("https://www.ftc.gov/feeds/press-release.xml",                         "FTC", False),
    ("https://www.gao.gov/rss/reports.xml",                                 "GAO", False),
    ("https://www.ustelecom.org/feed/",                                     "USTelecom", False),
    # Legal / FCC-bar (free RSS) — broad, so FCC-relevance filtered. Closes the
    # "no legal coverage" gap without any paid feed (Law360/Comm Daily stay paid).
    ("https://www.commlawblog.com/feed/",                                   "CommLawBlog", False),
    # National tech press (free RSS) — FCC-relevance filtered.
    ("https://techcrunch.com/feed/",                                        "TechCrunch", False),
    ("https://www.wired.com/feed/rss",                                      "Wired", False),
    ("https://www.engadget.com/rss.xml",                                    "Engadget", False),
    ("https://www.cnet.com/rss/news/",                                      "CNET", False),
    ("https://feeds.npr.org/1019/rss.xml",                                  "NPR", False),
    # Added 2026-06-29 for higher volume — broad sources, so FCC-relevance gated
    # (same _mentions_fcc filter as the other major-outlet feeds; non-FCC items dropped).
    ("https://www.techmeme.com/feed.xml",                                   "Techmeme", False),
    ("https://news.google.com/rss/search?q=FCC+OR+%22Federal+Communications+Commission%22&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+broadband+OR+spectrum+OR+telecommunications&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+5G+OR+wireless+OR+spectrum+auction&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=%22Federal+Communications+Commission%22+regulation+OR+policy&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+commissioner+OR+Carr+OR+Starks&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://hnrss.org/newest?q=FCC",                                      "Hacker News", False),
    # Broadcast/radio trade (free RSS) — FCC-relevance filtered.
    ("https://radioink.com/feed/",                                          "Radio Ink", False),
    ("https://current.org/feed/",                                          "Current", False),
    # Inside Radio (UAT missing-source fix): its direct feed is 429/bot-blocked and
    # paywalled, so there is no reliable free feed. Add a FCC-gated Google News
    # site-scoped fallback so its FCC stories can still surface (metadata/headline).
    ("https://news.google.com/rss/search?q=site:insideradio.com+FCC+OR+%22Federal+Communications+Commission%22&hl=en-US&gl=US&ceid=US:en", "Inside Radio", False),
    # (Politico is already covered by politicopicks in FCC_RSS_FEEDS; its
    #  topic feeds 403 bot traffic, so they're not added here. NTIA / White House /
    #  House E&C / Senate Commerce / NAB / NCTA feeds were dead or empty on check.)
    # ── Added 2026-06-30 (URGENT volume push → target 100+/day) ────────────────────
    # All gated (FCC-relevance filtered). Skipped as already present above / in
    # FCC_RSS_FEEDS: Fierce Wireless, CommLawBlog, TVNewsCheck, RadioInk, Radio World,
    # Current, GAO, The Hill Tech, WaPo Tech, NYT Tech.
    # Google News boolean searches (free, no key, high volume).
    ("https://news.google.com/rss/search?q=%22Federal+Communications+Commission%22+OR+FCC+regulation+OR+policy&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+spectrum+OR+auction+OR+5G+OR+wireless&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+broadband+OR+%22digital+equity%22+OR+BEAD&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+enforcement+OR+fine+OR+forfeiture+OR+%22pirate+radio%22&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+%22net+neutrality%22+OR+%22open+internet%22+OR+%22Section+230%22&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+broadcast+OR+%22media+ownership%22+OR+television+OR+radio+license&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+satellite+OR+%22Space+Bureau%22+OR+SpaceX+OR+Starlink+OR+NGSO&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+robocall+OR+TCPA+OR+spoofing+OR+%22consumer+protection%22&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+Carr+OR+Gomez+commissioner&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+%22E-Rate%22+OR+%22universal+service%22+OR+USF+OR+Lifeline&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+%22emergency+alert%22+OR+EAS+OR+NG911+OR+%22public+safety%22&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    ("https://news.google.com/rss/search?q=FCC+AI+OR+%22artificial+intelligence%22+OR+%22tower+siting%22+OR+%22pole+attachment%22&hl=en-US&gl=US&ceid=US:en", "Google News", False),
    # Telecom trade publications.
    ("https://www.telecompaper.com/rss",                                    "Telecompaper", False),
    ("https://www.telecoms.com/feed",                                       "Telecoms.com", False),
    ("https://www.totaltele.com/feed",                                      "Total Telecom", False),
    ("https://www.mobileworldlive.com/feed",                               "Mobile World Live", False),
    ("https://www.telecomramblings.com/feed",                              "Telecom Ramblings", False),
    # Policy / advocacy.
    ("https://www.benton.org/rss",                                          "Benton Institute", False),
    ("https://publicknowledge.org/feed/",                                  "Public Knowledge", False),
    ("https://www.freepress.net/feed",                                      "Free Press", False),
    ("https://www.eff.org/rss/updates.xml",                                 "EFF", False),
    # Broadcasting.
    ("https://www.mediapost.com/publications/feed/",                       "MediaPost", False),
    ("https://thedesk.net/feed/",                                           "The Desk", False),
    ("https://www.mediaite.com/feed/",                                      "Mediaite", False),
    # Satellite / space.
    ("https://spacenews.com/feed/",                                         "SpaceNews", False),
    ("https://www.satellitetoday.com/feed/",                               "Via Satellite", False),
    # Mainstream (high volume, FCC-gated).
    ("https://thehill.com/regulation/feed/",                               "The Hill Regulation", False),
    ("https://rss.politico.com/morningtech.xml",                           "Politico Morning Tech", False),
    ("https://feeds.reuters.com/reuters/technologyNews",                   "Reuters Tech", False),
    ("https://www.cnbc.com/id/19854910/device/rss/rss.html",               "CNBC Tech", False),
    # Government / regulatory.
    ("https://www.federalregister.gov/api/v1/documents.rss?conditions%5Bagencies%5D%5B%5D=federal-communications-commission", "Federal Register FCC", False),
    ("https://www.congress.gov/rss/search-results.xml?query=%7B%22source%22%3A%22all%22%2C%22search%22%3A%22Federal+Communications+Commission%22%7D", "Congress.gov FCC", False),
    # State / local.
    ("https://www.route-fifty.com/rss/technology/",                        "Route Fifty Tech", False),
]


# ── Deterministic FCC relevance scoring + categorization (client spec) ─────────
# These AUGMENT the LLM classifier — they never replace it. The point score nudges
# relevance up for clear FCC signals; the strict gate is OFF by default so the
# bulletin's minimum-volume floor is never at risk.
FCC_COMMISSIONERS = ("brendan carr", "anna gomez", "olivia trusty")

# Flip BULLETIN_STRICT_FCC_GATE=true to hard-reject any story with no explicit FCC
# mention. Default OFF (current behavior preserved).
STRICT_FCC_GATE = os.getenv("BULLETIN_STRICT_FCC_GATE", "false").strip().lower() == "true"


def _fcc_blob(art) -> str:
    return " ".join([
        getattr(art, "title", "") or "",
        getattr(art, "summary", "") or "",
        getattr(art, "full_text", "") or "",
        getattr(art, "outlet", "") or "",
    ]).lower()


def _has_fcc_mention(art) -> bool:
    """True if FCC / Federal Communications Commission / a commissioner appears
    anywhere in the article's text (headline, summary, body, outlet). Uses the
    3-tier gate: title is the Article title; summary carries the rest of the text
    so Tier-3 (title-only) programs are still judged against the real headline."""
    title = getattr(art, "title", "") or ""
    rest = " ".join([getattr(art, "summary", "") or "", getattr(art, "full_text", "") or ""])
    return _is_fcc_relevant_v2(title, rest)


def fcc_relevance_points(art) -> int:
    """The client's additive FCC relevance score. Higher = more clearly FCC."""
    title = ((getattr(art, "title", "") or "")).lower()
    text = _fcc_blob(art)
    pts = 0
    if "fcc" in title or "federal communications commission" in title:
        pts += 10
    if any(c in text for c in FCC_COMMISSIONERS) or "commissioner" in text or "fcc chair" in text:
        pts += 8
    if "docket" in text or "rulemaking" in text or "proceeding" in text or "filing" in text:
        pts += 7
    if "spectrum" in text:
        pts += 6
    if "broadband" in text:
        pts += 6
    if "enforcement" in text or "forfeiture" in text or "consent decree" in text:
        pts += 5
    if "fcc vote" in text or ("vote" in text and "fcc" in text) or "order" in text:
        pts += 5
    if "satellite" in text or "space" in text:
        pts += 4
    if "911" in text or "emergency alert" in text or "public safety" in text:
        pts += 4
    return pts


def apply_fcc_relevance_boost(articles) -> int:
    """Additively nudge relevance up for clear FCC signals (commissioners, dockets,
    enforcement, etc.). NEVER lowers a score; clamped to 1.0. Returns count boosted."""
    n = 0
    for a in articles:
        pts = fcc_relevance_points(a)
        if pts <= 0:
            continue
        boost = min(0.25, pts / 200.0)   # gentle: max +0.25
        try:
            new = min(1.0, (a.relevance_score or 0.0) + boost)
            if new > (a.relevance_score or 0.0):
                a.relevance_score = new
                n += 1
        except Exception:
            pass
    return n


# Finer FCC categories (client spec) — derived deterministically for the coverage
# report. Does NOT change the displayed AGT_SECTIONS; purely analytics metadata.
FCC_CATEGORY_RULES = [
    ("Commissioners",            ("brendan carr", "anna gomez", "olivia trusty", "commissioner", "fcc chair")),
    ("Enforcement Actions",      ("enforcement", "forfeiture", "consent decree", "notice of apparent liability")),
    ("Robocalls / TCPA",         ("robocall", "tcpa", "stir-shaken", "spoofing", "robotext")),
    ("Net Neutrality",           ("net neutrality", "open internet", "title ii")),
    ("Spectrum",                 ("spectrum", "auction", "megahertz", "gigahertz", " ghz", " mhz")),
    ("Broadband",                ("broadband", "bead", "affordable connectivity", "lifeline", "e-rate")),
    ("Wireless",                 ("wireless", "5g", "cell tower", "mobile carrier", "small cell")),
    ("Satellite / Space",        ("satellite", "starlink", "ngso", "earth station", "space bureau", "kuiper")),
    ("Undersea Cables",          ("undersea cable", "subsea cable", "submarine cable")),
    ("Broadcast / Media",        ("broadcast", "radio station", "tv station", "media ownership", "license renewal")),
    ("Telecom Mergers",          ("merger", "acquisition", "deal review", "antitrust")),
    ("Public Safety / Emergency",("911", "e911", "emergency alert", "psap", "public safety")),
    ("AI and Telecom",           ("artificial intelligence", "machine learning")),
    ("International Telecom",     ("itu", "world radiocommunication", "foreign carrier", "international affairs")),
    ("Congressional Oversight",  ("senate commerce", "house energy", "subcommittee", "oversight hearing", "congress")),
]


def fcc_category(art) -> str:
    """Map an article to one of the finer FCC categories (analytics only)."""
    text = _fcc_blob(art)
    for label, terms in FCC_CATEGORY_RULES:
        if any(t in text for t in terms):
            return label
    return "General"


# Last coverage report per agency (for GET /coverage/{agency_id}).
_last_coverage: Dict[str, Dict[str, Any]] = {}


def _build_coverage_report(agency_id, all_articles, unique, classified, briefing_arts) -> Dict[str, Any]:
    """Daily source/coverage analytics (additive; never affects the briefing).

    Content stats (by_category/section, subscription, outlets) are computed on the
    NEWS that actually renders — social posts are excluded here exactly as the
    briefing renderer excludes them, so the report matches the delivered bulletin.
    """
    from collections import Counter
    news = [a for a in briefing_arts if (getattr(a, "source_type", "") or "") != "social"]
    social_collected = sum(1 for a in all_articles if (getattr(a, "source_type", "") or "") == "social")
    sources_scanned = dict(Counter(getattr(a, "source", "?") or "?" for a in all_articles))
    by_category = dict(Counter(fcc_category(a) for a in news))
    by_section = dict(Counter(_section_of(a) for a in news))
    subs = sum(1 for a in news if getattr(a, "is_paywalled", False))
    top_outlets = Counter((getattr(a, "outlet", "") or "?") for a in news).most_common(10)
    expected = [r[0] for r in FCC_CATEGORY_RULES]
    missing = [c for c in expected if by_category.get(c, 0) == 0]
    provider_analytics = _build_provider_analytics(all_articles, unique, briefing_arts)
    # NewsAPI.ai vs other providers — real unique/duplicate/additional-FCC counts.
    try:
        from app.bulletin_intelligence.provider_analysis import compare_provider_coverage
        provider_coverage = compare_provider_coverage(all_articles, target_provider="NewsAPI.ai")
    except Exception as _e:
        logger.debug(f"provider coverage comparison skipped: {_e}")
        provider_coverage = None
    return {
        "generated_at": _now(),
        "agency_id": agency_id,
        "sources_scanned": sources_scanned,
        "source_count": len(sources_scanned),
        "stories_collected": len(all_articles),
        "after_dedup": len(unique),
        "duplicates_removed": max(0, len(all_articles) - len(unique)),
        "classified": len(classified),
        "social_collected": social_collected,
        "in_briefing": len(news),
        "rejected": max(0, len(classified) - len(briefing_arts)),
        "subscription_stories": subs,
        "by_category": by_category,
        "by_section": by_section,
        "missing_category_warnings": missing,
        "top_outlets": top_outlets,
        "strict_fcc_gate": STRICT_FCC_GATE,
        "provider_analytics": provider_analytics,
        "provider_coverage_comparison": provider_coverage,
    }


def _build_provider_analytics(all_articles, unique, briefing_arts) -> Dict[str, Any]:
    """Per-provider stats for the Operations Dashboard, computed from REAL stamped
    provider data. Response time is reported as null (pending per-provider timing
    instrumentation) — never fabricated. Additive; never affects the briefing."""
    from collections import Counter

    def _prov(a) -> str:
        return (getattr(a, "provider", "") or "").strip() or "Unknown"

    unique_ids = {getattr(a, "article_id", id(a)) for a in unique}
    briefing_ids = {getattr(a, "article_id", id(a)) for a in briefing_arts}

    collected = Counter(_prov(a) for a in all_articles)
    uniq = Counter(_prov(a) for a in all_articles if getattr(a, "article_id", id(a)) in unique_ids)
    accepted = Counter(_prov(a) for a in briefing_arts)
    # Sum relevance of accepted articles per provider for a real average.
    rel_sum: Dict[str, float] = {}
    for a in briefing_arts:
        p = _prov(a)
        rel_sum[p] = rel_sum.get(p, 0.0) + float(getattr(a, "relevance_score", 0.0) or 0.0)

    out: Dict[str, Any] = {}
    for prov in collected:
        c = collected[prov]
        u = uniq.get(prov, 0)
        acc = accepted.get(prov, 0)
        out[prov] = {
            "articles_collected": c,
            "unique": u,
            "duplicates": max(0, c - u),
            "accepted": acc,
            "rejected": max(0, u - acc),
            "average_relevance": round(rel_sum.get(prov, 0.0) / acc, 3) if acc else None,
            "unique_pct": round(100.0 * u / c, 1) if c else None,
            "response_time_ms": None,   # honest: per-provider timing not yet instrumented
        }
    return out


FCC_RSS_FEEDS = {
    "fcc_news_events": [
        ("https://www.fcc.gov/news-events/rss", "FCC"),
        ("https://www.fcc.gov/rss/headlines", "FCC"),
        # FCC.gov official feeds — added 2026-06-30 (URGENT volume push). Always
        # on-topic (FCC source) so ungated. Free, no key. Any IDs that 404 are
        # skipped gracefully by ingest_rss (non-200 → continue).
        ("https://www.fcc.gov/news-events/headlines/rss.xml", "FCC Headlines"),
        ("https://www.fcc.gov/news-events/rss-feed/37521", "FCC Daily Digest"),
        ("https://www.fcc.gov/news-events/rss-feed/37516", "FCC News Releases"),
        ("https://www.fcc.gov/news-events/rss-feed/37511", "FCC Orders"),
        ("https://www.fcc.gov/news-events/rss-feed/37506", "FCC Public Notices"),
        ("https://www.fcc.gov/news-events/rss-feed/37496", "FCC NOPRs"),
        ("https://www.fcc.gov/news-events/rss-feed/37486", "FCC Citations"),
        ("https://www.fcc.gov/news-events/rss-feed/37476", "FCC Reports"),
        ("https://www.fcc.gov/news-events/rss-feed/47491", "FCC Enforcement"),
        ("https://www.fcc.gov/news-events/rss-feed/47496", "FCC Media Bureau"),
        ("https://www.fcc.gov/news-events/rss-feed/47501", "FCC Wireless Bureau"),
        ("https://www.fcc.gov/news-events/rss-feed/47506", "FCC Wireline Bureau"),
        ("https://www.fcc.gov/news-events/rss-feed/47511", "FCC Public Safety Bureau"),
        ("https://www.fcc.gov/news-events/rss-feed/47516", "FCC Space Bureau"),
        ("https://www.fcc.gov/news-events/rss-feed/47521", "FCC International"),
        ("https://www.fcc.gov/news-events/rss-feed/47526", "FCC Engineering"),
        ("https://www.fcc.gov/news-events/rss-feed/47531", "FCC Economics"),
        ("https://www.fcc.gov/news-events/rss-feed/47536", "FCC General Counsel"),
        ("https://www.fcc.gov/news-events/rss-feed/47476", "FCC Consumer Affairs"),
        ("https://www.fcc.gov/news-events/rss-feed/45291", "FCC Commissioner Statements"),
        ("https://www.fcc.gov/news-events/rss-feed/37541", "FCC Broadcast Actions"),
        ("https://www.fcc.gov/news-events/rss-feed/37546", "FCC Broadcast Applications"),
    ],
    "wireless_mobile": [
        ("https://www.fiercewireless.com/rss/xml", "FierceWireless"),
        ("https://www.rcrwireless.com/feed", "RCR Wireless"),
        # Added 2026-06-29 for higher volume — telecom-trade feeds (free, no key).
        ("https://www.fierce-network.com/rss/xml", "Fierce Network"),
        ("https://insidetowers.com/feed/", "Inside Towers"),
    ],
    "media_broadcasting": [
        ("https://www.radioworld.com/feed", "Radio World"),
        ("https://www.tvtechnology.com/rss/all", "TV Technology"),
        # ("https://www.multichannel.com/rss/all", "Multichannel News"),  # CEASED PUBLICATION Sept 2024
        ("https://rbr.com/feed/", "RBR"),
        # Added 2026-07-08 (UAT missing-source fix): Radio Insight was never wired
        # into any feed list — its feed is live (200/valid RSS). Ungated broadcast.
        ("https://radioinsight.com/feed/", "Radio Insight"),
    ],
    "consumers_advocacy": [
        ("https://broadbandbreakfast.com/feed/", "Broadband Breakfast"),
        ("https://www.telecompetitor.com/feed/", "Telecompetitor"),
    ],
    "spectrum_policy": [
        ("https://broadbandbreakfast.com/feed/", "Broadband Breakfast"),
        ("https://www.fiercewireless.com/rss/xml", "FierceWireless"),
    ],
    "public_safety_emergency": [
        ("https://www.fcc.gov/news-events/rss", "FCC"),
    ],
    "business_industry": [
        ("https://thehill.com/policy/technology/feed/", "The Hill"),
        ("https://www.politico.com/rss/politicopicks.xml", "Politico"),
    ],
    "international_affairs": [
        ("https://www.telegeography.com/feed/", "TeleGeography"),
    ],
    "ai_emerging_tech": [
        ("https://thehill.com/policy/technology/feed/", "The Hill"),
    ],
}

# ── Extended feeds (ADD-only) ─────────────────────────────────────────────────
# Merge the client's extended source list in, deduplicated by URL against every
# feed already configured above so nothing is fetched twice. Ungated FCC.gov feeds
# join FCC_RSS_FEEDS; everything else joins the FCC-relevance-gated MAJOR_OUTLET_FEEDS.
_existing_feed_urls = set()
for _flist in FCC_RSS_FEEDS.values():
    for _it in _flist:
        if isinstance(_it, tuple) and _it:
            _existing_feed_urls.add(_it[0])
for _it in MAJOR_OUTLET_FEEDS:
    if isinstance(_it, tuple) and _it:
        _existing_feed_urls.add(_it[0])

if EXTENDED_FCC_OFFICIAL:
    _new_official = [(u, n) for (u, n) in EXTENDED_FCC_OFFICIAL if u not in _existing_feed_urls]
    for u, _n in _new_official:
        _existing_feed_urls.add(u)
    if _new_official:
        FCC_RSS_FEEDS.setdefault("extended_official", []).extend(_new_official)

if EXTENDED_OUTLET_FEEDS:
    _new_outlet = [f for f in EXTENDED_OUTLET_FEEDS if f[0] not in _existing_feed_urls]
    for f in _new_outlet:
        _existing_feed_urls.add(f[0])
    MAJOR_OUTLET_FEEDS.extend(_new_outlet)

logger.info(
    "Extended feeds merged: +%d official, +%d outlet",
    len(FCC_RSS_FEEDS.get("extended_official", [])),
    len([f for f in MAJOR_OUTLET_FEEDS if f[0] in {x[0] for x in EXTENDED_OUTLET_FEEDS}]),
)


async def _resolve_google_news_url(client, url: str) -> str:
    """Google News RSS items link to a news.google.com redirect, not the publisher.
    Follow it so dedup keys off the REAL article URL — otherwise the same story
    counts twice (once via a Google News query, once via its direct feed). Falls
    back to the original URL on any failure (incl. Google's JS interstitial, which
    returns 200 without an HTTP redirect — in that case the link stays as-is)."""
    if "news.google.com" not in (url or ""):
        return url
    try:
        r = await client.head(url, headers=HTTP_HEADERS, follow_redirects=True, timeout=5.0)
        final = str(r.url)
        if final and "news.google.com" not in final:
            return final
    except Exception:
        pass
    return url


async def ingest_rss(agency: AgencyConfig, lookback_hours: int = 24) -> list:
    """
    Ingest RSS feeds from Appendix B sources.
    Always FCC-relevant — no filtering needed.
    FREE — no API key required.
    """
    import xml.etree.ElementTree as ET

    articles = []
    # Collect anything at/after the business-day window start (ET). The precise
    # [start, end) window — including dropping today's items — is enforced in the
    # render pool (_prepare_briefing_sections), which also counts in/out-of-window.
    _win_start, _ = get_briefing_window()
    cutoff = _win_start
    seen = set()

    # Each feed: (url, outlet, topic, relevance_required, paywalled).
    # Appendix B feeds are always on-topic. Major-outlet feeds (NYT/WaPo/WSJ/...)
    # are broad, so they're FCC-relevance filtered before anything is kept.
    all_feeds = []
    for topic, feeds in FCC_RSS_FEEDS.items():
        for url, outlet in feeds:
            all_feeds.append((url, outlet, topic, False, False))
    for url, outlet, paywalled in MAJOR_OUTLET_FEEDS:
        all_feeds.append((url, outlet, "other", True, paywalled))

    # Fetch all feeds CONCURRENTLY (bounded). Sequential fetching over 100+ feeds at
    # a 30s timeout meant one slow/dead feed stalled the whole cycle for minutes;
    # with the expanded source list that pushed cycles past 15 min. asyncio is
    # single-threaded so the shared `seen`/`articles` updates below stay atomic
    # (no await between the dedup check and the append).
    sem = asyncio.Semaphore(30)   # raised 12 -> 30 for the expanded (100s of) feed list

    async def _process_feed(client, feed_url, outlet, topic, relevance_required, paywalled):
        async with sem:
            # fcc.gov aggressively blocks non-browser/cloud clients (403 or a hung
            # connection), so give it a longer timeout and one retry with a browser UA.
            # Dead feeds are logged (not silently dropped) so the source list can be pruned.
            is_fccgov = "fcc.gov" in feed_url
            feed_timeout = 20.0 if is_fccgov else 10.0
            ua_attempts = [HTTP_HEADERS] + ([_BROWSER_HEADERS] if is_fccgov else [])
            resp = None
            drop_reason = "no response"
            for _hdrs in ua_attempts:
                try:
                    r = await client.get(feed_url, headers=_hdrs, follow_redirects=True, timeout=feed_timeout)
                    if r.status_code == 200:
                        resp = r
                        break
                    drop_reason = f"HTTP {r.status_code}"
                except Exception as e:
                    drop_reason = f"{type(e).__name__}: {str(e)[:60]}"
            if resp is None:
                logger.debug(f"Feed dropped [{outlet}] {drop_reason}: {feed_url[:80]}")
                return

            try:
                root = ET.fromstring(resp.text)
                ns = {"atom": "http://www.w3.org/2005/Atom"}

                # Handle both RSS and Atom feeds
                items = (root.findall(".//item") or root.findall(".//atom:entry", ns))[:25]

                # Pre-resolve Google News redirect links CONCURRENTLY (not per-item
                # sequentially — that serialized ~170 HEAD calls and made each cycle
                # take 15-20+ min). Bounded to ~one HEAD timeout per feed. Non-Google
                # feeds skip this entirely (helper returns instantly).
                _raw_links = []
                for item in items:
                    _raw_links.append((
                        getattr(item.find("link"), "text", "") or
                        (item.find("atom:link", ns).get("href") if item.find("atom:link", ns) is not None else "") or ""
                    ).strip())
                _link_map = {}
                if any("news.google.com" in (lk or "") for lk in _raw_links):
                    _resolved = await asyncio.gather(*[_resolve_google_news_url(client, lk) for lk in _raw_links])
                    _link_map = {raw: res for raw, res in zip(_raw_links, _resolved)}

                for item in items:
                    # Extract fields handling both RSS and Atom
                    title = (
                        getattr(item.find("title"), "text", "") or
                        getattr(item.find("atom:title", ns), "text", "") or ""
                    ).strip()

                    link = (
                        getattr(item.find("link"), "text", "") or
                        (item.find("atom:link", ns).get("href") if item.find("atom:link", ns) is not None else "") or ""
                    ).strip()

                    pub_date = (
                        getattr(item.find("pubDate"), "text", "") or
                        getattr(item.find("atom:published", ns), "text", "") or
                        getattr(item.find("atom:updated", ns), "text", "") or ""
                    ).strip()

                    description = (
                        getattr(item.find("description"), "text", "") or
                        getattr(item.find("atom:summary", ns), "text", "") or ""
                    ).strip()[:400]

                    if not title or not link:
                        continue

                    # Use the pre-resolved real article URL for Google News links so
                    # cross-feed dedup works (same story via Google News + direct feed).
                    link = _link_map.get(link, link)

                    # Client-excluded outlets (e.g. techdirt.com) never enter.
                    if _is_excluded_domain(link):
                        continue

                    # Broad major-outlet feeds: keep only genuine FCC stories.
                    if relevance_required and not _is_fcc_relevant_v2(title, description):
                        continue

                    # Skip duplicates
                    dedup = _hash(link, title)
                    if dedup in seen:
                        continue
                    seen.add(dedup)

                    # Date filter — require a VERIFIABLE recent date. An unparseable
                    # or missing date is treated as stale (skipped), never stamped with
                    # now(): that fallback was letting weeks-old Atom/ISO-dated items in.
                    pub_dt = _parse_pub_dt(pub_date)
                    if pub_dt is None or pub_dt < cutoff:
                        continue
                    pub_iso = pub_dt.isoformat()

                    art = Article(
                        article_id=f"{agency.agency_id}_rss_{dedup}",
                        agency_id=agency.agency_id,
                        source="rss",
                        source_type="news",
                        title=title,
                        url=link,
                        published_at=pub_iso,
                        summary=description,
                        full_text=description,
                        author="",
                        outlet=outlet,
                        topic=topic,           # pre-assigned by feed category
                        relevance_score=0.75,  # RSS feeds are always on-topic
                        is_paywalled=paywalled,
                        ingested_at=_now(),
                        dedup_hash=dedup,
                    )
                    articles.append(art)

            except Exception as e:
                logger.debug(f"RSS parse error [{outlet}] {feed_url[:80]}: {e}")

    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        await asyncio.gather(*[
            _process_feed(client, u, o, t, rr, pw)
            for (u, o, t, rr, pw) in all_feeds
        ])

    logger.info(f"RSS: {len(articles)} articles for {agency.agency_id}")
    return articles

# ── INGESTION: GDELT Project (FREE — no key needed) ──────────────────────────
# GDELT monitors 300+ languages, 65+ countries, updates every 15 minutes
# Perfect for FCC broadcast, international, and US domestic news

async def ingest_gdelt(agency: AgencyConfig, lookback_hours: int = 24) -> List[Article]:
    """
    GDELT Project API — completely free, real-time global news monitoring.
    Updates every 15 minutes. No API key required.
    """
    articles = []
    # GDELT: exact phrase only - strict FCC filter
    query = "Federal Communications Commission sourcelang:eng"

    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.get(
                "https://api.gdeltproject.org/api/v2/doc/doc",
                params={
                    "query": query,
                    "mode": "artlist",
                    "maxrecords": 25,
                    "timespan": f"{min(lookback_hours, 24)}H",
                    "sort": "DateDesc",
                    "format": "json",
                },
                headers=HTTP_HEADERS
            )
            if resp.status_code == 200:
                data = resp.json()
                for art_data in data.get("articles", []):
                    title = art_data.get("title", "")
                    url   = art_data.get("url", "")
                    if not title or not url:
                        continue
                    dedup = _hash(url, title)
                    art = Article(
                        article_id=f"{agency.agency_id}_gdelt_{dedup}",
                        agency_id=agency.agency_id,
                        source="gdelt",
                        source_type="news",
                        title=title,
                        url=url,
                        published_at=art_data.get("seendate", _now()),
                        summary=title,
                        full_text=title,
                        author="",
                        outlet=art_data.get("domain", "News"),
                        relevance_score=0.6,
                        ingested_at=_now(),
                        dedup_hash=dedup,
                    )
                    articles.append(art)
    except Exception as e:
        logger.error(f"GDELT ingestion error: {e}")

    logger.info(f"GDELT: {len(articles)} articles for {agency.agency_id}")
    return articles


# ── INGESTION: Tavily AI Search (1,000 free searches/month) ──────────────────
# AI-optimized search built specifically for AI agents — returns clean summaries

async def ingest_tavily(agency: AgencyConfig, lookback_hours: int = 24) -> List[Article]:
    """
    Tavily AI Search — purpose-built for AI agents.
    Free tier: 1,000 searches/month. Sign up at tavily.com.
    Set TAVILY_API_KEY in Railway.
    """
    if not TAVILY_KEY:
        return []

    articles = []
    for query in agency.search_queries[:3]:
        try:
            async with httpx.AsyncClient(timeout=TIMEOUT) as client:
                resp = await client.post(
                    "https://api.tavily.com/search",
                    headers={"Content-Type": "application/json"},
                    json={
                        "api_key": TAVILY_KEY,
                        "query": query,
                        "search_depth": "advanced",
                        "include_answer": True,
                        "include_raw_content": False,
                        "max_results": 8,
                        "include_domains": [],
                        "exclude_domains": [],
                        "topic": "news",
                    }
                )
                if resp.status_code == 200:
                    data = resp.json()
                    for r in data.get("results", []):
                        dedup = _hash(r.get("url",""), r.get("title",""))
                        art = Article(
                            article_id=f"{agency.agency_id}_tavily_{dedup}",
                            agency_id=agency.agency_id,
                            source="tavily",
                            source_type="news",
                            title=r.get("title",""),
                            url=r.get("url",""),
                            published_at=r.get("published_date", _now()),
                            summary=r.get("content","")[:400],
                            full_text=r.get("content",""),
                            author="",
                            outlet=r.get("url","").split("/")[2] if r.get("url") else "Web",
                            relevance_score=r.get("score", 0.7),
                            ingested_at=_now(),
                            dedup_hash=dedup,
                        )
                        articles.append(art)
        except Exception as e:
            logger.error(f"Tavily error for query '{query}': {e}")

    logger.info(f"Tavily: {len(articles)} articles for {agency.agency_id}")
    return articles


# ── INGESTION: NewsAPI (free dev tier, 80K+ sources) ─────────────────────────

async def ingest_newsapi(agency: AgencyConfig, lookback_hours: int = 24) -> List[Article]:
    """
    NewsAPI — 80,000+ news sources. Free dev tier.
    Sign up at newsapi.org. Set NEWSAPI_KEY in Railway.
    """
    if not NEWSAPI_KEY:
        return []

    articles = []
    from_date = (datetime.now() - timedelta(hours=lookback_hours)).strftime("%Y-%m-%dT%H:%M:%S")
    query = " OR ".join([f'"{q.split()[0]}"' for q in agency.search_queries[:2]])

    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.get(
                "https://newsapi.org/v2/everything",
                params={
                    "apiKey": NEWSAPI_KEY,
                    "q": query,
                    "from": from_date,
                    "sortBy": "publishedAt",
                    "language": "en",
                    "pageSize": 20,
                },
                headers=HTTP_HEADERS
            )
            if resp.status_code == 200:
                for r in resp.json().get("articles", []):
                    dedup = _hash(r.get("url",""), r.get("title",""))
                    art = Article(
                        article_id=f"{agency.agency_id}_newsapi_{dedup}",
                        agency_id=agency.agency_id,
                        source="newsapi",
                        source_type="news",
                        title=r.get("title",""),
                        url=r.get("url",""),
                        published_at=r.get("publishedAt", _now()),
                        summary=r.get("description","")[:400],
                        full_text=r.get("content","")[:800],
                        author=r.get("author",""),
                        outlet=r.get("source",{}).get("name","News"),
                        is_paywalled=False,
                        ingested_at=_now(),
                        dedup_hash=dedup,
                    )
                    articles.append(art)
    except Exception as e:
        logger.error(f"NewsAPI error: {e}")

    logger.info(f"NewsAPI: {len(articles)} articles for {agency.agency_id}")
    return articles


# ── INGESTION: NewsAPI.ai (Event Registry) — additive collector ──────────────
# Auto-detected via NEWSAPI_AI_KEY. Absent → returns [] (graceful skip, no crash).
# Flows through the SAME pipeline as every other source (normalize → boolean →
# AI relevance → dedup → editorial → category). It never bypasses any gate.

# Provider metadata registry — maps an Article.source value to (provider label,
# provider URL, collection method) so per-provider analytics + the export provider
# column are computed from REAL stamped data, not inferred at report time.
PROVIDER_REGISTRY = {
    "rss":              ("RSS",              "",                                 "rss"),
    "gdelt":            ("GDELT",            "https://www.gdeltproject.org",     "news_index"),
    "gdelt_doc":        ("GDELT",            "https://www.gdeltproject.org",     "news_index"),
    "gdelt_tv":         ("GDELT TV",         "https://www.gdeltproject.org",     "news_index"),
    "tavily":           ("Tavily",           "https://tavily.com",               "search_api"),
    "newsapi":          ("NewsAPI.org",      "https://newsapi.org",              "news_api"),
    "newsapi_ai":       ("NewsAPI.ai",       "https://newsapi.ai",               "news_api"),
    "federal_register": ("Federal Register", "https://www.federalregister.gov",  "gov_api"),
    "congress_gov":     ("Congress.gov",     "https://www.congress.gov",         "gov_api"),
    "primary":          ("Primary Sources",  "https://www.fcc.gov",              "gov_api"),
    "govinfo":          ("GovInfo",          "https://www.govinfo.gov",          "gov_api"),
    "social":           ("Social",           "",                                 "social"),
    "bluesky":          ("BlueSky",          "https://bsky.app",                 "social"),
    "reddit":           ("Reddit",           "https://www.reddit.com",           "social"),
    "youtube":          ("YouTube",          "https://www.youtube.com",          "social"),
}


def stamp_providers(articles: List["Article"]) -> None:
    """Stamp provider tracking fields on every collected article (in place).

    Idempotent and additive: only fills blank fields, so an ingester that already
    set richer provider data wins. Unknown sources fall back to a title-cased
    label so nothing is left unattributed. Never raises."""
    for a in articles:
        try:
            src = (getattr(a, "source", "") or "").strip().lower()
            label, url, method = PROVIDER_REGISTRY.get(
                src, (src.replace("_", " ").title() if src else "Unknown", "", "")
            )
            if not getattr(a, "provider", ""):
                a.provider = label
            if not getattr(a, "provider_url", ""):
                a.provider_url = url
            if not getattr(a, "collection_method", ""):
                a.collection_method = method or (src or "unknown")
            if not getattr(a, "source_name", ""):
                a.source_name = getattr(a, "outlet", "") or label
            if not getattr(a, "collection_time", ""):
                a.collection_time = getattr(a, "ingested_at", "") or _now()
        except Exception:
            continue


async def ingest_newsapi_ai(agency: AgencyConfig, lookback_hours: int = 24) -> List["Article"]:
    """NewsAPI.ai / Event Registry — additive news-index collector.

    Detects NEWSAPI_AI_KEY automatically. If the key is missing the function
    returns [] immediately (graceful skip — the cycle proceeds with all other
    providers, no exception). Results are returned as standard Article objects so
    they pass through the identical downstream pipeline.
    """
    if not NEWSAPI_AI_KEY:
        return []

    articles: List["Article"] = []
    # Build a keyword query from the agency's own search terms (same terms the
    # other collectors use). Event Registry treats a list as OR by default.
    keywords = []
    for q in (agency.search_queries or [])[:3]:
        kw = (q or "").strip()
        if kw:
            keywords.append(kw)
    if not keywords:
        keywords = ["Federal Communications Commission", "FCC"]

    date_start = (datetime.now(timezone.utc) - timedelta(hours=min(lookback_hours, 48))).strftime("%Y-%m-%d")
    payload = {
        "apiKey": NEWSAPI_AI_KEY,
        "keyword": keywords,
        "keywordOper": "or",
        "lang": "eng",
        "dateStart": date_start,
        "articlesSortBy": "date",
        "articlesCount": 50,
        "resultType": "articles",
        "dataType": ["news", "pr"],
        "includeArticleImage": False,
        "includeArticleCategories": False,
    }

    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.post(
                "https://eventregistry.org/api/v1/article/getArticles",
                headers={"Content-Type": "application/json"},
                json=payload,
            )
            if resp.status_code != 200:
                logger.warning(f"NewsAPI.ai HTTP {resp.status_code}: {resp.text[:160]}")
                return []
            data = resp.json()
            results = ((data or {}).get("articles") or {}).get("results") or []
            for r in results:
                url = (r.get("url") or "").strip()
                title = (r.get("title") or "").strip()
                if not url or not title:
                    continue
                dedup = _hash(url, title)
                outlet = ((r.get("source") or {}).get("title")) or "NewsAPI.ai"
                authors = r.get("authors") or []
                author = ", ".join(a.get("name", "") for a in authors if isinstance(a, dict))[:200]
                body = (r.get("body") or "")[:800]
                art = Article(
                    article_id=f"{agency.agency_id}_newsapiai_{dedup}",
                    agency_id=agency.agency_id,
                    source="newsapi_ai",
                    source_type="news",
                    title=title,
                    url=url,
                    published_at=_normalize_pub(r.get("dateTimePub") or r.get("dateTime") or r.get("date") or ""),
                    summary=(body[:400] or title),
                    full_text=body or title,
                    author=author,
                    outlet=outlet,
                    relevance_score=0.7,   # neutral prior; AI relevance decides, same as peers
                    ingested_at=_now(),
                    dedup_hash=dedup,
                    provider="NewsAPI.ai",
                    provider_url="https://newsapi.ai",
                    source_name=outlet,
                    collection_method="news_api",
                    collection_time=_now(),
                )
                articles.append(art)
    except Exception as e:
        logger.error(f"NewsAPI.ai error: {e}")
        return []

    logger.info(f"NewsAPI.ai: {len(articles)} articles for {agency.agency_id}")
    return articles


# ── INGESTION: Claude Web Search → News ──────────────────────────────────────
# TO REPLACE WITH PERIGON LATER: swap this function body only.
# Function signature stays identical so nothing else changes.

async def ingest_news(agency: AgencyConfig, lookback_hours: int = 24) -> List[Article]:
    """
    Uses Claude web_search tool to find real current news articles.
    Returns same Article format as Perigon integration would.
    Swap body with Perigon API call when ready.
    """
    if not ANTHROPIC_KEY:
        logger.warning("ANTHROPIC_API_KEY not set")
        return []

    client = _get_client()
    articles = []
    queries_to_run = agency.search_queries  # run ALL topic queries for full coverage

    for query in queries_to_run:
        try:
            # Step 1: Claude searches the web with FCC-specific query
            fcc_query = f"site:reuters.com OR site:politico.com OR site:broadbandbreakfast.com OR site:fiercewireless.com {query}"
            search_response = await client.messages.create(
                model="claude-haiku-4-5",
                max_tokens=2000,
                tools=[{"type": "web_search_20250305", "name": "web_search"}],
                messages=[{
                    "role": "user",
                    "content": f"Search for FCC Federal Communications Commission news from the last {lookback_hours} hours. Query: {query}. Focus on telecom industry news sources like Reuters, Politico, Broadband Breakfast, FierceWireless, Broadcasting+Cable, Multichannel News."
                }]
            )

            # Step 2: Extract structured article data
            extract_response = await client.messages.create(
                model="claude-haiku-4-5",
                max_tokens=2000,
                messages=[
                    {"role": "user", "content": f"Search for news about: {query}"},
                    {"role": "assistant", "content": search_response.content},
                    {"role": "user", "content": """From those search results, extract up to 8 news articles as a JSON array.
Each object: title, url, outlet, author, published_at (ISO date or today), summary (2 sentences), is_paywalled (bool).
Return ONLY the JSON array. No explanation."""}
                ]
            )

            raw = _extract_text(extract_response.content)
            extracted = _parse_json_safe(raw)

            if not isinstance(extracted, list):
                extracted = [extracted] if isinstance(extracted, dict) else []

            for item in extracted:
                if not item.get('title') or not item.get('url'):
                    continue
                dedup = _hash(item.get('url', ''), item.get('title', ''))
                art = Article(
                    article_id=f"{agency.agency_id}_news_{dedup}",
                    agency_id=agency.agency_id,
                    source="claude_search",
                    source_type="news",
                    title=item.get('title', ''),
                    url=item.get('url', ''),
                    published_at=item.get('published_at', _now()),
                    summary=item.get('summary', ''),
                    full_text=item.get('summary', ''),
                    author=item.get('author', ''),
                    outlet=item.get('outlet', 'Web'),
                    is_paywalled=bool(item.get('is_paywalled', False)),
                    ingested_at=_now(),
                    dedup_hash=dedup,
                )
                articles.append(art)

            await asyncio.sleep(0.3)  # rate limit courtesy

        except Exception as e:
            logger.error(f"News search error for query '{query}': {e}")

    logger.info(f"News ingestion: {len(articles)} articles for {agency.agency_id}")
    return articles


# ── INGESTION: Claude Web Search → Broadcast ─────────────────────────────────
# TO REPLACE WITH TVEYES LATER: swap this function body only.

async def ingest_broadcast(agency: AgencyConfig, lookback_hours: int = 24) -> List[Article]:
    """
    Uses Claude to find and summarize broadcast TV/radio coverage.
    Returns broadcast-format Articles with clip placeholders.
    Swap body with TVEyes API call when ready.
    """
    if not ANTHROPIC_KEY or not agency.include_broadcast:
        return []

    client = _get_client()
    articles = []
    query = agency.search_queries[0] if agency.search_queries else agency.short_name

    try:
        search_response = await client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=1500,
            tools=[{"type": "web_search_20250305", "name": "web_search"}],
            messages=[{
                "role": "user",
                "content": f"Search for TV news broadcasts, radio coverage, and video news segments about: {query}. Look for CNN, Fox News, MSNBC, C-SPAN, NPR, ABC News, CBS News, NBC News coverage."
            }]
        )

        extract_response = await client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=1500,
            messages=[
                {"role": "user", "content": f"Search for TV/radio broadcast coverage of: {query}"},
                {"role": "assistant", "content": search_response.content},
                {"role": "user", "content": """Extract up to 4 broadcast/video news items as a JSON array.
Each object: title, url, outlet (TV/radio station), published_at, summary (1-2 sentences about what was broadcast).
Return ONLY the JSON array."""}
            ]
        )

        raw = _extract_text(extract_response.content)
        extracted = _parse_json_safe(raw)

        if isinstance(extracted, list):
            for item in extracted:
                dedup = _hash(item.get('url', ''), item.get('title', ''))
                art = Article(
                    article_id=f"{agency.agency_id}_tv_{dedup}",
                    agency_id=agency.agency_id,
                    source="claude_broadcast_search",
                    source_type="broadcast",
                    title=f"[Broadcast] {item.get('title', '')}",
                    url=item.get('url', ''),
                    published_at=item.get('published_at', _now()),
                    summary=item.get('summary', ''),
                    full_text=item.get('summary', ''),
                    author="",
                    outlet=item.get('outlet', 'Broadcast'),
                    broadcast_clip_url=item.get('url', ''),
                    ingested_at=_now(),
                    dedup_hash=dedup,
                )
                articles.append(art)

    except Exception as e:
        logger.error(f"Broadcast search error: {e}")

    logger.info(f"Broadcast: {len(articles)} clips for {agency.agency_id}")
    return articles


# ── INGESTION: Claude Web Search → Social ─────────────────────────────────────
# TO REPLACE WITH TWITTER/REDDIT APIS LATER: swap this function body only.

async def ingest_social(agency: AgencyConfig) -> List[Article]:
    """
    Uses Claude to find social media discussions about the agency's topics.
    Swap body with Twitter/Reddit API calls when ready.
    """
    if not ANTHROPIC_KEY or not agency.include_social:
        return []

    client = _get_client()
    articles = []
    query = f"{agency.short_name} {agency.search_queries[0][:60] if agency.search_queries else ''}"

    try:
        search_response = await client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=1500,
            tools=[{"type": "web_search_20250305", "name": "web_search"}],
            messages=[{
                "role": "user",
                "content": f"Search for social media discussions, tweets, Reddit posts, and public commentary about: {query}. Focus on significant public discussion."
            }]
        )

        extract_response = await client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=1500,
            messages=[
                {"role": "user", "content": f"Search social media about: {query}"},
                {"role": "assistant", "content": search_response.content},
                {"role": "user", "content": """Extract up to 4 social media items (tweets, Reddit posts, LinkedIn posts, YouTube comments) as JSON array.
Each object: title (brief description), url, outlet (Twitter/Reddit/etc.), published_at, summary, author.
Return ONLY the JSON array."""}
            ]
        )

        raw = _extract_text(extract_response.content)
        extracted = _parse_json_safe(raw)

        if isinstance(extracted, list):
            for item in extracted:
                dedup = _hash(item.get('url', ''), item.get('title', ''))
                art = Article(
                    article_id=f"{agency.agency_id}_social_{dedup}",
                    agency_id=agency.agency_id,
                    source="claude_social_search",
                    source_type="social",
                    title=item.get('title', ''),
                    url=item.get('url', ''),
                    published_at=item.get('published_at', _now()),
                    summary=item.get('summary', ''),
                    full_text=item.get('summary', ''),
                    author=item.get('author', ''),
                    outlet=item.get('outlet', 'Social'),
                    ingested_at=_now(),
                    dedup_hash=dedup,
                )
                articles.append(art)

    except Exception as e:
        logger.error(f"Social search error: {e}")

    logger.info(f"Social: {len(articles)} posts for {agency.agency_id}")
    return articles


# ── INGESTION: Federal Register + Congress.gov (REAL — FREE APIs) ─────────────
# These stay as real APIs — no cost, no key required for Federal Register.

async def ingest_regulatory(agency: AgencyConfig) -> List[Article]:
    """Real Federal Register and Congress.gov APIs — free, no API key for Federal Register."""
    if not agency.include_regulatory:
        return []

    articles = []

    # Federal Register — completely free, no auth
    for query in agency.search_queries[:2]:
        try:
            async with httpx.AsyncClient(timeout=TIMEOUT) as client:
                resp = await client.get(
                    "https://www.federalregister.gov/api/v1/documents.json",
                    params={
                        "conditions[term]": query[:80],
                        "conditions[publication_date][gte]": (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%d"),
                        "fields[]": ["title", "type", "abstract", "html_url", "publication_date", "agencies"],
                        "per_page": 5,
                        "order": "newest",
                    },
                    headers=HTTP_HEADERS
                )
                if resp.status_code == 200:
                    for doc in resp.json().get("results", []):
                        dedup = _hash(doc.get("html_url", ""), doc.get("title", ""))
                        agency_names = ", ".join([a.get("name", "") for a in doc.get("agencies", [])])
                        art = Article(
                            article_id=f"{agency.agency_id}_fr_{dedup}",
                            agency_id=agency.agency_id,
                            source="federal_register",
                            source_type="regulatory",
                            title=f"[Federal Register] {doc.get('title', '')}",
                            url=doc.get("html_url", ""),
                            published_at=doc.get("publication_date", datetime.now().strftime("%Y-%m-%d")),
                            summary=doc.get("abstract", "")[:400],
                            full_text=doc.get("abstract", ""),
                            author="",
                            outlet=f"Federal Register — {agency_names}",
                            article_type="regulatory",
                            relevance_score=0.85,
                            ingested_at=_now(),
                            dedup_hash=dedup,
                        )
                        articles.append(art)
        except Exception as e:
            logger.error(f"Federal Register error: {e}")

    # Congress.gov — free, optional API key for higher rate limits
    if CONGRESS_KEY:
        try:
            async with httpx.AsyncClient(timeout=TIMEOUT) as client:
                resp = await client.get(
                    "https://api.congress.gov/v3/bill",
                    params={
                        "apiKey": CONGRESS_KEY,
                        "format": "json",
                        "limit": 5,
                        "fromDateTime": (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%dT00:00:00Z"),
                    },
                    headers=HTTP_HEADERS
                )
                if resp.status_code == 200:
                    for bill in resp.json().get("bills", []):
                        dedup = _hash(bill.get("url", ""), bill.get("title", ""))
                        art = Article(
                            article_id=f"{agency.agency_id}_cg_{dedup}",
                            agency_id=agency.agency_id,
                            source="congress_gov",
                            source_type="regulatory",
                            title=f"[Congress] {bill.get('title', 'Untitled Bill')}",
                            url=bill.get("url", ""),
                            published_at=bill.get("latestAction", {}).get("actionDate", _now()),
                            summary=bill.get("latestAction", {}).get("text", ""),
                            full_text=bill.get("title", ""),
                            author="",
                            outlet=f"Congress.gov — {bill.get('type', '')} {bill.get('number', '')}",
                            article_type="regulatory",
                            relevance_score=0.80,
                            ingested_at=_now(),
                            dedup_hash=dedup,
                        )
                        articles.append(art)
        except Exception as e:
            logger.error(f"Congress.gov error: {e}")

    logger.info(f"Regulatory: {len(articles)} items for {agency.agency_id}")
    return articles


# ── Deduplication ──────────────────────────────────────────────────────────────
def _final_score(art) -> float:
    """Composite relevance+authority+recency rank (scoring.py). Falls back to
    relevance alone if the scoring module is unavailable, so dedup never breaks."""
    try:
        from app.bulletin_intelligence import scoring
        return scoring.final_score(
            getattr(art, "relevance_score", 0.5) or 0.0,
            getattr(art, "outlet", "") or "",
            getattr(art, "published_at", "") or "",
        )
    except Exception:
        return (getattr(art, "relevance_score", 0.0) or 0.0) * 100.0


def deduplicate(articles: List[Article]) -> List[Article]:
    seen_hashes, seen_titles, unique = set(), {}, []
    # Sort by composite score so the highest-AUTHORITY copy of a duplicated story
    # is the one retained (e.g. Reuters kept over an unknown blog), not merely the
    # highest LLM-relevance copy.
    for art in sorted(articles, key=_final_score, reverse=True):
        if art.dedup_hash in seen_hashes:
            continue
        title_key = art.title[:60].lower().strip()
        if title_key in seen_titles:
            continue
        seen_hashes.add(art.dedup_hash)
        seen_titles[title_key] = True
        unique.append(art)
    logger.info(f"Dedup: {len(articles)} → {len(unique)} articles")
    return unique


# ── Classification: Claude Haiku ──────────────────────────────────────────────
async def classify_articles(articles: List[Article], agency: AgencyConfig) -> List[Article]:
    if not ANTHROPIC_KEY or not articles:
        return articles

    client = _get_client()
    topics_str = "\n".join([f"  {t}: {TOPIC_LABELS.get(t, t)}" for t in agency.topics])
    batch_size = 8
    classified = []

    for i in range(0, len(articles), batch_size):
        batch = articles[i:i + batch_size]
        items = [{"id": a.article_id, "title": a.title, "outlet": a.outlet, "summary": a.summary[:400]} for a in batch]

        prompt = f"""You are classifying news for a daily intelligence briefing about {agency.name} ({agency.short_name}).
Return a JSON array, one object per article:
  id, topic (from list), section (from SECTIONS), article_type (news/opinion/analysis/editorial/press_release/regulatory), sentiment (positive/negative/neutral), relevance_score (0.0-1.0)

relevance_score = how directly the article concerns {agency.name} ({agency.short_name}) itself — its
leadership/officials, decisions, proceedings, rulings, votes, enforcement, or a matter that genuinely
requires its approval. Be STRICT; do not inflate.

CRITICAL — score 0.0-0.2 (NOT relevant) if ANY of these is true, no matter how on-topic it seems:
  - "{agency.short_name}" refers to something OTHER than {agency.name} (e.g. a football/soccer club, a
    church, a college, a company ticker, or a person's initials). "{agency.short_name}" has many meanings;
    only the U.S. {agency.name} counts.
  - "{agency.short_name}" appears only incidentally — in an image/photo caption, an advertisement, a
    navigation menu or sidebar, a "related links" list, or as a single passing mention — with NO actual
    {agency.short_name} action, official, ruling, or decision discussed in the body.
  - There is no genuine {agency.short_name} involvement at all (generic tech/business/world news that just
    happens to contain a related word).

When it IS genuinely about {agency.name}:
  0.85-1.0 : explicitly about {agency.short_name} actions, officials, or proceedings
  0.55-0.80: a topic {agency.short_name} regulates, with {agency.short_name} clearly involved
  0.40-0.54: {agency.short_name}-adjacent and genuinely implicated (e.g. a merger that needs
             {agency.short_name} approval) but the agency is not the central subject
  0.00-0.39: no real {agency.short_name} connection
When in doubt about whether "{agency.short_name}" even refers to {agency.name}, score LOW (below 0.4).

Topics:
{topics_str}
  other: Does not fit above topics

SECTIONS — pick the ONE display bucket each story belongs in (use these EXACT names):
  General                                 : {agency.short_name} leadership/commissioners, meetings, enforcement, votes, orders, general agency news
  Consumers                               : robocalls/TCPA, scams, accessibility, consumer protection, E-Rate, Lifeline
  Media & Broadcasting                    : TV, radio, cable, satellite TV/radio, broadcast licenses, media mergers
  Space Policy                            : satellites, NGSO/GSO, earth stations, space bureau, launch + spectrum
  Public Safety / Cybersecurity / Privacy : 911/E911, emergency alerts, outages, cybersecurity, data breaches, privacy
  Wireless & Spectrum                     : spectrum, auctions, 5G, wireless/mobile carriers, cell towers, small cells
  Broadband & Infrastructure              : fiber/broadband deployment, BEAD, USF, pole attachments, Connect America, digital equity
  AI / Machine Learning                   : artificial intelligence, machine learning, emerging tech
  Business & Tech                         : net neutrality, internet policy, telecom industry/markets, big tech, mergers
  Enforcement & Consumer                  : FCC fines/forfeitures, penalties, pirate radio, consent decrees, enforcement actions
  International                           : foreign telecom, undersea/subsea cables, ITU, treaties, {agency.short_name} international affairs

Articles:
{json.dumps(items)}

Return ONLY the JSON array."""

        try:
            resp = await client.messages.create(
                model="claude-haiku-4-5",
                max_tokens=1500,
                messages=[{"role": "user", "content": prompt}]
            )
            # Phase 1 cost tracking — best-effort, never raises, no-op unless
            # BULLETIN_COST_TRACKING_ENABLED=true. Does not alter classification.
            await _record_llm_cost(resp, operation="classify_articles", model="claude-haiku-4-5")
            results = _parse_json_safe(_extract_text(resp.content))
            rmap = {r["id"]: r for r in (results if isinstance(results, list) else [])}

            for art in batch:
                r = rmap.get(art.article_id, {})
                art.topic          = r.get("topic", "other")
                art.article_type   = r.get("article_type", "news")
                art.sentiment      = r.get("sentiment", "neutral")
                art.relevance_score = float(r.get("relevance_score", 0.5))
                sect = (r.get("section") or "").strip()
                art.section = sect if sect in AGT_SECTIONS else TOPIC_TO_SECTION.get(art.topic, "General")
                classified.append(art)

        except Exception as e:
            logger.error(f"Classification error: {e}")
            classified.extend(batch)

    logger.info(f"Classified {len(classified)} articles for {agency.agency_id}")
    return classified


# ── Briefing Generator: AGT FCC Daily News format (deterministic template) ─────
# The HTML layout is rendered in code (ported from the client's fcc_digest.py) so
# it always matches the deliverable exactly; the AI only writes the per-story
# summaries. Groups stories into the 6 client display sections (AGT_SECTIONS),
# news only (social posts feed nothing here — they'd go in a separate social block).
async def generate_briefing_html(agency: AgencyConfig, articles: List[Article], briefing_date: str) -> str:
    try:
        sections = await _prepare_briefing_sections(agency, articles)
        return _render_agt_html(agency, briefing_date, sections)
    except Exception as e:
        logger.error(f"AGT briefing render failed: {e}")
        return _simple_html(agency, articles, briefing_date)


async def build_briefing_outputs(agency: AgencyConfig, articles: List[Article], briefing_date: str):
    """Build the briefing sections ONCE and render BOTH the HTML and the editable
    .docx from them (so summaries aren't generated twice). Returns
    (html, docx_bytes, sections). `sections` is what actually rendered, so the caller
    can report a story count that matches the document instead of the pre-render
    candidate count; it is [] when the AGT render failed and the simple fallback was
    used. docx_bytes is b'' if the Word render fails — HTML is always returned. Never
    raises: any failure falls back to a simple HTML so the cycle can't crash."""
    try:
        sections = await _prepare_briefing_sections(agency, articles)
        html = _render_agt_html(agency, briefing_date, sections)
    except Exception as e:
        logger.error(f"AGT briefing build failed; using simple HTML: {e}")
        return _simple_html(agency, articles, briefing_date), b"", []
    try:
        docx_bytes = _render_agt_docx(agency, briefing_date, sections)
    except Exception as e:
        logger.warning(f"DOCX render failed: {e}")
        docx_bytes = b""
    return html, docx_bytes, sections


# ── Boolean section matching (the client's Appendix A spec) ────────────────────
try:
    # Phase 2: bound to the profiles cache, which is itself seeded verbatim from
    # fcc_boolean_search.FCC_SEARCH_TOPICS. Same object identity is retained across
    # refresh_from_db() (it mutates in place), so _boolean_section below needs no
    # change and picks up DB-driven queries automatically. Falls back to the
    # hardcoded constants whenever the table is empty or the flag is off.
    from app.bulletin_intelligence.profiles.boolean_profiles import PROFILES as _FCC_BOOL
except Exception:
    try:
        from app.bulletin_intelligence.fcc_boolean_search import FCC_SEARCH_TOPICS as _FCC_BOOL
    except Exception:
        _FCC_BOOL = {}

# Map the boolean-spec topic keys -> our AGT_SECTIONS display names, in spec order.
_BOOL_KEY_TO_SECTION = {
    "FCC_NEWS":            "General",
    "CONSUMERS":           "Consumers",
    "MEDIA_BROADCASTING":  "Media & Broadcasting",
    "SPACE_POLICY":        "Space Policy",
    "PUBLIC_SAFETY_CYBER": "Public Safety / Cybersecurity / Privacy",
    "WIRELESS_SPECTRUM":   "Wireless & Spectrum",
    "AI_MACHINE_LEARNING": "AI / Machine Learning",
    "BUSINESS_TECH":       "Business & Tech",
    "INTERNATIONAL":       "International",
}


def _phrase_in(phrase: str, text: str) -> bool:
    import re
    p = phrase.lower().strip().strip('"')
    if p.startswith("title:"):
        p = p[6:].strip().strip('"')
    if not p:
        return False
    if re.fullmatch(r"[a-z0-9][a-z0-9\-]*", p):          # word-bounded for short tokens
        return re.search(r"(?<![a-z0-9])" + re.escape(p) + r"(?![a-z0-9])", text) is not None
    return p in text


def _boolean_matches(expr: str, text: str) -> bool:
    """Evaluate a boolean expression against an article's lowercased text.

    Grammar (Phase 2 — NOT added; precedence NOT > AND > OR):
        or   := and (OR and)*
        and  := not_ (AND not_)*
        not_ := (NOT | !) not_ | atom
        atom := "(" or ")" | token

    NOT binds tighter than AND, so `a AND NOT b` parses as `a AND (NOT b)` and
    `NOT a OR b` as `(NOT a) OR b` — standard Boolean precedence, and what an editor
    writing `spectrum AND NOT sports` expects. Without negation there was no way to
    exclude a false-positive class, which is why the US/FCC-focus requirement needed
    this.

    Backward compatible: an expression containing no NOT/! token takes exactly the
    same path as before. A bare unquoted `not` in a query is now an operator — quote
    it ("not") to match it as a literal word.
    """
    import re
    if not expr:
        return False
    tokens = re.findall(r'"[^"]*"|\(|\)|[^\s()]+', expr)

    def atom(pos):
        if pos >= len(tokens):
            return False, pos
        if tokens[pos] == "(":
            val, pos = _or(pos + 1)
            if pos < len(tokens) and tokens[pos] == ")":
                pos += 1
            return val, pos
        tok = tokens[pos]
        return _phrase_in(tok, text), pos + 1

    def _not(pos):
        # Right-associative so `NOT NOT x` folds correctly back to `x`.
        if pos < len(tokens) and (tokens[pos].upper() == "NOT" or tokens[pos] == "!"):
            val, pos = _not(pos + 1)
            return (not val), pos
        return atom(pos)

    def _and(pos):
        val, pos = _not(pos)
        while pos < len(tokens) and tokens[pos].upper() == "AND":
            rhs, pos = _not(pos + 1)
            val = val and rhs
        return val, pos

    def _or(pos):
        val, pos = _and(pos)
        while pos < len(tokens) and tokens[pos].upper() == "OR":
            rhs, pos = _and(pos + 1)
            val = val or rhs
        return val, pos


    try:
        return bool(_or(0)[0])
    except Exception:
        return False


# Check SPECIFIC topics before broad catch-alls (Wireless/FCC News) so e.g. an
# international undersea-cable story isn't grabbed by the broad "FCC AND telecom".
_BOOL_MATCH_ORDER = [
    "INTERNATIONAL", "SPACE_POLICY", "AI_MACHINE_LEARNING", "PUBLIC_SAFETY_CYBER",
    "CONSUMERS", "MEDIA_BROADCASTING", "BUSINESS_TECH", "WIRELESS_SPECTRUM", "FCC_NEWS",
]


def _boolean_section(title: str, summary: str):
    """Return the AGT section whose boolean matches (specific topics first), or None."""
    text = f"{title} {summary}".lower()
    for key in _BOOL_MATCH_ORDER:
        sec = _BOOL_KEY_TO_SECTION.get(key)
        spec = _FCC_BOOL.get(key, {})
        expr = spec.get("boolean", "")
        if expr:
            if _boolean_matches(expr, text):
                return sec
        else:  # keyword-only topic (AI / Machine Learning)
            if any(_phrase_in(kw, text) for kw in spec.get("keywords", [])):
                return sec
    return None


# ── FCC-org category classifier (2026-07-02) ─────────────────────────────────
# Assigns a story to one AGT section by FCC org structure. General signals
# (governance, courts, leadership) win over sector keywords, so e.g. "SCOTUS FCC
# fine ruling" lands in General, not Enforcement. Returns None when nothing
# strong matches, so the client's boolean spec still handles the remainder.
_CAT_GENERAL = ("scotus", "supreme court", "appeals court", "d.c. circuit", "court ruling",
    "lawsuit", "litigation", "chairman", "chairwoman", "commissioner", "brendan carr",
    "anna gomez", "olivia trusty", "nomination", "confirmation", "oversight hearing",
    "general counsel", "inspector general", "gao report", "reorganization", "governance",
    "resign", "sworn in", "testimony", "agency budget")
_CAT_ENFORCEMENT = ("forfeiture", "notice of apparent liability", "consent decree",
    "fine", "penalty", "pirate radio", "enforcement action", "monetary penalty",
    "cease and desist")
_CAT_PUBLIC_SAFETY = ("911", "e911", "ng911", "next generation 911", "emergency alert",
    "wireless emergency alert", "network outage", "cybersecurity", "data breach",
    "privacy", "psap", "public safety")
_CAT_SPACE = ("satellite", "ngso", "gso", "earth station", "space bureau", "starlink",
    "spacex", "low earth orbit", "leo constellation", "orbital", "amazon kuiper")
_CAT_MEDIA = ("broadcast", "television", " tv ", "radio station", "cable ", "retransmission",
    "media ownership", "indecency", "license renewal", "sinclair", "nexstar", "atsc", "nextgen tv")
_CAT_CONSUMERS = ("robocall", "tcpa", "scam", "accessibility", "lifeline", "e-rate",
    "spoofing", "disability access", "consumer protection")
_CAT_BROADBAND = ("broadband", "fiber", "bead", "pole attachment", "connect america",
    "digital equity", "rural broadband", "affordable connectivity", "universal service")
_CAT_WIRELESS = ("spectrum auction", "spectrum", "5g", "6g", "cell tower", "small cell",
    "c-band", "cbrs", "aws-3", "mid-band", "wireless carrier", "millimeter wave")
_CAT_AI = ("artificial intelligence", "machine learning", " ai ", "generative ai", "deepfake")
_CAT_INTL = ("undersea cable", "subsea cable", "submarine cable", "itu ", "treaty",
    "foreign carrier", "team telecom")
_CAT_BUSINESS = ("net neutrality", "merger", "acquisition", "section 230", "big tech",
    "open internet", "antitrust")


def get_category(title: str, summary: str = "") -> Optional[str]:
    """Best FCC-org section for a story, General-first. None => defer to the
    client's boolean spec / model section downstream."""
    text = f" {title} {summary} ".lower()
    if any(k in text for k in _CAT_GENERAL):
        return "General"
    for terms, section in (
        (_CAT_ENFORCEMENT, "Enforcement & Consumer"),
        (_CAT_PUBLIC_SAFETY, "Public Safety / Cybersecurity / Privacy"),
        (_CAT_SPACE, "Space Policy"),
        (_CAT_MEDIA, "Media & Broadcasting"),
        (_CAT_CONSUMERS, "Consumers"),
        (_CAT_BROADBAND, "Broadband & Infrastructure"),
        (_CAT_WIRELESS, "Wireless & Spectrum"),
        (_CAT_AI, "AI / Machine Learning"),
        (_CAT_INTL, "International"),
        (_CAT_BUSINESS, "Business & Tech"),
    ):
        if any(k in text for k in terms):
            return section
    return None


# ── Chairman / Commissioner activity tagging (2026-07-02) ────────────────────
_LEADERSHIP_PREFIXES = (
    ("chairman carr", "CHAIRMAN CARR"), ("brendan carr", "CHAIRMAN CARR"),
    ("commissioner gomez", "COMMISSIONER GOMEZ"), ("anna gomez", "COMMISSIONER GOMEZ"),
    ("commissioner trusty", "COMMISSIONER TRUSTY"), ("olivia trusty", "COMMISSIONER TRUSTY"),
)


def _leadership_prefix(art) -> str:
    """'CHAIRMAN CARR' / 'COMMISSIONER GOMEZ' / 'COMMISSIONER TRUSTY' when a named
    FCC official appears in the story, else ''. Surfaces leadership activity at the
    top of the General section with a distinguishing headline prefix."""
    blob = f"{getattr(art, 'title', '')} {getattr(art, 'summary', '')}".lower()
    for needle, label in _LEADERSHIP_PREFIXES:
        if needle in blob:
            return label
    return ""


# ── Paywall / subscription detection (2026-07-02) ────────────────────────────
# Rendered as "SUBSCRIPTION REQUIRED" in the TOC + story body (see _render_agt_*).
PAYWALL_DOMAINS = (
    "wsj.com", "bloomberg.com", "ft.com", "nytimes.com", "law360.com",
    "washingtonpost.com", "politicopro.com", "theinformation.com", "economist.com",
)


def _is_paywalled_url(url: str) -> bool:
    u = (url or "").lower()
    return any(dom in u for dom in PAYWALL_DOMAINS)


def _section_of(art: "Article") -> str:
    # FCC-org category (General-first) leads; then the client's boolean spec, the
    # model's section, and the topic map. Never returns an invalid section.
    cat = get_category(art.title or "", art.summary or "")
    if cat in AGT_SECTIONS:
        return cat
    bs = _boolean_section(art.title or "", art.summary or "")
    if bs in AGT_SECTIONS:
        return bs
    if art.section in AGT_SECTIONS:
        return art.section
    sec = TOPIC_TO_SECTION.get(art.topic, "General")
    return sec if sec in AGT_SECTIONS else "General"   # never drop a story


def _clean_headline(title: str) -> str:
    """Strip a leading source tag like '[Federal Register] ' or '[Broadcast] ' that
    some ingesters prepend — the outlet is already shown separately."""
    t = (title or "").strip()
    if t.startswith("[") and "]" in t[:40]:
        t = t[t.index("]") + 1:].strip()
    return t


async def _prepare_briefing_sections(agency: AgencyConfig, articles: List[Article]):
    # Aggregate THIS cycle's articles with everything collected for this agency,
    # then keep only what falls inside the PREVIOUS-BUSINESS-DAY window (ET).
    win_start, win_end = get_briefing_window()

    def _in_window(a) -> bool:
        # Fail-closed: an unparseable/missing date is treated as out-of-window
        # (excluded), never rendered as if fresh. Aware-datetime comparison, so a
        # UTC article date is judged as the same instant as the ET boundaries.
        pub = _parse_pub_dt(a.published_at or "")
        return pub is not None and win_start <= pub < win_end

    # Candidate set = this cycle + everything else stored for this agency.
    candidates = list(articles)
    _have = {a.article_id for a in articles}
    for a in _articles.values():
        if a.agency_id == agency.agency_id and a.article_id not in _have:
            candidates.append(a)

    pool = {a.article_id: a for a in candidates if _in_window(a)}
    excluded = len(candidates) - len(pool)
    # Actual publish-date range of what was admitted — proves everything sits
    # inside the window (no manual date inspection needed).
    _pub_dts = [d for d in (_parse_pub_dt(a.published_at or "") for a in pool.values()) if d]
    _last_window_stats[agency.agency_id] = {
        "window_start_et": win_start.isoformat(),
        "window_end_et": win_end.isoformat(),
        "candidates": len(candidates),
        "in_window": len(pool),
        "excluded_out_of_window": excluded,
        "pool_pub_min": min(_pub_dts).isoformat() if _pub_dts else None,
        "pool_pub_max": max(_pub_dts).isoformat() if _pub_dts else None,
        "computed_at": _now(),
    }
    logger.info(
        f"Briefing window (ET) {win_start.isoformat()} .. {win_end.isoformat()}: "
        f"{len(pool)} in-window, {excluded} excluded of {len(candidates)} candidates"
    )

    # Label known subscription outlets so the briefing shows [SUBSCRIPTION REQUIRED]
    # (deterministic pre-pass; safe no-op if the module is unavailable).
    try:
        from app.bulletin_intelligence.editorial_rules import flag_subscriptions
        flag_subscriptions(list(pool.values()))
    except Exception as _e:
        logger.debug(f"flag_subscriptions skipped: {_e}")

    # NEWS only — relevant, non-social — then drop duplicate stories across cycles.
    # Relevance gate matches the briefing filter in run_daily_cycle: an article that
    # the classifier gave a real topic is kept regardless of score, and only untyped
    # ("other") items must clear 0.4. This used to require >= 0.4 unconditionally,
    # so a classified-but-low-scoring story was counted in article_count yet never
    # rendered — the cause of the reported-vs-visible gap (146 reported / 41 shown).
    news = [
        a for a in pool.values()
        if (a.topic != "other" or a.relevance_score >= 0.4) and a.source_type != "social"
    ]

    # Lenient spam/junk removal (press releases, malformed URLs, listicles). Uses a
    # LOW threshold so only clear junk is dropped — volume is preserved.
    try:
        from app.bulletin_intelligence.clustering import quality_score
        news = [a for a in news if quality_score(a) >= 0.35]
    except Exception as _e:
        logger.debug(f"quality filter skipped: {_e}")

    news = deduplicate(news)
    # Authority-aware ordering: the fcc.gov cap and cluster-primary selection below
    # both consume this order, so the most authoritative stories lead each cluster.
    news.sort(key=_final_score, reverse=True)

    # Cap fcc.gov to 10 (client gets FCC.gov directly) and total stories to 200.
    capped, fcc_gov = [], 0
    for a in news:
        if "fcc.gov" in (a.url or "").lower():
            if fcc_gov >= 10:
                continue
            fcc_gov += 1
        capped.append(a)
        if len(capped) >= 200:
            break

    # Cluster same-story-different-outlet into one primary + 'Similar Stories',
    # then summarize ONLY the primaries (fewer Claude calls too).
    clusters = _cluster_stories(capped)
    primaries = [c[0] for c in clusters]
    summaries = await _summaries_for(primaries, agency)
    return _collect_sections(clusters, summaries)


def _cluster_stories(articles: List[Article]) -> List[List[Article]]:
    """Group articles covering the SAME story across outlets (near-duplicate
    headlines like 'RBR:' vs 'RBR.COM:' on the June Open Meeting). Input is sorted
    by relevance, so each cluster's first member is the primary; the rest become
    'Similar Stories'."""
    import re
    STOP = {"the", "a", "an", "of", "to", "in", "on", "for", "and", "or", "at", "by",
            "with", "as", "is", "are", "be", "new", "says", "said", "after", "over",
            "from", "amid", "its", "it", "fcc", "federal", "communications", "commission"}

    def toks(title):
        words = re.findall(r"[a-z0-9]+", _clean_headline(title).lower())
        return {w for w in words if len(w) > 2 and w not in STOP}

    clusters = []  # each: {"toks": set, "members": [Article]}
    for a in articles:
        tk = toks(a.title)
        placed = False
        for cl in clusters:
            ct = cl["toks"]
            union = len(tk | ct) if (tk and ct) else 0
            inter = len(tk & ct) if union else 0
            # Related-story clustering (Jaccard overlap >= 0.30, min 2 shared tokens).
            # Groups multi-outlet coverage of one story into a single primary + its
            # RELATED list, so ~60 items render as ~25-30 story clusters. (Loosens the
            # prior 0.75 gate per the 2026-07-02 "related story clustering" spec.)
            if union and inter >= 2 and (inter / union) >= 0.30:
                cl["members"].append(a)
                cl["toks"] = ct | tk
                placed = True
                break
        if not placed:
            clusters.append({"toks": tk, "members": [a]})
    return [cl["members"] for cl in clusters]


def _collect_sections(clusters, summaries: Dict[str, str]):
    """Build all sections (always present). Each cluster -> one primary story plus
    its RELATED coverage. Named-leadership items are prefixed (e.g. 'CHAIRMAN CARR:')
    and floated to the top of the General section. Empty section headers are kept."""
    idx = 0
    by_section = {s: [] for s in AGT_SECTIONS}
    for members in clusters:
        primary = members[0]
        sec = _section_of(primary)
        if sec not in by_section:
            sec = "General"
        idx += 1
        prefix = _leadership_prefix(primary)
        headline = _clean_headline(primary.title)
        # Surface chairman/commissioner activity with a headline prefix (General only).
        if prefix and sec == "General":
            headline = f"{prefix}: {headline}"
        by_section[sec].append({
            "source": (primary.outlet or primary.source or "NEWS").strip(),
            "headline": headline,
            "url": (primary.url or "").strip(),
            "anchor": f"story_{idx}",
            # Private (leading underscore, like _leader): carried so the caller can
            # build topic_counts from what actually rendered. Renderers read keys
            # explicitly, so an extra key is inert.
            "_topic": getattr(primary, "topic", "other") or "other",
            "summary": (summaries.get(primary.article_id) or primary.summary or "").strip(),
            "is_paywalled": bool(primary.is_paywalled) or _is_paywalled_url(primary.url),
            "_leader": bool(prefix and sec == "General"),
            "similar": [{
                "source": (m.outlet or m.source or "NEWS").strip(),
                "headline": _clean_headline(m.title),
                "url": (m.url or "").strip(),
                "is_paywalled": bool(m.is_paywalled) or _is_paywalled_url(m.url),
            } for m in members[1:]],
        })
    # Float leadership activity to the top of General (stable order otherwise).
    by_section["General"].sort(key=lambda s: not s.get("_leader"))
    return [(sec, by_section[sec]) for sec in AGT_SECTIONS]


async def _summaries_for(articles: List[Article], agency: AgencyConfig) -> Dict[str, str]:
    """Claude-written 2-4 sentence factual summaries, keyed by article_id."""
    fallback = {a.article_id: (a.summary or a.title or "")[:400] for a in articles}
    if not ANTHROPIC_KEY or not articles:
        return fallback
    client = _get_client()

    async def _summary_batch(batch: List[Article]) -> Dict[str, str]:
        items = [{"id": a.article_id, "title": a.title, "outlet": a.outlet,
                  "text": (a.full_text or a.summary or "")[:1500]} for a in batch]
        prompt = (
            f"You write summaries for the {agency.short_name} Daily News Monitoring briefing — a "
            "government news-clipping service. For each item write a 2-4 sentence FACTUAL summary: "
            "lead with the concrete news (who/what/when, dates, dollar amounts, votes, notable quotes). "
            "No opinion, no marketing language. If the provided text is thin, summarize only what is "
            'given — do NOT invent facts. Return ONLY a JSON array of {"id":"...","summary":"..."}.\n\n'
            "Items:\n" + json.dumps(items)
        )
        res: Dict[str, str] = {}
        try:
            resp = await client.messages.create(
                model="claude-haiku-4-5", max_tokens=2400,
                messages=[{"role": "user", "content": prompt}],
            )
            # Phase 1 cost tracking — best-effort, never raises, no-op unless
            # BULLETIN_COST_TRACKING_ENABLED=true. Does not alter summaries.
            await _record_llm_cost(resp, operation="summaries", model="claude-haiku-4-5")
            for r in (_parse_json_safe(_extract_text(resp.content)) or []):
                if isinstance(r, dict) and r.get("id") and (r.get("summary") or "").strip():
                    res[r["id"]] = r["summary"].strip()
        except Exception as e:
            logger.warning(f"Summary batch failed: {e}")
        return res

    # Run the batches concurrently so 60 stories don't serialize into a slow cycle.
    batches = [articles[i:i + 8] for i in range(0, len(articles), 8)]
    results = await asyncio.gather(*[_summary_batch(b) for b in batches])
    out: Dict[str, str] = {}
    for r in results:
        out.update(r)
    for k, v in fallback.items():
        out.setdefault(k, v)
    return out


# Header logo = the AGENCY's own logo (e.g. the FCC seal), taken from
# AgencyConfig.logo_url. This env is just a fallback if the config has none.
# Must be a publicly-hosted https image URL; blank = no logo (never a broken image).
_AGT_LOGO_URL = os.getenv("AGENCY_LOGO_URL", "").strip()


def _render_agt_html(agency: AgencyConfig, briefing_date: str, sections) -> str:
    """Render the briefing as Outlook-safe HTML with INLINE styles. Outlook ignores
    <style> blocks (and strips them on copy/paste into a new email), so every style
    is inlined. Layout matches the client's fcc_digest.py output."""
    import html as _h
    esc = _h.escape

    S_LINK = "color:#003087;text-decoration:underline"
    S_SRC  = "font-weight:bold;color:#003087"
    SUB    = (' <span style="font-weight:bold;color:#555555;text-decoration:none">'
              'SUBSCRIPTION REQUIRED</span>')

    def extlink(url, label):
        if not url:
            return esc(label)
        return (f'<a href="{esc(url)}" target="_blank" rel="noopener noreferrer" '
                f'style="{S_LINK}">{esc(label)}</a>')

    toc = []
    for name, stories in sections:
        toc.append('<div style="font-size:10px;font-weight:bold;color:#003087;letter-spacing:1.5px;'
                   'text-transform:uppercase;margin:14px 0 7px;border-left:4px solid #003087;'
                   f'padding-left:8px">{esc(name)}</div>')
        for s in stories:
            sub = SUB if s["is_paywalled"] else ""
            toc.append(
                '<div style="font-size:13px;padding:3px 0 3px 12px;line-height:1.5;margin-bottom:4px">'
                f'<span style="{S_SRC}">{esc(s["source"])}:</span> '
                f'<a href="#{esc(s["anchor"])}" style="{S_LINK}">{esc(s["headline"])}</a>{sub}</div>'
            )

    body = []
    for name, stories in sections:
        body.append('<div style="background:#dde6f5;padding:9px 30px;font-size:10px;font-weight:bold;'
                    'color:#003087;letter-spacing:1.5px;text-transform:uppercase;'
                    f'border-left:5px solid #003087">{esc(name)}</div>')
        for s in stories:
            sub = SUB if s["is_paywalled"] else ""
            similar = ""
            if s.get("similar"):
                items = "".join(
                    '<div style="font-size:12px;line-height:1.5;margin-bottom:4px">'
                    f'<span style="{S_SRC}">RELATED &mdash; {esc(x["source"])}:</span> '
                    f'{extlink(x["url"], x["headline"])}{SUB if x.get("is_paywalled") else ""}</div>'
                    for x in s["similar"]
                )
                similar = ('<div style="background:#f5f8fd;border-left:4px solid #003087;padding:10px 12px;'
                           'margin-top:12px"><div style="font-size:10px;font-weight:bold;color:#003087;'
                           'letter-spacing:1.2px;text-transform:uppercase;margin-bottom:6px">Related Coverage</div>'
                           f'{items}</div>')
            body.append(
                f'<div id="{esc(s["anchor"])}" style="padding:16px 30px 12px;border-bottom:1px solid #e4e9f2">'
                '<div style="font-size:11px;font-weight:bold;color:#003087;text-transform:uppercase;'
                f'letter-spacing:.5px;margin-bottom:4px">{esc(s["source"])}</div>'
                '<div style="font-size:15px;font-weight:bold;color:#111111;line-height:1.4;margin-bottom:9px">'
                f'{extlink(s["url"], s["headline"])}{sub}</div>'
                '<div style="font-size:13px;line-height:1.75;color:#444444">'
                f'{esc(s["summary"])}</div>{similar}</div>'
                '<div style="text-align:right;font-size:10px;padding:3px 30px 8px">'
                f'<a href="#doctop" style="{S_LINK}">&#8593; Back to Top</a></div>'
            )

    logo_src = (getattr(agency, "logo_url", "") or _AGT_LOGO_URL or "").strip()
    logo = (f'<img src="{esc(logo_src)}" alt="{esc(agency.short_name)} logo" '
            'style="max-height:60px;margin:0 0 12px"><br>'
            if logo_src else "")

    return (
        '<!DOCTYPE html>\n<html lang="en">\n<head>\n<meta charset="UTF-8">\n'
        f'<title>{esc(agency.short_name)} Daily News Summary</title>\n'
        '<style>html{scroll-behavior:smooth}</style>\n</head>\n'
        '<body style="margin:0;padding:0;background:#eef1f6">\n'
        '<table role="presentation" width="100%" cellpadding="0" cellspacing="0" '
        'style="background:#eef1f6"><tr><td align="center">\n'
        '<div id="doctop" style="max-width:760px;margin:0 auto;background:#ffffff;'
        'font-family:Arial,Helvetica,sans-serif;font-size:14px;color:#222222">\n'
        '<div style="background:#003087;padding:26px 30px 22px;text-align:center;'
        'border-bottom:5px solid #7eb4ea">'
        + logo +
        '<div style="color:#ffffff;font-size:18px;font-weight:bold;letter-spacing:4px;'
        f'text-transform:uppercase;margin:0 0 8px">{esc(agency.name)}</div>'
        '<div style="color:#ffffff;font-size:28px;font-weight:bold;margin:0 0 8px">Daily News Summary</div>'
        f'<div style="color:#d7e8ff;font-size:15px;font-weight:bold;margin:0 0 12px">{esc(briefing_date)}</div>'
        '<div style="border-top:1px solid #4f7fbd;margin:12px auto;max-width:560px"></div>'
        '<div style="color:#d7e8ff;font-size:12px">Prepared by Alliance Global Tech, Inc. (AGT) | '
        'FCC Daily News Monitoring</div>'
        '</div>\n'
        '<div style="background:#f5f8fd;padding:18px 30px 10px;border-bottom:3px solid #003087">'
        '<div style="font-size:15px;font-weight:bold;color:#003087;margin:0 0 8px">'
        'Today&rsquo;s Wire &mdash; Contents</div>\n'
        + "\n".join(toc) + '\n</div>\n'
        '<div style="background:#003087;color:#ffffff;font-size:11px;font-weight:bold;letter-spacing:2px;'
        'text-transform:uppercase;padding:8px 30px">&#9658;&nbsp; Story Summaries</div>\n'
        + "\n".join(body) + '\n'
        '<div style="background:#003087;padding:14px 30px;text-align:center;font-size:11px;color:#a8c8f0">'
        'Prepared by Alliance Global Tech, Inc. (AGT) | FCC Daily News Monitoring | '
        '<a href="https://agtbi.com" target="_blank" rel="noopener noreferrer" '
        'style="color:#7eb4ea">agtbi.com</a></div>\n'
        '</div>\n</td></tr></table>\n</body>\n</html>'
    )


# ── Editable Word (.docx) export — matches the client's fcc_digest.py input ────
def _docx_hyperlink(paragraph, text, url):
    """Add a real Word hyperlink run to a paragraph (python-docx has no native API)."""
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rPr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve"); r.append(t)
    h.append(r); paragraph._p.append(h)


def _render_agt_docx(agency: AgencyConfig, briefing_date: str, sections) -> bytes:
    """Render the briefing as an editable .docx in the structure the client's
    fcc_digest.py parses: title block, TOC (all section headers + 'SOURCE: Headline'
    with hyperlinks), a 'Story Summaries' split marker, then per-section stories
    with summaries / Similar Stories / 'Back to Top', and the AGT footer."""
    import io
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    def heading(text):
        doc.add_paragraph().add_run(text).bold = True

    def src_headline(source, headline, url, paywalled):
        p = doc.add_paragraph()
        run = p.add_run(f"{(source or 'NEWS').upper()}: "); run.bold = True
        if url:
            _docx_hyperlink(p, headline, url)
        else:
            p.add_run(headline)
        if paywalled:
            p.add_run("  (SUBSCRIPTION REQUIRED)").italic = True

    # Title block (filtered out as header/junk by fcc_digest.py)
    r = doc.add_paragraph().add_run(agency.name); r.bold = True; r.font.size = Pt(16)
    r = doc.add_paragraph().add_run("Daily News Summary"); r.bold = True; r.font.size = Pt(14)
    doc.add_paragraph(briefing_date)
    doc.add_paragraph("Prepared by Alliance Global Tech, Inc. (AGT) | FCC Daily News Monitoring | agtbi.com")
    doc.add_paragraph("")

    # Table of contents — every section header always present
    for name, stories in sections:
        heading(name)
        for st in stories:
            src_headline(st["source"], st["headline"], st["url"], st["is_paywalled"])
    doc.add_paragraph("")

    # Split marker
    r = doc.add_paragraph().add_run("Story Summaries"); r.bold = True; r.font.size = Pt(13)

    # Summaries — every section header always present
    for name, stories in sections:
        heading(name)
        for st in stories:
            src_headline(st["source"], st["headline"], st["url"], st["is_paywalled"])
            if st["summary"]:
                doc.add_paragraph(st["summary"])
            for sim in st.get("similar", []):
                src_headline("RELATED — " + sim.get("source", ""), sim.get("headline", ""),
                             sim.get("url", ""), sim.get("is_paywalled", False))
            doc.add_paragraph("Back to Top")

    doc.add_paragraph("Prepared by Alliance Global Tech, Inc. (AGT) | FCC Daily News Monitoring")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ── Legacy LLM briefing generator (kept for reference; no longer called) ───────
async def _legacy_generate_briefing_html(agency: AgencyConfig, articles: List[Article], briefing_date: str) -> str:
    if not ANTHROPIC_KEY:
        return _simple_html(agency, articles, briefing_date)

    client = _get_client()

    # Group by topic, top 5 per topic.
    # NEWS sections contain NEWS only. Social posts (BlueSky/Reddit/YouTube) must
    # NOT appear as news stories — they feed the Social Media Summary at the end.
    by_topic: Dict[str, List[Article]] = {}
    for art in articles:
        if art.relevance_score >= 0.4 and art.source_type != "social":
            by_topic.setdefault(art.topic, []).append(art)

    context = []
    for topic, arts in sorted(by_topic.items()):
        label = TOPIC_LABELS.get(topic, topic)
        context.append(f"\n=== {label} ===")
        for art in sorted(arts, key=lambda a: a.relevance_score, reverse=True)[:5]:
            clip = f"[BROADCAST CLIP: {art.broadcast_clip_url}]" if art.broadcast_clip_url else ""
            lock = "[SUBSCRIPTION REQUIRED]" if art.is_paywalled else ""
            reg  = "[REGULATORY]" if art.source_type == "regulatory" else ""
            context.append(f"""
TITLE: {art.title}
TYPE: {art.article_type} | SENTIMENT: {art.sentiment} | SCORE: {art.relevance_score:.2f} {reg}
OUTLET: {art.outlet} | AUTHOR: {art.author}
URL: {art.url} {clip} {lock}
SUMMARY: {art.summary[:300]}
---""")

    # Coverage window
    from datetime import datetime as _dt, timedelta as _td
    _today = _dt.now()
    _start = _us_date_short(_today - _td(days=3))
    _window = f"{_start} - {_us_date(_today)}"

    # Real social posts only (BlueSky/Reddit/YouTube) — passed separately so the
    # social summary is built from actual data, never fabricated platforms/numbers.
    total_articles = len(articles)
    social_arts = [a for a in articles if a.source_type == "social"]
    _social_lines = [f"- [{(a.source or a.outlet or 'social')}] {a.title}"
                     for a in sorted(social_arts, key=lambda x: x.relevance_score, reverse=True)[:10]]
    social_context = "\n".join(_social_lines) if _social_lines else "NONE"

    prompt = f"""You are generating the FCC Daily News Briefing exactly matching this official format.

OUTPUT FORMAT — follow this exactly:

===HTML HEADER===
<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8">
<title>{agency.short_name} Daily News Summary — {briefing_date}</title>
<style>
  body {{font-family: Arial, sans-serif; max-width: 760px; margin: 40px auto; padding: 0 20px; color: #333; font-size: 13px; line-height: 1.5;}}
  h1 {{color: #0B3C5D; font-size: 22px; border-bottom: 2px solid #0078D4; padding-bottom: 8px; margin-bottom: 4px;}}
  h2 {{color: #0B3C5D; font-size: 15px; margin-top: 24px; margin-bottom: 6px; padding-bottom: 4px; border-bottom: 1px solid #ddd;}}
  .date {{font-size: 18px; font-weight: bold; color: #0B3C5D; margin-bottom: 0;}}
  .agency {{font-size: 20px; font-weight: bold; color: #0B3C5D;}}
  .toc-item {{margin: 3px 0; padding-left: 0;}}
  .outlet {{font-weight: bold; text-transform: uppercase;}}
  .sub-required {{color: #666; font-style: italic;}}
  .similar {{padding-left: 20px; margin: 2px 0; color: #555; font-size: 12px; list-style: none;}}
  .similar li::before {{content: "▪ "; color: #0078D4;}}
  .story {{margin-top: 18px; padding-bottom: 12px; border-bottom: 1px solid #f0f0f0;}}
  .story-headline {{font-weight: bold; font-size: 13px; margin-bottom: 6px;}}
  .story-body {{color: #333; margin-bottom: 8px;}}
  .backtotop {{font-size: 11px; color: #0078D4; text-decoration: none;}}
  .social-section {{background: #f9f9f9; border: 1px solid #ddd; padding: 14px; margin-top: 24px; border-radius: 4px;}}
  .footer {{margin-top: 30px; padding-top: 12px; border-top: 1px solid #ddd; font-size: 11px; color: #888;}}
</style>
</head>
<body>
<p class="date">{briefing_date}</p>
<p class="agency">{agency.name}</p>
<h1>Daily News Summary</h1>
<p style="color:#666;font-size:12px;">Coverage window: {_window}</p>

===TABLE OF CONTENTS (TOC)===
For EACH topic section that has articles, write:

<h2>[Topic Name]</h2>
<p class="toc-item"><span class="outlet">OUTLET NAME:</span> <a href="#story-[id]">Article Title</a> [add (SUBSCRIPTION REQUIRED) if paywalled]</p>

RULES FOR TOC:
- Outlet name in ALL CAPS followed by colon
- If same story covered by multiple outlets, list PRIMARY outlet only, group others as similar stories in summaries section
- FCC.gov as source: include MAXIMUM 2 items total across entire briefing
- Mark paywalled: (SUBSCRIPTION REQUIRED) in parentheses after title
- Each article gets unique anchor id like #story-1, #story-2 etc.

===STORY SUMMARIES===
After the complete TOC, write the summaries section:

<h2>Story Summaries</h2>

For each PRIMARY story:
<div class="story" id="story-[id]">
<p class="story-headline"><span class="outlet">OUTLET:</span> <a href="URL">Title</a>[if paywalled: <span class="sub-required"> (SUBSCRIPTION REQUIRED)</span>]</p>
[If NOT paywalled: <p class="story-body">2-3 factual sentences summarizing the article.</p>]
[If paywalled: <p class="sub-required">[SUBSCRIPTION REQUIRED]</p>]
[If has similar stories:]
<p style="font-size:12px;font-weight:bold;margin-bottom:3px;">Similar stories:</p>
<ul class="similar">
  <li><span class="outlet">OUTLET:</span> <a href="URL">Title</a></li>
</ul>
<a href="#top" class="backtotop">Back to Top</a>
</div>

===SOCIAL MEDIA SUMMARY===
At the end, write a social media summary section:

<div class="social-section">
<h2>Social Media Summary</h2>
<p>[Summarize social activity using ONLY the real posts in SOCIAL POSTS below. Name ONLY platforms that actually appear there (e.g. BlueSky, YouTube) — NEVER mention X/Twitter, Facebook, or LinkedIn unless a post from it is listed, and NEVER invent reach, engagement, or post-count numbers. If SOCIAL POSTS is NONE, write exactly: "No notable social media activity was captured for this period."]</p>

SOCIAL POSTS:
{social_context}
</div>

===FOOTER===
<div class="footer">
<p>Prepared by Alliance Global Tech, Inc. (AGT) | FCC Daily News Monitoring | agtbi.com</p>
</div>
</body></html>

CRITICAL RULES:
1. NO DUPLICATE STORIES — if multiple outlets cover same story, pick the most authoritative outlet as primary, list others as similar stories
2. FCC.GOV SOURCE — include maximum 2 articles from fcc.gov across the entire briefing
3. SUBSCRIPTION REQUIRED — mark clearly any paywalled article, do not summarize paywalled content
4. FACTUAL ONLY — no editorial opinion, exactly like an official government clipping service
5. OUTLET NAMES — always in ALL CAPS (REUTERS, BROADBAND BREAKFAST, etc.)
6. ALL CAPS for section headers matching FCC taxonomy

Articles to process:
{"".join(context)}

Write the complete HTML document now — exact FCC Daily News Briefing format."""

    try:
        resp = await client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=8000,
            messages=[{"role": "user", "content": prompt}]
        )
        html = _extract_text(resp.content)
        # Ensure it's clean HTML
        if not html.strip().startswith('<!') and not html.strip().startswith('<html'):
            html = f"<!DOCTYPE html>\n{html}"
        return html
    except Exception as e:
        logger.error(f"Briefing generation error: {e}")
        return _simple_html(agency, articles, briefing_date)


def _simple_html(agency: AgencyConfig, articles: List[Article], briefing_date: str) -> str:
    rows = "".join(
        f'<tr><td><a href="{a.url}">{a.title}</a></td><td>{a.outlet}</td><td>{TOPIC_LABELS.get(a.topic, a.topic)}</td></tr>'
        for a in articles[:20]
    )
    return f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<title>{agency.name} Briefing {briefing_date}</title></head>
<body style="font-family:Arial,sans-serif;max-width:800px;margin:auto;padding:20px">
<h1 style="color:{agency.primary_color}">{agency.name} Morning Briefing — {briefing_date}</h1>
<table border="1" cellpadding="8" cellspacing="0" width="100%">{rows}</table>
<p style="color:#666;font-size:12px">Generated by DocuAction AI — Alliance Global Tech, Inc.</p>
</body></html>"""


# ── Email Delivery: SendGrid ───────────────────────────────────────────────────
async def deliver_briefing(agency: AgencyConfig, html: str, briefing_date: str) -> Dict[str, Any]:
    if not SENDGRID_KEY:
        logger.warning("SENDGRID_API_KEY not set — dry run mode")
        return {"status": "dry_run", "recipients": len(agency.distribution_list)}

    subject = f"{agency.short_name} Morning Intelligence Briefing — {briefing_date}"
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.post(
                "https://api.sendgrid.com/v3/mail/send",
                headers={"Authorization": f"Bearer {SENDGRID_KEY}", "Content-Type": "application/json"},
                json={
                    "personalizations": [{"to": [{"email": e} for e in agency.distribution_list]}],
                    "from": {"email": agency.distribution_email, "name": f"{agency.name} Intelligence"},
                    "subject": subject,
                    "content": [{"type": "text/html", "value": html}],
                }
            )
            resp.raise_for_status()
            return {"status": "delivered", "recipients": len(agency.distribution_list), "subject": subject}
    except Exception as e:
        logger.error(f"SendGrid error: {e}")
        return {"status": "error", "error": str(e)}


# ── LLM Visibility Tracker — Multi-AI (UNIQUE FEATURE) ───────────────────────
# No competitor tracks what AI engines say about a federal agency.
# This answers: "When a journalist asks ChatGPT about the FCC, what does it say?"

async def _query_openai(question: str) -> str:
    """Query ChatGPT for LLM visibility check."""
    if not OPENAI_KEY:
        return "OpenAI API key not configured (set OPENAI_API_KEY in Railway)"
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.post(
                "https://api.openai.com/v1/chat/completions",
                headers={"Authorization": f"Bearer {OPENAI_KEY}", "Content-Type": "application/json"},
                json={
                    "model": "gpt-4o-mini",
                    "max_tokens": 400,
                    "messages": [{"role": "user", "content": question}]
                }
            )
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"ChatGPT error: {e}"


async def _query_perplexity(question: str) -> str:
    """Query Perplexity (has real-time web access via sonar model)."""
    if not PERPLEXITY_KEY:
        return "Perplexity API key not configured (set PERPLEXITY_API_KEY in Railway)"
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.post(
                "https://api.perplexity.ai/chat/completions",
                headers={"Authorization": f"Bearer {PERPLEXITY_KEY}", "Content-Type": "application/json"},
                json={
                    "model": "sonar",
                    "max_tokens": 400,
                    "messages": [{"role": "user", "content": question}]
                }
            )
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"Perplexity error: {e}"


async def _query_gemini(question: str) -> str:
    """Query Google Gemini for LLM visibility check."""
    if not GEMINI_KEY:
        return "Gemini API key not configured (set GEMINI_API_KEY in Railway)"
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.post(
                f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={GEMINI_KEY}",
                headers={"Content-Type": "application/json"},
                json={"contents": [{"parts": [{"text": question}]}]}
            )
            resp.raise_for_status()
            return resp.json()["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        return f"Gemini error: {e}"


async def run_llm_visibility_check(agency: AgencyConfig) -> Dict[str, Any]:
    """
    Query all major AI engines to see what they say about the agency.
    Unique feature — no media monitoring vendor offers this.
    Results show: AI engine awareness, accuracy, recency of information.
    """
    questions = [
        f"What is the {agency.name} ({agency.short_name}) currently working on? What are their main priorities?",
        f"What are the most important recent decisions or actions taken by the {agency.name}?",
        f"What controversies, criticisms, or issues is the {agency.name} currently facing?",
        f"Who are the current key leaders and commissioners at the {agency.name}?",
    ]

    results = {}

    for question in questions:
        q_results = {}

        # Claude (existing)
        if ANTHROPIC_KEY:
            try:
                client = _get_client()
                resp = await client.messages.create(
                    model="claude-haiku-4-5", max_tokens=400,
                    messages=[{"role": "user", "content": question}]
                )
                q_results["Claude (Anthropic)"] = _extract_text(resp.content)
            except Exception as e:
                q_results["Claude (Anthropic)"] = f"Error: {e}"

        # ChatGPT
        q_results["ChatGPT (OpenAI)"] = await _query_openai(question)

        # Perplexity
        q_results["Perplexity"] = await _query_perplexity(question)

        # Gemini
        q_results["Gemini (Google)"] = await _query_gemini(question)

        results[question] = q_results

    return {
        "agency_id": agency.agency_id,
        "agency_name": agency.name,
        "checked_at": _now(),
        "engines_queried": ["Claude (Anthropic)", "ChatGPT (OpenAI)", "Perplexity", "Gemini (Google)"],
        "questions_asked": len(questions),
        "results": results,
        "summary": f"Queried 4 AI engines on {len(questions)} questions about {agency.name}. This report shows AI engine awareness and information recency — a unique competitive differentiator unavailable from any other media monitoring vendor.",
    }


# ── Archive search ─────────────────────────────────────────────────────────────
def search_archive(
    agency_id: str,
    keyword: Optional[str] = None,
    topic: Optional[str] = None,
    source_type: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    min_relevance: float = 0.0,
    page: int = 1,
    page_size: int = 50
) -> Dict[str, Any]:
    cutoff = datetime.now(timezone.utc) - timedelta(days=365)
    results = []

    for art in _articles.values():
        if art.agency_id != agency_id:
            continue
        try:
            pub = datetime.fromisoformat(art.published_at.replace("Z", "+00:00"))
            if pub < cutoff:
                continue
            if start_date:
                sd = datetime.fromisoformat(start_date).replace(tzinfo=timezone.utc)
                if pub < sd:
                    continue
            if end_date:
                ed = datetime.fromisoformat(end_date).replace(tzinfo=timezone.utc)
                if pub > ed:
                    continue
        except Exception:
            pass

        if topic and art.topic != topic:
            continue
        if source_type and art.source_type != source_type:
            continue
        if art.relevance_score < min_relevance:
            continue
        if keyword:
            kw = keyword.lower()
            if kw not in art.title.lower() and kw not in art.summary.lower():
                continue
        results.append(art)

    # Robust sort: published_at can be None/empty/non-string for some sources, and
    # mixing those in a sort key raises TypeError (None < str) → a 500 that breaks
    # the whole article list. Coerce to a string so it always orders cleanly.
    results.sort(key=lambda a: str(getattr(a, "published_at", "") or ""), reverse=True)
    total = len(results)
    start = (page - 1) * page_size

    def _safe_dict(a):
        try:
            return asdict(a)
        except Exception:
            return {k: getattr(a, k, None) for k in (
                "article_id", "agency_id", "title", "url", "outlet", "summary",
                "topic", "source_type", "published_at", "relevance_score",
                "is_paywalled", "broadcast_clip_url", "article_type")}

    return {
        "agency_id": agency_id,
        "total": total,
        "page": page,
        "page_size": page_size,
        "total_pages": max(1, (total + page_size - 1) // page_size),
        "articles": [_safe_dict(a) for a in results[start:start + page_size]],
    }


def get_archive_stats(agency_id: str) -> Dict[str, Any]:
    arts = [a for a in _articles.values() if a.agency_id == agency_id]
    by_topic, by_source, by_type, monthly = {}, {}, {}, {}
    for a in arts:
        by_topic[a.topic] = by_topic.get(a.topic, 0) + 1
        by_source[a.source_type] = by_source.get(a.source_type, 0) + 1
        by_type[a.article_type] = by_type.get(a.article_type, 0) + 1
        month = (a.published_at or "")[:7]   # None/empty-safe
        if month:
            monthly[month] = monthly.get(month, 0) + 1
    return {
        "agency_id": agency_id,
        "total_articles": len(arts),
        "archive_months": 12,
        "by_topic": by_topic,
        "by_source_type": by_source,
        "by_article_type": by_type,
        "monthly_volume": dict(sorted(monthly.items())),
    }


async def _ingest_gdelt_doc_articles(agency: AgencyConfig, lookback_hours: int) -> List["Article"]:
    """GDELT DOC 2.0 online news (free, no API key, thousands of outlets) → Articles."""
    try:
        from app.bulletin_intelligence.gdelt_doc_ingest import ingest_gdelt_doc
        raw = await ingest_gdelt_doc(lookback_hours)
        out = [a for a in (_dict_to_article(d, agency.agency_id, "news") for d in raw) if a]
        logger.info(f"GDELT DOC → {len(out)} articles for {agency.agency_id}")
        return out
    except Exception as e:
        logger.warning(f"GDELT DOC ingest failed: {e}")
        return []


async def _ingest_reddit_articles(agency: AgencyConfig) -> List["Article"]:
    """Reddit posts (free; needs REDDIT_CLIENT_ID/SECRET) → Articles. Skips if unset."""
    try:
        from app.bulletin_intelligence.reddit_ingest import ingest_reddit
        raw = await ingest_reddit()
        out = [a for a in (_dict_to_article(d, agency.agency_id, "social") for d in raw) if a]
        if out:
            logger.info(f"Reddit → {len(out)} posts for {agency.agency_id}")
        return out
    except Exception as e:
        logger.warning(f"Reddit ingest failed: {e}")
        return []


async def _ingest_primary_source_articles(agency: AgencyConfig) -> List["Article"]:
    """FCC.gov daily digest/headlines + congressional hearing transcripts (govinfo).
    Free, no paid key (govinfo uses DEMO_KEY). Each carries real description/snippet
    text so the briefing's per-story summarizer has something concrete to work from."""
    try:
        from app.bulletin_intelligence.cspan_fcc_ingest import ingest_primary_sources
        raw = await ingest_primary_sources()
        out = [a for a in (_dict_to_article(d, agency.agency_id, "regulatory") for d in raw) if a]
        if out:
            logger.info(f"Primary sources → {len(out)} items for {agency.agency_id}")
        return out
    except Exception as e:
        logger.warning(f"Primary-source ingest failed: {e}")
        return []


async def _ingest_broadcast_tv_articles(agency: AgencyConfig, lookback_hours: int) -> List["Article"]:
    """GDELT TV 2.0 broadcast closed-caption clips (CNN/Fox/MSNBC/CSPAN/Bloomberg…).
    Free, no key. Each clip's caption snippet becomes the article summary, giving the
    briefing real broadcast coverage instead of the heuristic web_search guess only."""
    if not agency.include_broadcast:
        return []
    try:
        from app.bulletin_intelligence.gdelt_tv_ingest import ingest_broadcast_gdelt
        raw = await ingest_broadcast_gdelt(lookback_hours)
        # GDELT TV's clip matching is loose (it surfaces unrelated history/biz
        # programming), so keep only clips that actually mention the FCC. Cheap,
        # cuts the bulk of the noise, and saves Claude classification cost. Then
        # cap to the most recent 50 (GDELT returns newest-first) so a heavy news
        # day can't flood the classifier — the briefing only shows a handful anyway.
        raw = [d for d in raw if _is_fcc_relevant_v2(d.get('title',''), d.get('summary',''))][:50]
        out = [a for a in (_dict_to_article(d, agency.agency_id, "broadcast") for d in raw) if a]
        if out:
            logger.info(f"GDELT TV → {len(out)} FCC-relevant broadcast clips for {agency.agency_id}")
        return out
    except Exception as e:
        logger.warning(f"GDELT TV ingest failed: {e}")
        return []


# ── Master daily cycle ─────────────────────────────────────────────────────────
async def run_daily_cycle(
    agency_id: str,
    auto_deliver: bool = False,
    lookback_hours: int = 24
) -> Dict[str, Any]:
    agency = get_agency(agency_id)
    if not agency:
        return {"error": f"Agency {agency_id} not registered"}

    # Cost guard: refuse to start if a cycle for this agency is already running
    # (or started within the TTL). Prevents duplicate clicks / overlapping
    # scheduler+manual runs from multiplying Claude API spend.
    now_ts = datetime.now(timezone.utc).timestamp()
    last = _running_cycles.get(agency_id, 0)
    if last and (now_ts - last) < CYCLE_LOCK_TTL:
        logger.info(f"Cycle already in progress for {agency_id}; skipping duplicate")
        return {
            "agency_id": agency_id,
            "status": "already_running",
            "message": "A collection cycle is already in progress for this agency. Please wait for it to finish.",
        }
    _running_cycles[agency_id] = now_ts

    briefing_date = _us_date(datetime.now())
    briefing_id = f"{agency_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    # Phase 1: tag this task's Claude calls with the run id so cost rows attribute to
    # this cycle. Same id instrumentation.record_run() uses, so cost and run history
    # join on run_id. Best-effort; never blocks the cycle.
    try:
        from app.bulletin_intelligence.costs.cost_tracker import set_run_context
        set_run_context(briefing_id, agency_id)
    except Exception:
        pass
    # Phase 2: pick up any operator edits to the Boolean profiles for this cycle.
    # No-op unless BULLETIN_PROFILES_DB_ENABLED=true and the table has rows.
    try:
        from app.bulletin_intelligence.profiles.boolean_profiles import refresh_from_db
        await refresh_from_db(agency_id)
    except Exception:
        pass
    logger.info(f"Daily cycle starting: {agency.name}")

    # RSS ONLY — Appendix B sources, always FCC-relevant
    tasks = [ingest_rss(agency, lookback_hours)]
    # NewsAPI with FCC domain restrictions
    if NEWSAPI_KEY:
        tasks.append(ingest_newsapi(agency, lookback_hours))
    # Tavily for additional FCC coverage
    if TAVILY_KEY:
        tasks.append(ingest_tavily(agency, lookback_hours))
    # NewsAPI.ai (Event Registry) — additive collector. Auto-detected: runs only
    # when NEWSAPI_AI_KEY is set, otherwise skipped gracefully. Same pipeline.
    if NEWSAPI_AI_KEY:
        tasks.append(ingest_newsapi_ai(agency, lookback_hours))
    # Perigon (Phase 3) — Boolean-native news API. Auto-skips when PERIGON_API_KEY
    # is unset, exactly like the other optional collectors, so this is inert until a
    # key is configured. Queries the Phase 2 Boolean profiles verbatim (Perigon
    # speaks the same AND/OR/NOT/quoted-phrase dialect) with language=en&country=us
    # applied server-side, plus a NOT-exclusion for the FC Cincinnati sense of "FCC".
    try:
        from app.bulletin_intelligence.providers.perigon import (
            PERIGON_ENABLED as _PERIGON_ON, ingest_perigon,
        )
        if _PERIGON_ON:
            tasks.append(ingest_perigon(agency, lookback_hours))
    except Exception as e:
        logger.warning(f"Perigon source unavailable: {e}")
    # Claude web_search ingest_news disabled — too noisy/expensive
    # tasks.append(ingest_news(agency, lookback_hours))

    # ── Free broad-coverage sources (no Claude cost; gather ignores failures) ──
    # GDELT DOC 2.0 — thousands of online outlets, free, no key
    tasks.append(_ingest_gdelt_doc_articles(agency, lookback_hours))
    # Primary sources — FCC.gov daily digest/headlines + congressional hearing
    # transcripts (govinfo). Free, no paid key.
    tasks.append(_ingest_primary_source_articles(agency))
    # GDELT TV — broadcast closed-caption transcripts (CNN/Fox/MSNBC/CSPAN/Bloomberg).
    # Free, no key; auto-skips if agency.include_broadcast is False.
    tasks.append(_ingest_broadcast_tv_articles(agency, lookback_hours))
    # BlueSky public search — free social, no key
    try:
        from app.bulletin_intelligence.bluesky_ingest import ingest_bluesky
        tasks.append(ingest_bluesky(agency, lookback_hours, make_article=Article, hasher=_hash, now_iso=_now))
    except Exception as e:
        logger.warning(f"BlueSky source unavailable: {e}")
    # YouTube — free video/broadcast (auto-skips if YOUTUBE_API_KEY is unset)
    try:
        from app.bulletin_intelligence.youtube_ingest import ingest_youtube
        tasks.append(ingest_youtube(agency, lookback_hours, make_article=Article, hasher=_hash, now_iso=_now))
    except Exception as e:
        logger.warning(f"YouTube source unavailable: {e}")
    # Reddit — free social (auto-skips if REDDIT_CLIENT_ID/SECRET unset)
    try:
        from app.bulletin_intelligence.reddit_ingest import ingest_reddit
        tasks.append(_ingest_reddit_articles(agency))
    except Exception as e:
        logger.warning(f"Reddit source unavailable: {e}")

    if agency.include_broadcast:
        tasks.append(ingest_broadcast(agency, lookback_hours))
    if agency.include_social:
        tasks.append(ingest_social(agency))
    if agency.include_regulatory:
        tasks.append(ingest_regulatory(agency))

    gathered = await asyncio.gather(*tasks, return_exceptions=True)
    all_articles = []
    for r in gathered:
        if isinstance(r, list):
            all_articles.extend(r)

    # Stamp provider tracking on every article (idempotent; only fills blanks) so
    # per-provider analytics, coverage comparison, and the export provider column
    # are computed from real stamped data. Additive — no article is filtered here.
    stamp_providers(all_articles)

    # Drop client-excluded outlets (e.g. techdirt.com) from EVERY source, not
    # just RSS — GDELT/NewsAPI/Tavily can surface them too.
    before = len(all_articles)
    # Reporting only: retain the blocked-domain articles for the editor audit.
    _excluded_domain_arts = [a for a in all_articles if _is_excluded_domain(getattr(a, "url", ""))]
    all_articles = [a for a in all_articles if not _is_excluded_domain(getattr(a, "url", ""))]
    if before != len(all_articles):
        logger.info(f"Excluded {before - len(all_articles)} article(s) from blocked domains")

    # Process pipeline. Dedup: opt-in stronger URL+fingerprint pass wraps the
    # existing deduplicate (BULLETIN_BETTER_DEDUP; default OFF → identical
    # behavior). Never removes unique stories; falls back on any error.
    try:
        from app.bulletin_intelligence.coverage_hotfix import better_deduplicate
        unique = better_deduplicate(all_articles, deduplicate)
    except Exception as _e:
        logger.debug(f"better_deduplicate wrapper skipped: {_e}")
        unique = deduplicate(all_articles)

    # Safety cap (ADD-only): the extended source list can surface far more unique
    # articles per cycle, and classify_articles makes one LLM call per 8 of them,
    # sequentially — so cost/time scale with volume. Cap what we hand the classifier
    # (highest cheap-relevance first) so a large fetch can't blow up cost or push the
    # cycle past the scheduler window. Override with BULLETIN_MAX_CLASSIFY.
    _max_classify = int(os.getenv("BULLETIN_MAX_CLASSIFY", "600"))
    to_classify = unique
    if len(unique) > _max_classify:
        to_classify = sorted(
            unique, key=lambda a: getattr(a, "relevance_score", 0) or 0, reverse=True
        )[:_max_classify]
        logger.warning(f"Classify cap: {len(unique)} unique -> {_max_classify} (BULLETIN_MAX_CLASSIFY)")

    try:
        classified = await classify_articles(to_classify, agency)
    except Exception as classify_err:
        logger.error(f"Classification failed: {classify_err}")
        for art in unique:
            art.topic = art.topic or "fcc_news_events"
            art.relevance_score = 0.7
        classified = unique

    # ── Enhancement pass (deterministic, additive — never lowers volume by default) ──
    # 1) Flag known subscription outlets → briefing shows [SUBSCRIPTION REQUIRED].
    try:
        from app.bulletin_intelligence.editorial_rules import flag_subscriptions
        flagged = flag_subscriptions(classified)
    except Exception as _e:
        flagged = 0
        logger.debug(f"flag_subscriptions skipped: {_e}")
    # 2) Boost relevance for clear FCC signals (commissioners, dockets, enforcement).
    boosted = apply_fcc_relevance_boost(classified)
    # 2b) Reject corporate-announcement noise (UAT false-positive fix, e.g. a
    #     "T-Mobile executive appointment") unless it has a real FCC nexus.
    #     Flag-reversible (BULLETIN_EDITORIAL_STRICT=false). Rejected items stay
    #     archived upstream; only the briefing set is trimmed.
    corp_rejected = 0
    _corp = []
    try:
        from app.bulletin_intelligence.editorial_rules import filter_corporate_noise
        classified, _corp = filter_corporate_noise(classified)
        corp_rejected = len(_corp)
        if corp_rejected:
            logger.info(f"Editorial: rejected {corp_rejected} corporate-noise story(ies)")
    except Exception as _e:
        logger.debug(f"filter_corporate_noise skipped: {_e}")
    # ── Coverage hotfix: re-admit any HIGH-PRIORITY story an editorial filter
    #    rejected (AT&T/Robocall/TCPA/… never dropped on low confidence alone).
    #    Additive + flagged (BULLETIN_PRIORITY_PROTECT) + fail-safe. ──
    try:
        from app.bulletin_intelligence.coverage_hotfix import reclaim_priority
        _reclaimed = reclaim_priority(_corp)
        if _reclaimed:
            classified = classified + _reclaimed
    except Exception as _e:
        logger.debug(f"priority reclaim skipped: {_e}")
    # 3) Optional STRICT FCC gate — OFF by default so the minimum-volume floor is
    #    never at risk. When BULLETIN_STRICT_FCC_GATE=true, drop any story with no
    #    explicit FCC mention anywhere in its text.
    if STRICT_FCC_GATE:
        kept = [a for a in classified if _has_fcc_mention(a)]
        logger.info(f"STRICT_FCC_GATE: {len(classified)} -> {len(kept)} FCC-mention-only")
        classified = kept
    logger.info(f"Enhancement pass: {flagged} subscription-flagged, {boosted} relevance-boosted")

    # Store in archive
    for art in classified:
        _articles[art.article_id] = art

    # Filter for briefing. Cap raised 80 -> 150 (2026-06-30) so the delivered
    # briefing can actually carry the higher collection volume (client wants 100+/day).
    # Coverage hotfix: high-priority stories survive the "other + low score"
    # cutoff (additive — only widens the briefing set). Fail-safe.
    try:
        from app.bulletin_intelligence.coverage_hotfix import briefing_keep_ids
        _prio_keep = briefing_keep_ids(classified)
    except Exception:
        _prio_keep = set()
    briefing_arts = sorted(
        [a for a in classified if (getattr(a, "article_id", None) in _prio_keep) or not (a.topic == "other" and a.relevance_score < 0.4)],
        key=lambda a: (a.topic != "other", a.relevance_score), reverse=True
    )[:150]
    logger.info(f"Briefing: {len(briefing_arts)} articles from {len(classified)} classified")

    # Coverage analytics (additive; surfaced at GET /coverage/{agency_id}).
    try:
        coverage = _build_coverage_report(agency_id, all_articles, unique, classified, briefing_arts)
        _last_coverage[agency_id] = coverage
        logger.info(
            f"Coverage: {coverage['stories_collected']} collected, "
            f"{coverage['duplicates_removed']} dupes removed, {coverage['in_briefing']} in briefing, "
            f"{coverage['subscription_stories']} subscription; missing categories: "
            f"{coverage['missing_category_warnings'] or 'none'}"
        )
    except Exception as _e:
        coverage = None
        logger.warning(f"Coverage report failed: {_e}")

    # ── Coverage hotfix: append gap detection + editorial-review queue to the
    #    coverage report and log. Advisory only — flags, never rejects. Fail-safe. ──
    try:
        from app.bulletin_intelligence.coverage_hotfix import build_coverage_extra
        _extra = build_coverage_extra(all_articles, unique, briefing_arts)
        if isinstance(coverage, dict):
            coverage["coverage_hotfix"] = _extra
            _last_coverage[agency_id] = coverage
        logger.info(
            f"Coverage hotfix: {_extra['priority_in_briefing']} priority-protected in briefing, "
            f"{_extra['coverage_gap_count']} coverage gap(s), "
            f"{_extra['editorial_review_count']} flagged for editorial review"
        )
    except Exception as _e:
        logger.warning(f"Coverage hotfix analytics skipped: {_e}")

    # New Source Discovery → Editorial Queue: compare NewsAPI.ai-discovered outlets
    # against the 194-source registry and classify. Advisory only — NEVER auto-imports.
    # Best-effort: registry may be empty (pending Appendix A) → skipped gracefully.
    if coverage is not None and NEWSAPI_AI_KEY:
        try:
            from app.bulletin_intelligence import bulletin_store
            from app.bulletin_intelligence.provider_analysis import compare_against_registry
            registry = await bulletin_store.load_source_registry()
            coverage["registry_editorial_queue"] = compare_against_registry(
                all_articles, registry, target_provider="NewsAPI.ai"
            )
        except Exception as _e:
            logger.debug(f"registry editorial queue skipped: {_e}")

    # Generate briefing — HTML + editable Word (.docx), built from the same sections
    html, docx_bytes, sections = await build_briefing_outputs(agency, briefing_arts, briefing_date)
    import base64 as _b64
    docx_b64 = _b64.b64encode(docx_bytes).decode() if docx_bytes else ""

    # Report what the reader can actually see. article_count/topic_counts used to be
    # taken from briefing_arts (the pre-render candidate set), which is a superset of
    # what _prepare_briefing_sections renders — the header claimed 146 stories over a
    # document containing 41. Count the rendered stories instead. One rendered story =
    # one cluster primary; its RELATED links are additional coverage of that same
    # story, so they are deliberately not counted again here.
    rendered = [s for _sec_name, _stories in sections for s in _stories]
    # Test `sections`, not `rendered`: _collect_sections always returns all nine
    # section tuples (empty headers included), so a truthy `sections` means the AGT
    # render ran. Only the fallback returns []. A successful render with zero stories
    # must therefore report 0, not fall back to the candidate count.
    if sections:
        article_count = len(rendered)
        topic_counts = {}
        for s in rendered:
            _t = s.get("_topic") or "other"
            topic_counts[_t] = topic_counts.get(_t, 0) + 1
    else:
        # AGT render failed -> _simple_html, which lists every briefing article, so
        # briefing_arts IS the visible set on that path.
        article_count = len(briefing_arts)
        topic_counts = {}
        for art in briefing_arts:
            topic_counts[art.topic] = topic_counts.get(art.topic, 0) + 1
    logger.info(
        f"Briefing counts: {article_count} rendered stories "
        f"(from {len(briefing_arts)} briefing candidates)"
    )

    # LIVE-FEED MODEL: every briefing is immediately available the moment it is
    # built — there is NO approval gate. status="delivered" here means "live /
    # viewable", not "emailed". Email is a separate, optional action (see
    # send_briefing_email / POST /send). auto_deliver still exists for the
    # scheduler and also emails the briefing, but an email failure NEVER makes the
    # briefing unavailable — it stays live regardless.
    delivery = {}
    delivered_at = ""
    if auto_deliver:
        delivery = await deliver_briefing(agency, html, briefing_date)
        if delivery.get("status") in ("delivered", "dry_run"):
            delivered_at = _now()

    status = "delivered"
    briefing = Briefing(
        briefing_id=briefing_id,
        agency_id=agency_id,
        briefing_date=briefing_date,
        status=status,
        html_content=html,
        article_count=article_count,
        topic_counts=topic_counts,
        generated_at=_now(),
        delivered_at=delivered_at,
        delivery_recipients=len(agency.distribution_list),
        docx_b64=docx_b64,
    )
    _briefings[briefing_id] = briefing

    # Persist to durable store (best-effort; never blocks/fails the cycle)
    try:
        from app.bulletin_intelligence import bulletin_store
        await bulletin_store.save_articles([asdict(a) for a in classified])
        await bulletin_store.save_briefing(asdict(briefing))
    except Exception as e:
        logger.warning(f"Persist after daily cycle failed: {e}")

    # ── Editor audit (REPORTING ONLY): after the bulletin is finalized, render the
    #    collect / remove-with-reason / high-priority / coverage-gap / editorial-queue
    #    breakdown to FCC_BULLETIN_EDITOR_AUDIT_YYYYMMDD.log. Reads artifacts only —
    #    no filtering/collection/AI change. Fail-safe — never affects the briefing. ──
    try:
        from app.bulletin_intelligence.editor_audit import write_editor_audit
        _uniq_ids = {getattr(a, "article_id", id(a)) for a in unique}
        _brief_ids = {getattr(a, "article_id", id(a)) for a in briefing_arts}
        _removed = {
            "Duplicate": [a for a in all_articles if getattr(a, "article_id", id(a)) not in _uniq_ids],
            "Corporate Noise": list(_corp),
            "Low Confidence": [a for a in classified if getattr(a, "article_id", id(a)) not in _brief_ids],
            "Non-FCC (blocked domain)": list(_excluded_domain_arts),
        }
        write_editor_audit(agency_id, all_articles, unique, classified, briefing_arts, _removed)
    except Exception as _e:
        logger.warning(f"Editor audit skipped: {_e}")

    # Phase 4 — best-effort, flag-gated run instrumentation (never breaks the cycle).
    try:
        from app.bulletin_intelligence.instrumentation import record_run
        _fin = datetime.now(timezone.utc)
        await record_run(
            agency_id, run_id=briefing_id, trigger="cycle",
            started_at=datetime.fromtimestamp(now_ts, timezone.utc).isoformat(),
            finished_at=_fin.isoformat(),
            duration_ms=int((_fin.timestamp() - now_ts) * 1000),
            ingested=len(all_articles), after_dedup=len(unique),
            in_briefing=len(briefing_arts), rejected=max(0, len(classified) - len(briefing_arts)),
            dupes_removed=max(0, len(all_articles) - len(unique)),
            coverage=coverage, sources_scanned=(coverage or {}).get("sources_scanned"),
            status="completed",
        )
    except Exception:
        pass

    # Release the cost guard so a fresh cycle can run once this one is done.
    _running_cycles.pop(agency_id, None)

    return {
        "agency_id": agency_id,
        "briefing_id": briefing_id,
        "briefing_date": briefing_date,
        "status": status,
        "ingested": len(all_articles),
        "after_dedup": len(unique),
        "in_briefing": len(briefing_arts),
        "topic_counts": topic_counts,
        "window": _last_window_stats.get(agency_id, {}),
        "delivery": delivery,
        "coverage_report": coverage,
        "preview_url": _briefing_preview_url(briefing_id),
        "message": f"Briefing LIVE now — view at {_briefing_preview_url(briefing_id)}"
    }


async def approve_and_deliver(briefing_id: str) -> Dict[str, Any]:
    briefing = _briefings.get(briefing_id)
    if not briefing:
        return {"error": "Briefing not found"}
    agency = get_agency(briefing.agency_id)
    if not agency:
        return {"error": "Agency not found"}

    result = await deliver_briefing(agency, briefing.html_content, briefing.briefing_date)
    briefing.status = "delivered" if result.get("status") in ("delivered", "dry_run") else "error"
    briefing.approved_at = _now()
    briefing.delivered_at = _now()

    # Persist the updated status (best-effort)
    try:
        from app.bulletin_intelligence import bulletin_store
        await bulletin_store.save_briefing(asdict(briefing))
    except Exception as e:
        logger.warning(f"Persist after approve failed: {e}")

    return {"briefing_id": briefing_id, "status": briefing.status, "delivery": result}


def get_editorial_queue(agency_id: str) -> List[Dict]:
    return [asdict(b) for b in _briefings.values() if b.agency_id == agency_id and b.status == "pending_approval"]


def get_briefing(briefing_id: str) -> Optional[Dict]:
    b = _briefings.get(briefing_id)
    return asdict(b) if b else None


def get_briefing_html(briefing_id: str) -> Optional[str]:
    b = _briefings.get(briefing_id)
    return b.html_content if b else None


def get_briefing_history(agency_id: str) -> List[Dict]:
    """All briefings for an agency (any status), newest first, without HTML payload."""
    items = [b for b in _briefings.values() if b.agency_id == agency_id]
    items.sort(key=lambda b: b.generated_at or "", reverse=True)
    history = []
    for b in items:
        d = asdict(b)
        d.pop("html_content", None)
        history.append(d)
    return history


# ── Live-feed access + summary email ───────────────────────────────────────────
# Public base URL used to build absolute preview links (in API responses and in
# the "VIEW FULL BRIEFING" email button). Override per-environment if needed.
PUBLIC_BASE_URL = os.getenv("BULLETIN_PUBLIC_BASE_URL", "https://api-prod.docuaction.io").rstrip("/")
# From-address for the summary send (Step 10). Must be a SendGrid-verified sender.
SEND_FROM_EMAIL = os.getenv("BULLETIN_SEND_FROM", "news@agtbi.com")


def _briefing_preview_url(briefing_id: str) -> str:
    """Absolute URL to the full HTML preview of a specific briefing."""
    return f"{PUBLIC_BASE_URL}/api/v1/bulletin/briefings/{briefing_id}/preview"


def _latest_preview_url(agency_id: str) -> str:
    """Stable, bookmarkable URL that always redirects to the newest briefing."""
    return f"{PUBLIC_BASE_URL}/api/v1/bulletin/latest/{agency_id}/preview"


def get_latest_briefing(agency_id: str) -> Optional[Dict]:
    """Most recent briefing for an agency (any status), without the HTML payload."""
    items = [b for b in _briefings.values() if b.agency_id == agency_id]
    if not items:
        return None
    items.sort(key=lambda b: (b.generated_at or "", b.briefing_id), reverse=True)
    d = asdict(items[0])
    d.pop("html_content", None)
    return d


def get_today_briefing(agency_id: str) -> Optional[Dict]:
    """Newest briefing generated *today* (matches briefing_id date prefix, the
    same clock the scheduler/watchdog use), or None if none exists yet."""
    today = datetime.now().strftime("%Y%m%d")
    prefix = f"{agency_id}_{today}"
    todays = [b for b in _briefings.values()
              if b.agency_id == agency_id and b.briefing_id.startswith(prefix)]
    if not todays:
        return None
    todays.sort(key=lambda b: (b.generated_at or "", b.briefing_id), reverse=True)
    d = asdict(todays[0])
    d.pop("html_content", None)
    return d


def _build_summary_email_html(agency: AgencyConfig, briefing: Briefing, preview_url: str) -> str:
    """Compact summary email: headline count + top topics + a 'VIEW FULL BRIEFING'
    button that links to the full hosted preview. Intentionally lightweight so it
    renders in any mail client; the full report lives at preview_url."""
    color = agency.primary_color or "#0B3C5D"
    tc = briefing.topic_counts or {}
    top = sorted(tc.items(), key=lambda kv: kv[1], reverse=True)[:6]
    rows = "".join(
        f'<li style="margin:4px 0;color:#333;font-size:14px">'
        f'{k.replace("_", " ").title()} — <strong>{v}</strong></li>'
        for k, v in top
    ) or '<li style="color:#666">See the full briefing for today\'s coverage.</li>'
    return f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0"></head>
<body style="margin:0;padding:0;background:#f4f5f7;font-family:Arial,Helvetica,sans-serif">
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="background:#f4f5f7;padding:24px 0">
<tr><td align="center">
<table role="presentation" width="600" cellpadding="0" cellspacing="0" style="max-width:600px;background:#ffffff;border-radius:8px;overflow:hidden;border:1px solid #e5e7eb">
<tr><td style="background:{color};padding:20px 28px">
<div style="color:#ffffff;font-size:20px;font-weight:bold">{agency.name}</div>
<div style="color:#cbd5e1;font-size:14px;margin-top:2px">Daily Intelligence Briefing — {briefing.briefing_date}</div>
</td></tr>
<tr><td style="padding:24px 28px">
<p style="margin:0 0 12px;color:#111;font-size:16px">
<strong>{briefing.article_count}</strong> stories in today's briefing.</p>
<ul style="margin:0 0 20px;padding-left:20px">{rows}</ul>
<table role="presentation" cellpadding="0" cellspacing="0"><tr><td align="center"
style="border-radius:6px;background:{color}">
<a href="{preview_url}" target="_blank"
style="display:inline-block;padding:14px 32px;color:#ffffff;font-size:16px;font-weight:bold;text-decoration:none;border-radius:6px">
VIEW FULL BRIEFING →</a>
</td></tr></table>
<p style="margin:20px 0 0;color:#888;font-size:12px">
Or open it directly: <a href="{preview_url}" style="color:{color}">{preview_url}</a></p>
</td></tr>
<tr><td style="padding:16px 28px;background:#f9fafb;border-top:1px solid #e5e7eb;color:#9ca3af;font-size:12px">
Generated by DocuAction AI — Alliance Global Tech, Inc.</td></tr>
</table></td></tr></table></body></html>"""


async def send_briefing_email(briefing_id: str, *,
                              recipients: Optional[List[str]] = None) -> Dict[str, Any]:
    """Email a summary of one briefing (short summary + VIEW FULL BRIEFING button)
    from SEND_FROM_EMAIL to the agency's distribution list. SEPARATE from
    collection — collection makes a briefing live; this is an explicit send.
    Records delivered_at/recipients and re-persists on success."""
    briefing = _briefings.get(briefing_id)
    if not briefing:
        return {"error": "Briefing not found"}
    agency = get_agency(briefing.agency_id)
    if not agency:
        return {"error": "Agency not found"}

    to_list = recipients if recipients else list(agency.distribution_list)
    if not to_list:
        return {"error": "No recipients configured for this agency"}

    preview_url = _briefing_preview_url(briefing_id)
    html = _build_summary_email_html(agency, briefing, preview_url)
    subject = f"{agency.short_name} Daily Briefing — {briefing.briefing_date}"

    if not SENDGRID_KEY:
        logger.warning("SENDGRID_API_KEY not set — send_briefing_email dry run")
        return {"status": "dry_run", "recipients": len(to_list), "preview_url": preview_url,
                "subject": subject, "from": SEND_FROM_EMAIL}
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.post(
                "https://api.sendgrid.com/v3/mail/send",
                headers={"Authorization": f"Bearer {SENDGRID_KEY}", "Content-Type": "application/json"},
                json={
                    "personalizations": [{"to": [{"email": e} for e in to_list]}],
                    "from": {"email": SEND_FROM_EMAIL, "name": f"{agency.short_name} Daily Briefing"},
                    "subject": subject,
                    "content": [{"type": "text/html", "value": html}],
                },
            )
            resp.raise_for_status()
    except Exception as e:
        logger.error(f"send_briefing_email SendGrid error: {e}")
        return {"status": "error", "error": str(e), "preview_url": preview_url}

    briefing.delivered_at = _now()
    briefing.delivery_recipients = len(to_list)
    try:
        from app.bulletin_intelligence import bulletin_store
        await bulletin_store.save_briefing(asdict(briefing))
    except Exception as e:
        logger.warning(f"Persist after send failed: {e}")
    return {"status": "delivered", "recipients": len(to_list), "preview_url": preview_url,
            "subject": subject, "from": SEND_FROM_EMAIL}


# ── Pre-register FCC ───────────────────────────────────────────────────────────
register_agency(AgencyConfig(
    agency_id="fcc",
    name="Federal Communications Commission",
    short_name="FCC",
    primary_color="#0B3C5D",
    search_queries=[
        "FCC Federal Communications Commission spectrum broadband",
        "net neutrality internet access broadband deployment FCC",
        "media ownership broadcast license radio television FCC",
        "wireless 5G spectrum auction mobile telecommunications",
        "artificial intelligence AI telecom media FCC regulation",
        "public safety E911 emergency communications FCC",
        "robocalls TCPA consumer protection FCC",
        "satellite space communications SpaceX FCC",
    ],
    topics=FCC_TOPICS,
    distribution_email="intelligence@docuaction.io",
    distribution_list=["imran@agtbi.com"],
    delivery_time_et="07:30",
    archive_months=12,
))
