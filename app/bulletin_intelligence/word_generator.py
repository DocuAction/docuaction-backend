"""FCC Daily News Summary as a Word document.

WHY THIS EXISTS

The briefing is reviewed as an Excel workbook, cleaned up, and uploaded back.
What then has to leave the building is a Word document in the FCC's format,
pasted into an email. Producing that by hand from a spreadsheet is where
formatting drifts and stories get dropped.

WHAT IT READS

The briefing's articles, rehydrated by the same helper every other export uses
(`bulletin_download_routes._briefing_articles`). That matters for Task 2.3: a
reviewed workbook uploaded through POST /upload-reviewed/{id} applies its edits
to those same article records, so reading them here yields the reviewed text
without a second parsing path. One rehydration helper for Excel, HTML, email and
Word is the only way those four stay in agreement about which stories are in a
briefing.

FORMATTING NOTES WORTH KEEPING

python-docx has no hyperlink API. A link is a `w:hyperlink` element carrying a
relationship id, built here in `_add_hyperlink` — without it the headlines are
plain text and the document is useless as something to click through from.

"Page X of Y" is likewise not a string we can compute: the page count is not
known until Word paginates. The footer embeds PAGE and NUMPAGES field codes,
which Word evaluates on open.

Everything scraped from a feed is passed through `strip_html` before it reaches
a run. A stray tag in a delivered federal document is the visible failure; the
invisible one is markup surviving into text somebody pastes elsewhere.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

NAVY = RGBColor(0x00, 0x30, 0x87)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
NAVY_HEX = "003087"

BODY_FONT = "Calibri"
BODY_PT = 11
HEADING_PT = 14
SOURCE_PT = 10
TITLE_PT = 20

CONTRACT = "273FCC26F0061"
ORGANIZATION = "Alliance Global Tech, Inc."


def _strip(value: Any) -> str:
    """Plain text. Feed markup must never reach a run."""
    try:
        from app.bulletin_intelligence.excel_export import strip_html

        return strip_html(value)
    except Exception:  # pragma: no cover - defensive
        return str(value or "")


def _field(article: Any, name: str, default: str = "") -> Any:
    if isinstance(article, dict):
        return article.get(name, default)
    return getattr(article, name, default)


def _relevance_label(article: Any) -> str:
    try:
        score = float(_field(article, "relevance_score", 0) or 0)
    except (TypeError, ValueError):
        return "Low"
    if score >= 0.75:
        return "High"
    return "Medium" if score >= 0.45 else "Low"


def _add_hyperlink(paragraph, url: str, text: str, *, bold: bool = True,
                   size_pt: int = BODY_PT, color: RGBColor = NAVY):
    """A real clickable hyperlink.

    python-docx exposes no API for this, so the w:hyperlink element and its
    relationship are constructed directly. A headline that is not clickable
    makes the document a dead end for the reader.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)

    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), NAVY_HEX)
    props.append(color_el)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)

    if bold:
        props.append(OxmlElement("w:b"))

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(size_pt * 2))       # half-points
    props.append(size)

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    props.append(fonts)

    run.append(props)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    link.append(run)
    paragraph._p.append(link)
    return link


def _add_field(paragraph, instruction: str) -> None:
    """Embed a Word field code (PAGE, NUMPAGES).

    The page count does not exist until Word paginates, so "Page 2 of 7" cannot
    be written as a string here. Word evaluates these on open.
    """
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _bottom_border(paragraph, color: str = NAVY_HEX, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


class BulletinWordGenerator:
    """Generate the FCC Daily News Summary Word document."""

    def __init__(self, *, contract: str = CONTRACT,
                 organization: str = ORGANIZATION,
                 agency_name: str = "FCC"):
        self.contract = contract
        self.organization = organization
        self.agency_name = agency_name

    # ── grouping ─────────────────────────────────────────────────────────────

    @staticmethod
    def group_by_category(articles: Iterable[Any],
                          section_of=None) -> List[Tuple[str, List[Any]]]:
        """Categories ordered largest first, ties broken alphabetically.

        Deterministic ordering matters: two runs over the same briefing must
        produce the same document, or a reviewer cannot tell an edit from a
        reshuffle.
        """
        grouped: Dict[str, List[Any]] = {}
        for article in articles or []:
            name = None
            if section_of is not None:
                try:
                    name = section_of(article)
                except Exception:
                    name = None
            name = str(name or _field(article, "section", "")
                       or _field(article, "topic", "") or "Other")
            grouped.setdefault(name, []).append(article)
        return sorted(grouped.items(), key=lambda kv: (-len(kv[1]), kv[0].lower()))

    # ── document ─────────────────────────────────────────────────────────────

    def _setup(self, document) -> None:
        style = document.styles["Normal"]
        style.font.name = BODY_FONT
        style.font.size = Pt(BODY_PT)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.15
        paragraph_format.space_after = Pt(6)

        for section in document.sections:
            section.page_width = Pt(8.5 * 72)
            section.page_height = Pt(11 * 72)
            section.top_margin = Pt(72)
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)

    def _header(self, document, briefing_date: str, count: int) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{self.agency_name} Daily News Summary")
        run.font.size = Pt(TITLE_PT)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run.font.name = BODY_FONT

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = "Article" if count == 1 else "Articles"
        sub_run = subtitle.add_run(f"{briefing_date}  |  {count} {label}")
        sub_run.font.size = Pt(BODY_PT)
        sub_run.font.color.rgb = GRAY
        sub_run.font.name = BODY_FONT
        _bottom_border(subtitle)

    def _table_of_contents(self, document, grouped) -> None:
        heading = document.add_paragraph()
        heading_run = heading.add_run("Contents")
        heading_run.font.size = Pt(HEADING_PT)
        heading_run.font.bold = True
        heading_run.font.color.rgb = NAVY
        heading_run.font.name = BODY_FONT

        for name, items in grouped:
            line = document.add_paragraph()
            line.paragraph_format.space_after = Pt(2)
            dots = "." * max(3, 52 - len(name))
            run = line.add_run(f"{name} {dots} ({len(items)})")
            run.font.size = Pt(BODY_PT)
            run.font.name = BODY_FONT

    def _category_heading(self, document, name: str, count: int) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(14)
        plural = "article" if count == 1 else "articles"
        run = paragraph.add_run(f"{name} ({count} {plural})")
        run.font.size = Pt(HEADING_PT)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run.font.name = BODY_FONT
        _bottom_border(paragraph)

    def _article(self, document, article: Any) -> None:
        title = _strip(_field(article, "title"))
        kind = str(_field(article, "article_type", "") or "").lower()
        if kind in ("opinion", "editorial") and "[Opinion]" not in title:
            title = f"{title} [Opinion]"
        paywalled = bool(_field(article, "is_paywalled", False))
        if paywalled and "[Subscription Required]" not in title:
            title = f"{title} [Subscription Required]"

        url = _strip(_field(article, "url"))
        headline = document.add_paragraph()
        headline.paragraph_format.space_after = Pt(2)
        if url.startswith(("http://", "https://")):
            _add_hyperlink(headline, url, title)
        else:
            # A bare run, not an empty link: an <a> with no href renders as a
            # dead link that looks clickable.
            run = headline.add_run(title)
            run.font.bold = True
            run.font.color.rgb = NAVY
            run.font.name = BODY_FONT

        summary = _strip(_field(article, "summary"))
        if summary:
            body = document.add_paragraph()
            body.paragraph_format.space_after = Pt(2)
            run = body.add_run(summary)
            run.font.size = Pt(BODY_PT)
            run.font.color.rgb = BLACK
            run.font.name = BODY_FONT

        source = _strip(_field(article, "outlet")) or _strip(_field(article, "source"))
        meta = document.add_paragraph()
        run = meta.add_run(f"{source}  |  {_relevance_label(article)} Relevance")
        run.font.size = Pt(SOURCE_PT)
        run.font.italic = True
        run.font.color.rgb = GRAY
        run.font.name = BODY_FONT

    def _footer(self, document) -> None:
        """Footer on every page. Word repeats the section footer automatically,
        so this is written once per section rather than per page."""
        for section in document.sections:
            paragraph = section.footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lead = paragraph.add_run(
                f"Prepared by {self.organization}  |  Contract {self.contract}  |  Page ")
            lead.font.size = Pt(9)
            lead.font.color.rgb = GRAY
            lead.font.name = BODY_FONT
            _add_field(paragraph, " PAGE ")
            middle = paragraph.add_run(" of ")
            middle.font.size = Pt(9)
            middle.font.color.rgb = GRAY
            middle.font.name = BODY_FONT
            _add_field(paragraph, " NUMPAGES ")

    # ── entry points ─────────────────────────────────────────────────────────

    def build(self, articles: Iterable[Any], *, briefing_date: str,
              section_of=None) -> bytes:
        """Render the document. Returns .docx bytes."""
        rows = list(articles or [])
        grouped = self.group_by_category(rows, section_of)

        document = Document()
        self._setup(document)
        self._header(document, briefing_date, len(rows))

        if not rows:
            # An empty briefing must say so. A header with nothing under it
            # reads as a broken export rather than a quiet day.
            note = document.add_paragraph()
            run = note.add_run("No articles met the criteria for this briefing.")
            run.font.size = Pt(BODY_PT)
            run.font.name = BODY_FONT
            self._footer(document)
            return self._to_bytes(document)

        self._table_of_contents(document, grouped)
        for name, items in grouped:
            self._category_heading(document, name, len(items))
            for article in items:
                self._article(document, article)

        self._footer(document)
        return self._to_bytes(document)

    @staticmethod
    def _to_bytes(document) -> bytes:
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def generate(self, briefing_id: str) -> bytes:
        """Build the document for one briefing.

        Reads through the shared rehydration helper, so a reviewed workfile
        applied via POST /upload-reviewed/{id} is reflected here without a
        second parsing path.
        """
        from app.bulletin_intelligence.bulletin_download_routes import _briefing_articles

        briefing, agency, articles = _briefing_articles(briefing_id)
        try:
            from app.bulletin_intelligence.engine import _section_of

            section_of = _section_of
        except Exception:  # pragma: no cover
            section_of = None

        self.agency_name = getattr(agency, "short_name", self.agency_name)
        return self.build(
            articles,
            briefing_date=str(briefing.get("briefing_date") or ""),
            section_of=section_of)


def filename_for(briefing_date: str, agency_name: str = "FCC") -> str:
    """FCC_Bulletin_Aug09_2026.docx — a name that sorts and reads."""
    stamp = None
    for fmt in ("%B %d, %Y", "%Y-%m-%d", "%b %d, %Y"):
        try:
            stamp = datetime.strptime(str(briefing_date).strip(), fmt)
            break
        except (ValueError, TypeError):
            continue
    suffix = stamp.strftime("%b%d_%Y") if stamp else "briefing"
    return f"{agency_name}_Bulletin_{suffix}.docx"
