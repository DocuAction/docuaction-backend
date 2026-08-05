"""
DocuAction Bulletin Intelligence — Download Routes
Generates downloadable Word documents for FCC bulletins.
Supports: 1, 2, 3, 4, 5, 7, 30 day ranges.
Add to existing routes.py or register separately.
"""
import io
import re
import logging
from datetime import datetime, timedelta, timezone
from fastapi import APIRouter, HTTPException, Query

from .auth import guard
from fastapi.responses import StreamingResponse

try:
    from docx.oxml.ns import qn
except ImportError:
    qn = None

logger = logging.getLogger("docuaction.bulletin.download")
router = APIRouter(prefix="/api/v1/bulletin", tags=["Bulletin Downloads"])

ALLOWED_DAYS = [1, 2, 3, 4, 5, 7, 14, 30, 60, 90, 180]


def _us_date(dt) -> str:
    """US long date, no leading zero on the day: 'July 7, 2026' (never 'July 07'
    or '7 July 2026'). Portable — avoids non-portable %-d / %#d flags."""
    try:
        return f"{dt.strftime('%B')} {dt.day}, {dt.year}"
    except Exception:
        return str(dt)


def _us_date_short(dt) -> str:
    """US month + day, no year, no leading zero: 'July 7'."""
    try:
        return f"{dt.strftime('%B')} {dt.day}"
    except Exception:
        return str(dt)

# Terms that signal a genuine FCC connection (agency + current leadership).
# An article must mention one of these OR score highly with the classifier to be
# shown — this filters out tangential tech/telecom news that never involves the FCC.
FCC_TERMS = (
    "fcc", "federal communications commission", "f.c.c.",
    "brendan carr", "anna gomez", "olivia trusty", "geoffrey starks", "nathan simington",
)


def _is_valid_article(art):
    """Return True only for real, relevant FCC articles."""
    url = (getattr(art, 'url', '') or '').lower()
    title = (getattr(art, 'title', '') or '').lower()
    summary = (getattr(art, 'summary', '') or '').lower()
    relevance = getattr(art, 'relevance_score', 0) or 0
    topic = getattr(art, 'topic', '') or ''

    # Client-excluded outlets (e.g. techdirt.com) — hide even if already stored.
    try:
        from app.bulletin_intelligence.engine import _is_excluded_domain
        if _is_excluded_domain(url):
            return False
    except Exception:
        if 'techdirt.com' in url:
            return False

    # Remove demo articles
    if 'example.com' in url:
        return False
    if '[demo]' in title:
        return False
    if 'demonstration article' in summary:
        return False

    # Exclude social-media posts (BlueSky / Reddit / YouTube). The bulletin is
    # NEWS only; social belongs in the separate social summary, not the article
    # list. Matches the briefing pipeline, which already drops source_type=='social'.
    if (getattr(art, 'source_type', '') or '').lower() == 'social':
        return False

    # Remove low-relevance noise
    if relevance < 0.4:
        return False

    # Remove low-relevance "other" category (FIFA, Audi, drownings, etc.)
    if topic == 'other' and relevance < 0.6:
        return False

    # Require a real FCC connection: the article must mention the agency or its
    # leadership, OR earn a high relevance score (trust the classifier for clear
    # regulatory items that may not name the FCC verbatim).
    text = title + ' ' + summary
    if not any(term in text for term in FCC_TERMS) and relevance < 0.7:
        return False

    return True


def _parse_dt(raw):
    """Parse a publish date that may be ISO 8601 OR RFC 2822 (RSS feeds use the
    latter, e.g. 'Wed, 17 Jun 2026 12:00:00 GMT'). Returns a datetime or None."""
    if isinstance(raw, datetime):
        return raw
    s = str(raw).strip()
    if not s:
        return None
    try:
        return datetime.fromisoformat(s.replace('Z', '+00:00'))
    except Exception:
        pass
    try:
        from email.utils import parsedate_to_datetime
        return parsedate_to_datetime(s)
    except Exception:
        return None


def _published_within(art, cutoff):
    """True if the article was published on/after cutoff.

    Filters on publish date (when the news was published), not ingest date
    (when we collected it), so a briefing only contains articles actually
    published within the selected window. Articles with no parseable date
    are kept rather than silently dropped.
    """
    raw = getattr(art, 'published_at', '') or getattr(art, 'ingested_at', '') or ''
    if not raw:
        return True
    art_date = _parse_dt(raw)
    if art_date is None:
        return True
    if art_date.tzinfo is None:
        art_date = art_date.replace(tzinfo=timezone.utc)
    return art_date >= cutoff


@router.get("/download/{agency_id}", dependencies=guard("viewer"))
async def download_bulletin(
    agency_id: str,
    days: int = Query(1, description="Number of days: 1, 2, 3, 4, 5, 7, or 30"),
):
    """
    Download a formatted Word document bulletin for the specified time period.
    Articles are organized by topic with clickable links to source URLs.
    """
    if days not in ALLOWED_DAYS:
        raise HTTPException(400, f"Invalid days parameter. Allowed: {ALLOWED_DAYS}")

    try:
        from app.bulletin_intelligence.engine import _articles, _agencies
    except ImportError:
        raise HTTPException(500, "Bulletin engine not available")

    agency = _agencies.get(agency_id)
    if not agency:
        raise HTTPException(404, f"Agency {agency_id} not found")

    cutoff = datetime.now(timezone.utc) - timedelta(days=days)
    filtered = []
    for art in _articles.values():
        if art.agency_id != agency_id:
            continue
        if not _is_valid_article(art):
            continue
        if _published_within(art, cutoff):
            filtered.append(art)

    if not filtered:
        raise HTTPException(404, f"No articles found for the last {days} day(s)")

    topics = {}
    topic_labels = {
        "fcc_news_events": "General",
        "consumers_advocacy": "Consumers & Advocacy",
        "media_broadcasting": "Media & Broadcasting",
        "public_safety_emergency": "Public Safety & Emergency",
        "wireless_mobile": "Wireless & Mobile",
        "ai_emerging_tech": "AI & Emerging Technology",
        "business_industry": "Business & Industry",
        "international_affairs": "International Affairs",
        "space_communications": "Space & Communications",
        "spectrum_policy": "Spectrum Policy",
    }
    topic_colors = {
        "fcc_news_events": "2563EB",
        "consumers_advocacy": "D97706",
        "media_broadcasting": "DC2626",
        "public_safety_emergency": "EF4444",
        "wireless_mobile": "8B5CF6",
        "ai_emerging_tech": "0D9488",
        "business_industry": "0F172A",
        "international_affairs": "6366F1",
        "space_communications": "3B82F6",
        "spectrum_policy": "7C3AED",
    }

    for art in filtered:
        topic = getattr(art, 'topic', 'other') or 'other'
        if topic not in topics:
            topics[topic] = []
        topics[topic].append(art)

    for topic in topics:
        topics[topic].sort(key=lambda a: getattr(a, 'relevance_score', 0) or 0, reverse=True)

    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise HTTPException(500, "python-docx not installed. Add 'python-docx' to requirements.txt")

    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("FCC DAILY INTELLIGENCE BRIEFING")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    run.font.bold = True

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run(agency.name)
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now = datetime.now(timezone.utc)
    if days == 1:
        date_text = f"{_us_date(now)} | Last 24 Hours"
    else:
        start = _us_date_short(now - timedelta(days=days))
        end = _us_date(now)
        date_text = f"{start} — {end} | Last {days} Days"
    run3 = date_para.add_run(date_text)
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    stats = doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = stats.add_run(f"{len(filtered)} Articles | {len(topics)} Topics | AI-Classified")
    run4.font.size = Pt(8)
    run4.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run4.font.italic = True

    doc.add_paragraph()

    idx_title = doc.add_paragraph()
    run_idx = idx_title.add_run("TOPIC INDEX")
    run_idx.font.size = Pt(12)
    run_idx.font.bold = True
    run_idx.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    for topic_key, articles in sorted(topics.items(), key=lambda x: (x[0] != "fcc_news_events", -len(x[1]))):
        label = topic_labels.get(topic_key, topic_key.replace('_', ' ').title())
        idx_line = doc.add_paragraph()
        idx_line.paragraph_format.space_after = Pt(2)
        run_dot = idx_line.add_run("■ ")
        color_hex = topic_colors.get(topic_key, "333333")
        run_dot.font.color.rgb = RGBColor(int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
        run_dot.font.size = Pt(10)
        run_label = idx_line.add_run(f"{label} ({len(articles)} articles)")
        run_label.font.size = Pt(10)
        run_label.font.bold = True
        run_label.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_page_break()

    for topic_key, articles in sorted(topics.items(), key=lambda x: (x[0] != "fcc_news_events", -len(x[1]))):
        label = topic_labels.get(topic_key, topic_key.replace('_', ' ').title())
        color_hex = topic_colors.get(topic_key, "333333")

        header_para = doc.add_paragraph()
        header_para.paragraph_format.space_before = Pt(12)
        header_para.paragraph_format.space_after = Pt(8)
        run_h = header_para.add_run(f"━━━  {label.upper()}  ━━━")
        run_h.font.size = Pt(13)
        run_h.font.bold = True
        run_h.font.color.rgb = RGBColor(int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))

        for i, art in enumerate(articles):
            title_text = getattr(art, 'title', '') or 'Untitled'
            url = getattr(art, 'url', '') or ''
            outlet = getattr(art, 'outlet', '') or ''
            summary = getattr(art, 'summary', '') or ''
            pub_date = getattr(art, 'published_at', '') or ''
            relevance = getattr(art, 'relevance_score', 0) or 0

            title_text = title_text.split(' - ')[0].split(' | ')[0].strip()
            if len(title_text) > 120:
                title_text = title_text[:117] + '...'

            if summary:
                summary = summary.replace('\n', ' ').strip()
                summary = re.sub(r'<[^>]+>', '', summary)
                if len(summary) > 300:
                    cut = summary[:300].rfind('.')
                    if cut > 100:
                        summary = summary[:cut + 1]
                    else:
                        summary = summary[:300] + '...'

            art_para = doc.add_paragraph()
            art_para.paragraph_format.space_before = Pt(6)
            art_para.paragraph_format.space_after = Pt(2)
            run_num = art_para.add_run(f"{i + 1}. ")
            run_num.font.size = Pt(10)
            run_num.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

            if url:
                hyperlink = _add_hyperlink(art_para, url, title_text)
            else:
                run_title = art_para.add_run(title_text)
                run_title.font.size = Pt(10)
                run_title.font.bold = True
                run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

            if outlet or pub_date:
                source_para = doc.add_paragraph()
                source_para.paragraph_format.space_after = Pt(2)
                source_para.paragraph_format.left_indent = Cm(0.5)
                source_text = f"{outlet}"
                if pub_date:
                    try:
                        if 'T' in str(pub_date):
                            dt = datetime.fromisoformat(str(pub_date).replace('Z', '+00:00'))
                            source_text += f" | {_us_date(dt)}"
                        else:
                            source_text += f" | {pub_date}"
                    except Exception:
                        source_text += f" | {pub_date}"
                if relevance > 0:
                    source_text += f" | Relevance: {int(relevance * 100)}%"
                run_src = source_para.add_run(source_text)
                run_src.font.size = Pt(8)
                run_src.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
                run_src.font.italic = True

            if summary:
                sum_para = doc.add_paragraph()
                sum_para.paragraph_format.space_after = Pt(8)
                sum_para.paragraph_format.left_indent = Cm(0.5)
                run_sum = sum_para.add_run(summary)
                run_sum.font.size = Pt(9)
                run_sum.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        doc.add_paragraph()

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(20)
    run_f1 = footer_para.add_run("Prepared by Bulletin Intelligence — Powered by DocuAction AI\n")
    run_f1.font.size = Pt(8)
    run_f1.font.bold = True
    run_f1.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run_f2 = footer_para.add_run("Alliance Global Tech, Inc. | AI-Generated Classification | Human Review Required")
    run_f2.font.size = Pt(7)
    run_f2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    if days == 1:
        fname = f"FCC_Briefing_{now.strftime('%Y%m%d')}.docx"
    else:
        fname = f"FCC_Briefing_{days}day_{now.strftime('%Y%m%d')}.docx"

    logger.info(f"Bulletin download: agency={agency_id} days={days} articles={len(filtered)} topics={len(topics)}")

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


@router.get("/download-options/{agency_id}", dependencies=guard("viewer"))
async def download_options(agency_id: str):
    """Show available download options with article counts per time period."""
    try:
        from app.bulletin_intelligence.engine import _articles, _agencies
    except ImportError:
        raise HTTPException(500, "Bulletin engine not available")

    agency = _agencies.get(agency_id)
    if not agency:
        raise HTTPException(404, f"Agency {agency_id} not found")

    now = datetime.now(timezone.utc)
    options = []

    for d in ALLOWED_DAYS:
        cutoff = now - timedelta(days=d)
        count = 0
        for art in _articles.values():
            if art.agency_id != agency_id:
                continue
            if not _is_valid_article(art):
                continue
            if _published_within(art, cutoff):
                count += 1

        label = f"Last 24 Hours" if d == 1 else f"Last {d} Days"
        options.append({
            "days": d,
            "label": label,
            "article_count": count,
            "download_url": f"/api/v1/bulletin/download/{agency_id}?days={d}",
            "excel_url": f"/api/v1/bulletin/download-excel/{agency_id}?days={d}",
        })

    return {
        "agency_id": agency_id,
        "agency_name": agency.name,
        "options": options,
        "total_articles_in_archive": len([a for a in _articles.values() if a.agency_id == agency_id]),
    }


@router.get("/download-excel/{agency_id}", dependencies=guard("viewer"))
async def download_bulletin_excel(
    agency_id: str,
    days: int = Query(1, description="Number of days: 1, 2, 3, 4, 5, 7, or 30"),
):
    """Download the bulletin as an Excel (.xlsx) QA sheet — one row per article,
    sorted by category, with the SAME-story-across-outlets grouped together so you
    can spot duplicates and off-topic items at a glance.

    Columns: #, Category, Story Group, Relationship (Primary/Similar), Title,
    Summary, Source, Subscription Required (Yes/No), Relevance, URL.
    """
    if days not in ALLOWED_DAYS:
        raise HTTPException(400, f"Invalid days parameter. Allowed: {ALLOWED_DAYS}")

    try:
        from app.bulletin_intelligence.engine import _articles, _agencies
    except ImportError:
        raise HTTPException(500, "Bulletin engine not available")

    agency = _agencies.get(agency_id)
    if not agency:
        raise HTTPException(404, f"Agency {agency_id} not found")

    cutoff = datetime.now(timezone.utc) - timedelta(days=days)
    filtered = [
        art for art in _articles.values()
        if art.agency_id == agency_id and _is_valid_article(art) and _published_within(art, cutoff)
    ]
    if not filtered:
        raise HTTPException(404, f"No articles found for the last {days} day(s)")

    # Group same-story-across-outlets so 'Similar' rows sit together, and resolve
    # each article's display category — reuse the engine so this matches the
    # bulletin exactly. Fall back to one-row-per-article if anything is unavailable.
    try:
        from app.bulletin_intelligence.engine import _cluster_stories, _section_of, AGT_SECTIONS
        section_index = {s: i for i, s in enumerate(AGT_SECTIONS)}

        def _sec(a):
            try:
                return _section_of(a)
            except Exception:
                return "General"

        ordered = sorted(filtered, key=lambda a: getattr(a, "relevance_score", 0) or 0, reverse=True)
        clusters = _cluster_stories(ordered)
    except Exception as e:
        logger.warning(f"Excel clustering unavailable, flat list: {e}")
        section_index = {}
        _sec = lambda a: getattr(a, "topic", "") or "General"
        clusters = [[a] for a in filtered]

    # Order clusters by their primary's category (bulletin order), then relevance.
    clusters.sort(key=lambda m: (section_index.get(_sec(m[0]), 99),
                                 -(getattr(m[0], "relevance_score", 0) or 0)))

    try:
        buffer = _render_excel_workbook(agency, clusters, _sec)
    except ImportError:
        raise HTTPException(500, "openpyxl not installed. Add 'openpyxl' to requirements.txt")

    now = datetime.now(timezone.utc)
    fname = f"FCC_Bulletin_QA_{days}day_{now.strftime('%Y%m%d')}.xlsx"
    total = sum(len(m) for m in clusters)
    logger.info(f"Excel download: agency={agency_id} days={days} rows={total} groups={len(clusters)}")

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


def _briefing_articles(briefing_id: str):
    """Rehydrate a briefing's stories, in document order.

    A Briefing stores only its rendered HTML, so the stories are recovered by
    matching the URLs in that HTML back to the archive. That matters: it keeps
    every export showing the SAME set the reader sees in the preview, rather
    than each format re-deriving its own list and quietly disagreeing.

    Returns (briefing, agency, articles). Raises HTTPException on any miss.
    """
    try:
        from app.bulletin_intelligence.engine import (
            get_briefing, get_briefing_html, _articles, _agencies,
        )
    except ImportError:
        raise HTTPException(500, "Bulletin engine not available")

    briefing = get_briefing(briefing_id)
    if not briefing:
        raise HTTPException(404, f"Briefing {briefing_id} not found")
    agency_id = briefing.get("agency_id")
    agency = _agencies.get(agency_id)
    if not agency:
        raise HTTPException(404, f"Agency {agency_id} not found")

    html = get_briefing_html(briefing_id) or ""
    seen, ordered_urls = set(), []
    for u in re.findall(r'href="(https?://[^"#]+)"', html):
        if "agtbi.com" in u:            # skip the footer/brand link
            continue
        if u in seen:
            continue
        seen.add(u)
        ordered_urls.append(u)

    by_url = {}
    for a in _articles.values():
        if getattr(a, "agency_id", None) == agency_id and getattr(a, "url", ""):
            by_url.setdefault(a.url, a)
    arts = [by_url[u] for u in ordered_urls if u in by_url]

    # Fallback for old briefings whose stories have aged out of the archive.
    if not arts:
        cap = max(int(briefing.get("article_count") or 0), 50)
        arts = sorted(
            [a for a in _articles.values()
             if a.agency_id == agency_id and _is_valid_article(a)],
            key=lambda a: getattr(a, "relevance_score", 0) or 0, reverse=True,
        )[:cap]
    if not arts:
        raise HTTPException(404, "No articles available to export for this briefing")
    return briefing, agency, arts


@router.get("/briefings/{briefing_id}/excel")
async def download_briefing_excel_public(briefing_id: str):
    """Client-facing Excel for a briefing — PUBLIC, same access model as the HTML
    preview. FCC contacts open both from an emailed link and have no accounts.

    Columns are the reader's view (#, Topic, Source, Headline, Summary, URL) plus
    a Summary sheet of counts by topic. Internal QA signals — relevance score,
    story-group id, subscription flag — are deliberately absent: this endpoint is
    unauthenticated, and those belong to the guarded QA sheet at
    /briefings/{id}/excel-qa.
    """
    briefing, _agency, arts = _briefing_articles(briefing_id)

    try:
        from app.bulletin_intelligence.excel_export import create_bulletin_excel
        from app.bulletin_intelligence.engine import _section_of
    except ImportError as e:
        raise HTTPException(500, f"Excel export unavailable: {e}")

    def _sec(a):
        try:
            return _section_of(a)
        except Exception:
            return "General"

    # Sheet 2 needs the stories Google News carried that our own sources did not.
    # Those are exactly the articles the QA pass added, which the engine tags with
    # source_type "qa" — so the cross-check sheet reports what actually happened in
    # the cycle rather than re-deriving a comparison from scratch here.
    _gn_missing, _qa_report = [], None
    try:
        from app.bulletin_intelligence.google_news_collector import get_qa_report
        _qa_report = get_qa_report(str(briefing.get("agency_id") or "fcc"))
        _gn_missing = [a for a in arts
                       if (getattr(a, "source_type", "") or "") == "qa"]
    except Exception as _e:
        logger.debug(f"Google News cross-check data unavailable: {_e}")

    wb = create_bulletin_excel(briefing, arts, _sec,
                               google_news_missing=_gn_missing,
                               qa_report=_qa_report)
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    agency_tag = str(briefing.get("agency_id") or "fcc").upper()
    date_tag = re.sub(r"[^0-9A-Za-z]+", "_",
                      str(briefing.get("briefing_date", "") or "")).strip("_")[:24]
    fname = f"{agency_tag}_Bulletin_{date_tag or briefing_id}.xlsx"
    logger.info(f"Briefing Excel (public): id={briefing_id} rows={len(arts)}")

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.get("/briefings/{briefing_id}/excel-qa", dependencies=guard("viewer"))
async def download_briefing_excel(briefing_id: str):
    """Download a past briefing (from Run History) as the SAME QA Excel sheet as the
    Daily Briefing — identical columns: #, Category, Story Group, Relationship,
    Title, Summary, Source, Subscription Required, Relevance, URL.

    Stays role-guarded and moved off /excel (which is now the public client sheet)
    because relevance scores and subscription flags are internal QA signals.

    The briefing only stores its rendered HTML, so we rehydrate each story's full
    metadata from the archive by matching the URLs in that HTML. Falls back to the
    agency's current top articles if an old briefing's stories have aged out, so the
    sheet is never empty.
    """
    try:
        from app.bulletin_intelligence.engine import (
            _cluster_stories, _section_of, AGT_SECTIONS,
        )
    except ImportError:
        raise HTTPException(500, "Bulletin engine not available")

    # Steps 1-3 (extract URLs from the stored HTML, rehydrate from the archive,
    # fall back for aged-out briefings) are shared with the public sheet.
    briefing, agency, arts = _briefing_articles(briefing_id)

    # 4) Same clustering + ordering + render as the Daily Briefing Excel.
    section_index = {s: i for i, s in enumerate(AGT_SECTIONS)}

    def _sec(a):
        try:
            return _section_of(a)
        except Exception:
            return "General"

    ordered = sorted(arts, key=lambda a: getattr(a, "relevance_score", 0) or 0, reverse=True)
    try:
        clusters = _cluster_stories(ordered)
    except Exception as e:
        logger.warning(f"Briefing Excel clustering unavailable, flat list: {e}")
        clusters = [[a] for a in ordered]
    clusters.sort(key=lambda m: (section_index.get(_sec(m[0]), 99),
                                 -(getattr(m[0], "relevance_score", 0) or 0)))

    try:
        buffer = _render_excel_workbook(agency, clusters, _sec)
    except ImportError:
        raise HTTPException(500, "openpyxl not installed. Add 'openpyxl' to requirements.txt")

    date_tag = re.sub(r"[^0-9A-Za-z]+", "", str(briefing.get("briefing_date", "") or ""))[:16]
    fname = f"FCC_Briefing_QA_{date_tag or briefing_id}.xlsx"
    total = sum(len(m) for m in clusters)
    logger.info(f"Briefing Excel: id={briefing_id} rows={total} groups={len(clusters)}")

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


def _render_excel_workbook(agency, clusters, sec_of):
    """Build the QA spreadsheet from clustered articles. Each cluster = one primary
    story plus its similar/duplicate coverage, kept adjacent and sharing a Story
    Group number so duplicates are obvious."""
    import io
    import re as _re
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "FCC Bulletin QA"

    # Provider column (col 11) is appended LAST so every existing column index
    # (and its styling) is unchanged — additive, backward-compatible. "Google News
    # Match" (col 12) follows the same rule.
    #
    #   ADDED_BY_QA — Google News carried this and none of our own sources did.
    #                 The QA pass added it; without that pass we would have missed it.
    #   YES         — we already had it AND Google News carried it too.
    #   NO          — we had it; Google News did not surface it this cycle.
    #
    # There is deliberately no "MISSING" state: QA adds what it finds, so a story
    # Google News carried is never absent from the delivered briefing. The
    # distinction that matters to a reviewer is which rows we owe to the QA pass.
    headers = ["#", "Category", "Story Group", "Relationship", "Title", "Summary",
               "Source", "Subscription Required", "Relevance", "URL", "Provider",
               "Google News Match", "QA Score", "Duplicate Flag", "URL Status",
               "Word Count", "Classification Confidence", "Editorial Notes"]
    # Columns L-R are the internal review surface. Two notes on what is NOT here:
    #
    #   URL Status is left blank. Populating it means one HTTP request per article
    #   inside a download handler — slow, and it would make the spreadsheet's
    #   contents depend on whether third-party sites happen to be up. It is a
    #   column for a link-checker job to fill, not this one.
    #
    #   Classification Confidence is blank unless the classifier actually recorded
    #   one. Deriving it from relevance_score would be inventing a number and
    #   presenting it as a measurement.
    try:
        from app.bulletin_intelligence.url_dedup import duplicate_flag as _dup_flag
        from app.bulletin_intelligence.engine import _last_duplicate_groups as _dup_groups
    except Exception:  # noqa: BLE001 — the sheet must build regardless
        _dup_flag, _dup_groups = None, []

    def _qa_extras(article) -> list:
        summary = _re.sub(r"<[^>]+>", "", (getattr(article, "summary", "") or ""))
        words = len(summary.split())
        # QA Score flags rows a reviewer should look at rather than scoring quality:
        # a summary outside the 60-100 word band is the concrete, checkable defect.
        in_band = 60 <= words <= 100
        conf = getattr(article, "classification_confidence", None)
        return [
            "OK" if in_band else "REVIEW",
            (_dup_flag(article, _dup_groups) if _dup_flag else ""),
            "",                       # URL Status — see note above
            words,
            ("" if conf in (None, "") else conf),
            "",                       # Editorial Notes — reviewer writes these
        ]
    try:
        from app.bulletin_intelligence.google_news_collector import (
            get_qa_report, titles_match)
        _gn_titles = list((get_qa_report("fcc") or {}).get("google_titles") or [])
    except Exception:  # noqa: BLE001 — the spreadsheet must build regardless
        _gn_titles, titles_match = [], None

    def _gn_status(article) -> str:
        if (getattr(article, "source_type", "") or "") == "qa":
            return "ADDED_BY_QA"
        if not _gn_titles or titles_match is None:
            return ""  # no QA run this cycle — assert nothing rather than guess
        title = getattr(article, "title", "") or ""
        return "YES" if any(titles_match(title, g) for g in _gn_titles) else "NO"
    ws.append(headers)
    header_fill = PatternFill("solid", fgColor="0F172A")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    for c in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=c)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(vertical="center", horizontal="left", wrap_text=True)

    similar_font = Font(color="64748B", italic=True)
    group_fill = PatternFill("solid", fgColor="F1F5F9")  # tint multi-outlet groups
    row = 2
    num = 0
    for gi, members in enumerate(clusters, 1):
        is_multi = len(members) > 1
        for mi, art in enumerate(members):
            num += 1
            title = (getattr(art, "title", "") or "").split(" - ")[0].split(" | ")[0].strip()
            summary = _re.sub(r"<[^>]+>", "", (getattr(art, "summary", "") or "")).replace("\n", " ").strip()
            if len(summary) > 600:
                summary = summary[:600].rstrip() + "…"
            outlet = getattr(art, "outlet", "") or ""
            url = getattr(art, "url", "") or ""
            rel = getattr(art, "relevance_score", 0) or 0
            paywalled = bool(getattr(art, "is_paywalled", False))
            relationship = "Primary" if mi == 0 else "Similar"
            provider = getattr(art, "provider", "") or getattr(art, "source", "") or ""
            ws.append([num, sec_of(art), gi, relationship, title, summary, outlet,
                       "Yes" if paywalled else "No", f"{int(rel * 100)}%", url, provider,
                       _gn_status(art)] + _qa_extras(art))

            if url:
                uc = ws.cell(row=row, column=10)
                uc.hyperlink = url
                uc.font = Font(color="2563EB", underline="single")
            if paywalled:
                ws.cell(row=row, column=8).font = Font(bold=True, color="B45309")
            if relationship == "Similar":
                ws.cell(row=row, column=4).font = similar_font
                ws.cell(row=row, column=5).font = similar_font
            if is_multi:
                ws.cell(row=row, column=3).fill = group_fill
            row += 1

    widths = [5, 24, 11, 12, 50, 72, 22, 18, 10, 42, 16, 18,
              10, 14, 12, 11, 22, 30]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    for r in range(2, row):
        ws.cell(row=r, column=5).alignment = Alignment(wrap_text=True, vertical="top")
        ws.cell(row=r, column=6).alignment = Alignment(wrap_text=True, vertical="top")
    ws.freeze_panes = "A2"
    if row > 2:
        ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{row - 1}"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a python-docx paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = paragraph._element.makeelement(qn('w:hyperlink'), {qn('r:id'): r_id})
    new_run = paragraph._element.makeelement(qn('w:r'), {})
    rPr = paragraph._element.makeelement(qn('w:rPr'), {})

    color_elem = paragraph._element.makeelement(qn('w:color'), {qn('w:val'): '2563EB'})
    rPr.append(color_elem)
    u_elem = paragraph._element.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    rPr.append(u_elem)
    b_elem = paragraph._element.makeelement(qn('w:b'), {})
    rPr.append(b_elem)
    sz_elem = paragraph._element.makeelement(qn('w:sz'), {qn('w:val'): '20'})
    rPr.append(sz_elem)
    font_elem = paragraph._element.makeelement(qn('w:rFonts'), {qn('w:ascii'): 'Arial', qn('w:hAnsi'): 'Arial'})
    rPr.append(font_elem)

    new_run.append(rPr)
    text_elem = paragraph._element.makeelement(qn('w:t'), {})
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

    return hyperlink
