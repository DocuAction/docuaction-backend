"""
Document Text Extraction Service
Properly reads text from PDF, DOCX, XLSX, TXT, CSV, and images.
Updated: Claude Vision API for image OCR and scanned PDF OCR.
"""
import os
import base64
import logging
from pathlib import Path

logger = logging.getLogger("docuaction.extractor")


async def extract_text(file_path: str, file_type: str = None) -> str:
    """
    Extract readable text from any supported document type.
    Supported: PDF, DOCX, XLSX, XLS, CSV, TXT, PNG, JPG, JPEG, TIFF, BMP
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if not file_type:
        file_type = path.suffix.lower().replace(".", "")
    file_type = file_type.lower()

    logger.info(f"Extracting text from: {path.name} (type: {file_type})")

    try:
        if file_type == "txt":
            return _read_txt(path)
        elif file_type == "csv":
            return _read_csv(path)
        elif file_type == "pdf":
            return await _read_pdf(path)
        elif file_type in ("docx", "doc"):
            return _read_docx(path)
        elif file_type in ("xlsx", "xls"):
            return _read_xlsx(path)
        elif file_type in ("png", "jpg", "jpeg", "tiff", "bmp"):
            return await _read_image(path, file_type)
        else:
            return _read_txt(path)
    except Exception as e:
        logger.error(f"Text extraction failed for {path.name}: {e}")
        raise


# ═══════════════════════════════════════════════════════
# CLAUDE VISION OCR — used for images AND scanned PDFs
# ═══════════════════════════════════════════════════════

async def _ocr_with_claude_vision(image_data: bytes, media_type: str, source_name: str = "document") -> str:
    """
    Send an image to Claude Vision API and extract all text.
    Works for: JPG, PNG, scanned PDFs (converted to images), TIFF, BMP.
    """
    try:
        import anthropic
        from app.core.config import settings

        api_key = settings.ANTHROPIC_API_KEY
        if not api_key:
            logger.error("ANTHROPIC_API_KEY not set — cannot do OCR")
            return f"[OCR unavailable: Anthropic API key not configured. Please set ANTHROPIC_API_KEY.]"

        client = anthropic.Anthropic(api_key=api_key)

        # Encode image to base64
        b64_image = base64.b64encode(image_data).decode("utf-8")

        # Map file types to media types Claude accepts
        media_map = {
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "png": "image/png",
            "gif": "image/gif",
            "webp": "image/webp",
            "tiff": "image/png",  # Will be converted
            "bmp": "image/png",   # Will be converted
        }
        claude_media = media_map.get(media_type, "image/jpeg")

        logger.info(f"Sending {source_name} to Claude Vision for OCR ({len(image_data)} bytes, {claude_media})")

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": claude_media,
                                "data": b64_image,
                            },
                        },
                        {
                            "type": "text",
                            "text": "Extract ALL text from this image. Include every word, number, heading, label, and caption you can see. Preserve the structure — use headings, paragraphs, and line breaks as they appear in the document. If there are tables, format them with | separators. If there are handwritten notes, do your best to read them. Output ONLY the extracted text, nothing else."
                        },
                    ],
                }
            ],
        )

        # Extract text from response
        text = ""
        for block in response.content:
            if hasattr(block, "text"):
                text += block.text

        if text and len(text.strip()) > 10:
            logger.info(f"Claude Vision OCR: {len(text.split())} words extracted from {source_name}")
            return text.strip()
        else:
            logger.warning(f"Claude Vision returned minimal text for {source_name}")
            return f"[Claude Vision could not extract meaningful text from this image. The image may be too blurry, too small, or contain no readable text.]"

    except ImportError:
        logger.error("anthropic package not installed")
        return "[OCR unavailable: anthropic package not installed. Run: pip install anthropic]"
    except Exception as e:
        logger.error(f"Claude Vision OCR failed for {source_name}: {e}")
        return f"[OCR failed: {str(e)}. Please try uploading a clearer image or a text-based document.]"


# ═══════════════════════════════════════════════════════
# FILE TYPE READERS
# ═══════════════════════════════════════════════════════

def _read_txt(path: Path) -> str:
    """Read plain text files."""
    encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            text = path.read_text(encoding=enc)
            logger.info(f"TXT: {len(text.split())} words extracted")
            return text
        except (UnicodeDecodeError, UnicodeError):
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_csv(path: Path) -> str:
    """Read CSV files as structured text."""
    import csv
    rows = []
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i == 0:
                rows.append("HEADERS: " + " | ".join(row))
            else:
                rows.append(" | ".join(row))
            if i > 500:
                rows.append(f"... (truncated, {i}+ rows)")
                break
    text = "\n".join(rows)
    logger.info(f"CSV: {len(rows)} rows, {len(text.split())} words extracted")
    return text


async def _read_pdf(path: Path) -> str:
    """Read PDF files. Uses PyPDF2 for text PDFs, Claude Vision for scanned PDFs."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                pages.append(f"[Page {i+1}]\n{page_text.strip()}")

        if pages:
            text = "\n\n".join(pages)
            word_count = len(text.split())
            logger.info(f"PDF: {len(reader.pages)} pages, {word_count} words extracted")

            # If very few words extracted, it might be a scanned PDF
            if word_count < 50 and len(reader.pages) > 0:
                logger.info("PDF has very few words — attempting Claude Vision OCR on scanned pages")
                ocr_text = await _ocr_pdf_with_vision(path)
                if ocr_text and len(ocr_text.split()) > word_count:
                    return ocr_text
            return text
        else:
            # No text at all — definitely scanned/image-based PDF
            logger.info("PDF: No extractable text — using Claude Vision OCR")
            return await _ocr_pdf_with_vision(path)

    except ImportError:
        logger.error("PyPDF2 not installed")
        return _read_txt(path)


async def _ocr_pdf_with_vision(path: Path) -> str:
    """Convert PDF pages to images and OCR with Claude Vision."""
    try:
        # Try using pdf2image (requires poppler)
        from pdf2image import convert_from_path
        import io

        images = convert_from_path(str(path), first_page=1, last_page=5, dpi=200)
        all_text = []

        for i, img in enumerate(images):
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            img_bytes = buf.getvalue()

            page_text = await _ocr_with_claude_vision(img_bytes, "png", f"PDF page {i+1}")
            if page_text and not page_text.startswith("["):
                all_text.append(f"[Page {i+1}]\n{page_text}")

        if all_text:
            return "\n\n".join(all_text)

    except ImportError:
        logger.info("pdf2image not installed — sending entire PDF as document to Claude")

    # Fallback: send the raw PDF bytes to Claude as a document
    try:
        import anthropic
        from app.core.config import settings

        api_key = settings.ANTHROPIC_API_KEY
        if not api_key:
            return "[OCR unavailable: Anthropic API key not configured.]"

        client = anthropic.Anthropic(api_key=api_key)
        pdf_bytes = path.read_bytes()
        b64_pdf = base64.b64encode(pdf_bytes).decode("utf-8")

        # Check file size — Claude accepts up to ~30MB base64
        if len(pdf_bytes) > 20 * 1024 * 1024:
            return "[PDF too large for Vision OCR. Maximum 20MB. Please split the PDF or upload individual pages as images.]"

        logger.info(f"Sending entire PDF to Claude as document ({len(pdf_bytes)} bytes)")

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": b64_pdf,
                            },
                        },
                        {
                            "type": "text",
                            "text": "Extract ALL text from this PDF document. Include every word, number, heading, label, table, and caption. Preserve structure with headings and paragraphs. For tables, use | separators. Output ONLY the extracted text."
                        },
                    ],
                }
            ],
        )

        text = ""
        for block in response.content:
            if hasattr(block, "text"):
                text += block.text

        if text and len(text.strip()) > 10:
            logger.info(f"Claude PDF OCR: {len(text.split())} words extracted")
            return text.strip()

    except Exception as e:
        logger.error(f"PDF Vision OCR failed: {e}")

    return "[Could not extract text from this scanned PDF. Please try uploading individual pages as PNG/JPG images, or a text-based PDF.]"


def _read_docx(path: Path) -> str:
    """Read Word documents using python-docx."""
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if para.style and para.style.name and "Heading" in para.style.name:
                    parts.append(f"\n## {text}\n")
                else:
                    parts.append(text)

        for table_idx, table in enumerate(doc.tables):
            parts.append(f"\n[Table {table_idx + 1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

        text = "\n".join(parts)
        logger.info(f"DOCX: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(text.split())} words extracted")
        return text
    except ImportError:
        logger.error("python-docx not installed")
        return _read_txt(path)


def _read_xlsx(path: Path) -> str:
    """Read Excel files using openpyxl."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"\n[Sheet: {sheet_name}]")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                if any(c.strip() for c in cells):
                    parts.append(" | ".join(cells))
                    row_count += 1
                if row_count > 500:
                    parts.append("... (truncated)")
                    break
        wb.close()
        text = "\n".join(parts)
        logger.info(f"XLSX: {len(wb.sheetnames)} sheets, {len(text.split())} words extracted")
        return text
    except ImportError:
        logger.error("openpyxl not installed")
        return _read_txt(path)


async def _read_image(path: Path, file_type: str) -> str:
    """Read text from images using Claude Vision API."""
    image_bytes = path.read_bytes()

    # Check file size
    if len(image_bytes) > 20 * 1024 * 1024:
        return "[Image too large. Maximum 20MB for OCR processing.]"

    # Convert TIFF/BMP to PNG for Claude (it only accepts JPEG, PNG, GIF, WEBP)
    if file_type in ("tiff", "bmp"):
        try:
            from PIL import Image
            import io
            img = Image.open(path)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            image_bytes = buf.getvalue()
            file_type = "png"
            logger.info(f"Converted {file_type.upper()} to PNG for Claude Vision")
        except ImportError:
            logger.warning("Pillow not installed — sending raw image to Claude")

    return await _ocr_with_claude_vision(image_bytes, file_type, path.name)
