"""
Document Text Extraction Service
- Text PDFs → PyPDF2
- Scanned PDFs → Convert to images → Claude Vision OCR
- Images (PNG/JPG) → Claude Vision OCR
- DOCX → python-docx
- TXT/CSV → direct read
"""
import os
import io
import base64
import logging
from pathlib import Path

logger = logging.getLogger("docuaction.document_processor")


def extract_text(file_path: str) -> dict:
    """
    Extract text from any document type.
    Returns: {"text": str, "method": str, "pages": int, "word_count": int}
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif ext in (".txt", ".md", ".csv", ".tsv", ".json", ".xml", ".html"):
        return _extract_text_file(file_path)
    elif ext in (".xlsx", ".xls"):
        return _extract_excel(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"):
        return _extract_image(file_path)
    else:
        return {"text": "", "method": "unsupported", "pages": 0, "word_count": 0, "error": f"Unsupported format: {ext}"}


def _extract_pdf(file_path: str) -> dict:
    """Extract text from PDF. Falls back to vision OCR for scanned PDFs."""
    text = ""
    pages = 0

    # Strategy 1: PyPDF2 text extraction
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages = len(reader.pages)
        page_texts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                page_texts.append(t.strip())
        text = "\n\n".join(page_texts)
    except Exception as e:
        logger.warning(f"PyPDF2 failed: {e}")

    # Check if we got meaningful text
    word_count = len(text.split()) if text else 0

    if word_count > 50:
        logger.info(f"PDF text extraction: {word_count} words from {pages} pages")
        return {"text": text, "method": "text_extraction", "pages": pages, "word_count": word_count}

    # Strategy 2: Claude Vision OCR (for scanned/image PDFs)
    logger.info(f"PDF has minimal text ({word_count} words). Attempting vision OCR...")
    try:
        ocr_text = _claude_vision_pdf(file_path)
        if ocr_text and len(ocr_text.split()) > 20:
            logger.info(f"Vision OCR extracted {len(ocr_text.split())} words")
            return {"text": ocr_text, "method": "vision_ocr", "pages": pages, "word_count": len(ocr_text.split())}
    except Exception as e:
        logger.warning(f"Vision OCR failed: {e}")

    # Return whatever we have
    if text:
        return {"text": text, "method": "partial_extraction", "pages": pages, "word_count": word_count}

    return {"text": "", "method": "failed", "pages": pages, "word_count": 0, "error": "No text could be extracted"}


def _claude_vision_pdf(file_path: str) -> str:
    """Send PDF to Claude as a document for vision-based text extraction."""
    from app.core.config import settings

    with open(file_path, "rb") as f:
        pdf_data = base64.standard_b64encode(f.read()).decode("utf-8")

    # Check file size (Claude limit ~32MB for documents)
    file_size = os.path.getsize(file_path)
    if file_size > 30 * 1024 * 1024:
        logger.warning(f"PDF too large for vision OCR: {file_size // 1024}KB")
        return ""

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_data,
                        },
                    },
                    {
                        "type": "text",
                        "text": "Extract ALL text content from this document. Preserve the structure, headings, tables, and formatting as much as possible. Return only the extracted text, nothing else. If there are tables, format them with pipes (|). Include all numbers, dates, names, and data points.",
                    },
                ],
            }],
        )

        return response.content[0].text.strip()
    except Exception as e:
        logger.error(f"Claude vision PDF extraction failed: {e}")
        return ""


def _extract_image(file_path: str) -> dict:
    """Extract text from an image using Claude Vision."""
    from app.core.config import settings

    ext = Path(file_path).suffix.lower()
    media_types = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".gif": "image/gif", ".webp": "image/webp", ".bmp": "image/bmp"}
    media_type = media_types.get(ext, "image/png")

    with open(file_path, "rb") as f:
        image_data = base64.standard_b64encode(f.read()).decode("utf-8")

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": image_data,
                        },
                    },
                    {
                        "type": "text",
                        "text": "Extract ALL text visible in this image. Preserve structure and formatting. If there are tables, use pipe formatting. Return only the text content.",
                    },
                ],
            }],
        )

        text = response.content[0].text.strip()
        return {"text": text, "method": "vision_ocr", "pages": 1, "word_count": len(text.split())}
    except Exception as e:
        logger.error(f"Image OCR failed: {e}")
        return {"text": "", "method": "failed", "pages": 1, "word_count": 0, "error": str(e)}


def _extract_docx(file_path: str) -> dict:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                text += "\n" + row_text

        return {"text": text, "method": "docx_extraction", "pages": 1, "word_count": len(text.split())}
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return {"text": "", "method": "failed", "pages": 0, "word_count": 0, "error": str(e)}


def _extract_text_file(file_path: str) -> dict:
    """Extract text from plain text files."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return {"text": text, "method": "direct_read", "pages": 1, "word_count": len(text.split())}
    except Exception as e:
        return {"text": "", "method": "failed", "pages": 0, "word_count": 0, "error": str(e)}


def _extract_excel(file_path: str) -> dict:
    """Extract text from Excel files."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        text_parts = []

        for sheet in wb.sheetnames:
            ws = wb[sheet]
            text_parts.append(f"--- Sheet: {sheet} ---")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    text_parts.append(row_text)

        text = "\n".join(text_parts)
        return {"text": text, "method": "excel_extraction", "pages": len(wb.sheetnames), "word_count": len(text.split())}
    except Exception as e:
        logger.error(f"Excel extraction failed: {e}")
        return {"text": "", "method": "failed", "pages": 0, "word_count": 0, "error": str(e)}


async def ensure_document_text(document, upload_dir: str = "./uploads") -> str:
    """
    Ensure a document has extracted text. If not, extract it.
    Call this before AI processing.
    """
    # Check if document already has text content
    existing_text = getattr(document, 'content', '') or getattr(document, 'extracted_text', '') or ''
    if existing_text and len(existing_text.split()) > 50:
        return existing_text

    # Find the file
    file_path = None
    if hasattr(document, 'file_path') and document.file_path:
        file_path = document.file_path
    elif hasattr(document, 'id'):
        # Search for file in uploads directory
        doc_dir = os.path.join(upload_dir, "documents")
        if os.path.exists(doc_dir):
            for f in os.listdir(doc_dir):
                if str(document.id) in f:
                    file_path = os.path.join(doc_dir, f)
                    break

    if not file_path or not os.path.exists(file_path):
        logger.warning(f"Document file not found for ID: {getattr(document, 'id', 'unknown')}")
        return existing_text or ""

    # Extract text
    result = extract_text(file_path)

    if result["text"]:
        logger.info(f"Extracted {result['word_count']} words via {result['method']} from {Path(file_path).name}")
        return result["text"]

    return existing_text or ""
