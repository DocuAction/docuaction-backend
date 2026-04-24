"""
Professional Word Proposal Generator
Generates 16-40 page government proposal documents with:
- Cover page with AGT branding
- Table of Contents
- Multiple detailed sections
- Tables for compliance, staffing, deliverables
- Past performance references
- Proper formatting and styles
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
BLUE = RGBColor(0x00, 0x78, 0xD4)
GRAY = RGBColor(0x60, 0x5E, 0x5C)
DARK = RGBColor(0x32, 0x31, 0x30)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _setup_styles(doc):
    """Configure document styles."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size, color in [(1, 18, NAVY), (2, 14, BLUE), (3, 12, NAVY)]:
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(8)


def _add_cover_page(doc, analysis):
    """Add professional cover page."""
    for _ in range(4):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROPOSAL RESPONSE")
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_paragraph()

    # RFQ Title
    title = analysis.get('title', 'Request for Quote Response')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE

    # Solicitation
    sol = analysis.get('solicitation_number', '')
    if sol:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Solicitation: {sol}")
        run.font.size = Pt(14)
        run.font.color.rgb = GRAY

    # Agency
    agency = analysis.get('agency', '')
    if agency:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Submitted to: {agency}")
        run.font.size = Pt(13)
        run.font.color.rgb = GRAY

    for _ in range(4):
        doc.add_paragraph()

    # Company block
    lines = [
        ("Submitted by:", False, 12, GRAY),
        ("Alliance Global Tech, Inc.", True, 18, NAVY),
        ("5457 Twin Knolls Rd, Suite 300", False, 11, GRAY),
        ("Columbia, MD 21045", False, 11, GRAY),
        ("Phone: (443) 832-4278 | www.agtbi.com", False, 11, GRAY),
        ("", False, 11, GRAY),
        (f"Date: {datetime.now().strftime('%B %d, %Y')}", False, 11, GRAY),
        ("", False, 11, GRAY),
        ("SBA 8(a) Certified | CMMI Level 3 | ISO 27001/9001/20000", False, 10, BLUE),
        ("CAGE: 8ERE8 | UEI: MP2FLV1MAW93", False, 10, GRAY),
        ("GSA MAS: 47QTCA21D003M", False, 10, GRAY),
    ]
    for text, bold, size, color in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold

    # Confidential notice
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROPRIETARY AND CONFIDENTIAL")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.italic = True

    doc.add_page_break()


def _add_toc(doc):
    """Add Table of Contents page."""
    doc.add_heading('Table of Contents', level=1)
    sections = [
        ("1.", "Executive Summary"),
        ("2.", "Technical Approach"),
        ("   2.1", "Requirements Response"),
        ("   2.2", "Technical Methodology"),
        ("   2.3", "Deliverables"),
        ("3.", "Management Approach"),
        ("   3.1", "Project Management Framework"),
        ("   3.2", "Risk Mitigation"),
        ("   3.3", "Quality Assurance"),
        ("   3.4", "Communication Plan"),
        ("4.", "Staffing Plan"),
        ("   4.1", "Key Personnel"),
        ("   4.2", "Labor Categories"),
        ("   4.3", "Transition Plan"),
        ("5.", "Past Performance"),
        ("6.", "Compliance Matrix"),
        ("   6.1", "AGT Certifications"),
        ("7.", "Pricing Narrative"),
        ("Appendix A", "Evaluation Criteria Response"),
    ]
    for num, title in sections:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {title}")
        run.font.size = Pt(11)
        if not num.startswith(" "):
            run.bold = True

    doc.add_page_break()


def _write_section(doc, text):
    """Write AI-generated section content with proper formatting."""
    if not text:
        return
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('* '))
            run.bold = True
        elif line.startswith('- ') or line.startswith('• ') or line.startswith('* '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            doc.add_paragraph(line, style='List Number')
        else:
            doc.add_paragraph(line)


def _add_compliance_table(doc, analysis):
    """Add compliance matrix table."""
    compliance = analysis.get('compliance_requirements', [])

    # Always include standard AGT certifications
    standard = [
        {"requirement": "SBA 8(a) Certification", "description": "Small Business Administration 8(a) Business Development Program", "agt_status": "Met"},
        {"requirement": "CMMI Level 3", "description": "Capability Maturity Model Integration for Development and Services", "agt_status": "Met"},
        {"requirement": "ISO 27001", "description": "Information Security Management System", "agt_status": "Met"},
        {"requirement": "ISO 9001", "description": "Quality Management System", "agt_status": "Met"},
        {"requirement": "ISO 20000", "description": "IT Service Management", "agt_status": "Met"},
        {"requirement": "DoD Facility Clearance", "description": "Department of Defense Facility Security Clearance", "agt_status": "Met"},
    ]

    all_compliance = compliance + [s for s in standard if s['requirement'] not in [c.get('requirement', '') for c in compliance]]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Requirement'
    hdr[1].text = 'Description'
    hdr[2].text = 'AGT Status'

    for item in all_compliance:
        row = table.add_row().cells
        row[0].text = item.get('requirement', '')
        row[1].text = item.get('description', '')
        row[2].text = item.get('agt_status', 'Met')


def _add_labor_table(doc, analysis):
    """Add labor categories table."""
    labor = analysis.get('labor_categories', [])
    if not labor:
        labor = [
            {"title": "Program Manager", "level": "Senior", "estimated_hours_monthly": 160, "suggested_rate": 165},
            {"title": "Technical Lead", "level": "Senior", "estimated_hours_monthly": 160, "suggested_rate": 175},
            {"title": "Software Developer", "level": "Mid", "estimated_hours_monthly": 160, "suggested_rate": 120},
        ]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Labor Category'
    hdr[1].text = 'Level'
    hdr[2].text = 'Hours/Month'
    hdr[3].text = 'Hourly Rate'

    for l in labor:
        row = table.add_row().cells
        row[0].text = l.get('title', '')
        row[1].text = l.get('level', '')
        row[2].text = str(l.get('estimated_hours_monthly', 160))
        row[3].text = f"${l.get('suggested_rate', 0):.2f}"


def _add_past_performance(doc):
    """Add detailed past performance section."""
    performances = [
        {
            "title": "CMS OnePI/OPIDS Program Support",
            "agency": "Centers for Medicare & Medicaid Services (CMS)",
            "contract": "Prime Contract",
            "period": "2014 – Present (10+ years continuous)",
            "value": "$50M+ cumulative",
            "description": (
                "AGT serves as the prime contractor for the CMS One Program Integrator (OnePI) "
                "and Operational Process Improvement and Data Solutions (OPIDS) programs. This "
                "engagement encompasses enterprise data management using Oracle Siebel CRM, IBM "
                "Business Intelligence Publisher (BIP), IBM Cognos analytics, and AWS Virtual "
                "Data Center (VDC) operations. AGT's team manages the full lifecycle of data "
                "integration, reporting, and analytics supporting Medicare and Medicaid operations."
            ),
            "relevance": (
                "Demonstrates AGT's ability to manage large-scale federal IT programs with "
                "complex data environments, strict compliance requirements, and continuous "
                "delivery expectations over extended periods."
            ),
        },
        {
            "title": "IRS Contact Center Configuration Support",
            "agency": "Internal Revenue Service (IRS)",
            "contract": "Subcontract / Direct Support",
            "period": "Multi-year",
            "value": "Multiple task orders",
            "description": (
                "AGT provides contact center technology expertise supporting the IRS's enterprise "
                "communication infrastructure. Key activities include Cisco Finesse development, "
                "UCCE configuration management, call routing optimization, and real-time analytics "
                "dashboard development for call center performance monitoring."
            ),
            "relevance": (
                "Demonstrates AGT's expertise in contact center technologies, Cisco platforms, "
                "and high-availability systems supporting millions of citizen interactions."
            ),
        },
        {
            "title": "U.S. Air Force IT Infrastructure Support",
            "agency": "U.S. Air Force",
            "contract": "Prime / Subcontract",
            "period": "Multi-year",
            "value": "Multiple task orders",
            "description": (
                "AGT provides network engineering, cybersecurity, and IT infrastructure support "
                "for Air Force installations. Work includes DoD 8140.03 compliance, Risk Management "
                "Framework (RMF) implementation, classified environment support, and Secret-level "
                "cleared personnel management."
            ),
            "relevance": (
                "Demonstrates AGT's capability to operate in DoD classified environments, "
                "maintain security clearances, and deliver mission-critical IT services."
            ),
        },
    ]

    for pp in performances:
        doc.add_heading(pp['title'], level=2)

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        fields = [
            ("Agency:", pp['agency']),
            ("Contract Type:", pp['contract']),
            ("Period of Performance:", pp['period']),
            ("Contract Value:", pp['value']),
            ("Relevance:", pp['relevance']),
        ]
        for i, (label, value) in enumerate(fields):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    paragraph.style.font.size = Pt(10)

        doc.add_paragraph()
        doc.add_paragraph(pp['description'])
        doc.add_paragraph()


def generate_proposal_docx(analysis: dict, sections: dict = None, response_text: str = None) -> bytes:
    """Generate complete Word proposal document."""
    doc = Document()
    _setup_styles(doc)

    # Cover page
    _add_cover_page(doc, analysis)

    # Table of Contents
    _add_toc(doc)

    # ── Section 1: Executive Summary ──
    doc.add_heading('1. Executive Summary', level=1)
    if sections and sections.get('executive_summary'):
        _write_section(doc, sections['executive_summary'])
    else:
        summary = analysis.get('summary', '')
        agency = analysis.get('agency', 'the requesting agency')
        doc.add_paragraph(summary)
        doc.add_paragraph(
            f"Alliance Global Tech, Inc. (AGT) is pleased to submit this proposal in response "
            f"to the solicitation from {agency}. As an SBA 8(a) certified small business with "
            f"CMMI Level 3 appraisal and ISO 27001/9001/20000 certifications, AGT brings over "
            f"a decade of proven performance supporting federal IT initiatives across CMS, IRS, "
            f"DoD, Treasury, and other agencies."
        )
        strengths = analysis.get('agt_strengths', [])
        if strengths:
            doc.add_heading('Key Differentiators', level=2)
            for s in strengths:
                doc.add_paragraph(s, style='List Bullet')

    win_prob = analysis.get('win_probability', '')
    if win_prob:
        p = doc.add_paragraph()
        run = p.add_run(f"Win Assessment: {win_prob}")
        run.font.color.rgb = BLUE
        run.bold = True

    doc.add_page_break()

    # ── Section 2: Technical Approach ──
    doc.add_heading('2. Technical Approach', level=1)
    if sections and sections.get('technical_approach'):
        _write_section(doc, sections['technical_approach'])
    else:
        approach = analysis.get('recommended_approach', '')
        if approach:
            for para in approach.split('\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())

    # Requirements table
    requirements = analysis.get('requirements', [])
    if requirements:
        doc.add_heading('2.1 Requirements Response', level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = '#'
        hdr[1].text = 'Requirement'
        hdr[2].text = 'Category'
        hdr[3].text = 'Priority'
        for i, req in enumerate(requirements):
            row = table.add_row().cells
            row[0].text = str(req.get('id', i + 1))
            row[1].text = req.get('description', '')
            row[2].text = req.get('category', '')
            row[3].text = req.get('priority', '')

    # Deliverables
    deliverables = analysis.get('deliverables', [])
    if deliverables:
        doc.add_heading('2.2 Deliverables', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Deliverable'
        hdr[1].text = 'Description'
        hdr[2].text = 'Timeline'
        for d in deliverables:
            row = table.add_row().cells
            row[0].text = d.get('name', '')
            row[1].text = d.get('description', '')
            row[2].text = d.get('due', 'TBD')

    doc.add_page_break()

    # ── Section 3: Management Approach ──
    doc.add_heading('3. Management Approach', level=1)
    if sections and sections.get('management_approach'):
        _write_section(doc, sections['management_approach'])
    else:
        doc.add_paragraph(
            "AGT employs a hybrid project management methodology combining Agile/Scrum "
            "practices with traditional milestone-based oversight to ensure both flexibility "
            "and accountability throughout the contract lifecycle."
        )
        mgmt = [
            "Dedicated PMP-certified Program Manager with 10+ years federal experience",
            "Weekly status reports with KPIs, milestones, and risk register updates",
            "Monthly Executive Progress Reviews with government stakeholders",
            "CMMI Level 3 processes for requirements management and configuration control",
            "DevSecOps pipeline for continuous integration and automated testing",
            "Earned Value Management (EVM) for FFP and CPFF contracts",
        ]
        doc.add_heading('3.1 Project Management Framework', level=2)
        for item in mgmt:
            doc.add_paragraph(item, style='List Bullet')

    # Risk table
    risks = analysis.get('risks', [])
    if risks:
        doc.add_heading('3.2 Risk Mitigation', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Risk'
        hdr[1].text = 'Severity'
        hdr[2].text = 'Mitigation Strategy'
        for r in risks:
            row = table.add_row().cells
            row[0].text = r.get('risk', '')
            row[1].text = r.get('severity', '')
            row[2].text = r.get('mitigation', '')

    doc.add_page_break()

    # ── Section 4: Staffing Plan ──
    doc.add_heading('4. Staffing Plan', level=1)
    if sections and sections.get('staffing_plan'):
        _write_section(doc, sections['staffing_plan'])
    else:
        doc.add_paragraph(
            "AGT will staff this effort with qualified, experienced professionals drawn from "
            "our bench of cleared and credentialed employees. Our staffing approach ensures "
            "continuity, rapid onboarding, and minimal transition risk."
        )

    doc.add_heading('4.1 Labor Categories', level=2)
    _add_labor_table(doc, analysis)

    # Key personnel
    key_personnel = analysis.get('key_personnel_needed', [])
    if key_personnel:
        doc.add_heading('4.2 Key Personnel', level=2)
        for kp in key_personnel:
            p = doc.add_paragraph()
            run = p.add_run(f"{kp.get('role', '')}: ")
            run.bold = True
            skills = kp.get('skills', [])
            p.add_run(', '.join(skills) if skills else 'Requirements to be confirmed')
            if kp.get('clearance'):
                p.add_run(f" | Clearance: {kp['clearance']}")

    doc.add_page_break()

    # ── Section 5: Past Performance ──
    doc.add_heading('5. Past Performance', level=1)
    if sections and sections.get('past_performance'):
        _write_section(doc, sections['past_performance'])
    else:
        doc.add_paragraph(
            "AGT has a demonstrated track record of delivering high-quality IT solutions and "
            "services to federal agencies. The following past performance references are directly "
            "relevant to the requirements of this solicitation."
        )
    _add_past_performance(doc)

    doc.add_page_break()

    # ── Section 6: Compliance Matrix ──
    doc.add_heading('6. Compliance Matrix', level=1)
    if sections and sections.get('compliance'):
        _write_section(doc, sections['compliance'])
    doc.add_heading('6.1 Compliance Status', level=2)
    _add_compliance_table(doc, analysis)

    doc.add_page_break()

    # ── Section 7: Pricing Narrative ──
    doc.add_heading('7. Pricing Narrative', level=1)
    if sections and sections.get('pricing_narrative'):
        _write_section(doc, sections['pricing_narrative'])
    else:
        doc.add_paragraph(
            "AGT's pricing reflects our commitment to providing best value to the Government. "
            "Our rates are competitive and aligned with our GSA MAS Schedule. All labor rates "
            "are fully burdened to include fringe benefits, overhead, general and administrative "
            "costs, and profit."
        )
        pricing_points = [
            "Fully burdened labor rates inclusive of fringe, overhead, and G&A",
            "Competitive GSA Schedule pricing available for direct ordering",
            "No hidden fees, escalation clauses, or surprise charges",
            "Volume discounts for multi-year engagements",
            "Transparent cost breakdown available upon request",
            "Firm-fixed pricing for defined deliverables ensuring cost certainty",
        ]
        for item in pricing_points:
            doc.add_paragraph(item, style='List Bullet')

    # Evaluation criteria
    eval_criteria = analysis.get('evaluation_criteria', [])
    if eval_criteria:
        doc.add_page_break()
        doc.add_heading('Appendix A: Evaluation Criteria Response', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Factor'
        hdr[1].text = 'Weight'
        hdr[2].text = 'AGT Response Summary'
        for ec in eval_criteria:
            row = table.add_row().cells
            row[0].text = ec.get('factor', '')
            row[1].text = str(ec.get('weight', ''))
            row[2].text = ec.get('description', '')

    # ── FOOTER ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Alliance Global Tech, Inc. | SBA 8(a) Certified | CAGE: 8ERE8 | UEI: MP2FLV1MAW93\n"
        "5457 Twin Knolls Rd, Suite 300, Columbia, MD 21045 | www.agtbi.com\n"
        "PROPRIETARY AND CONFIDENTIAL"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
