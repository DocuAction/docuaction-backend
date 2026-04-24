"""
Document text extraction — supports PDF, DOCX, TXT files.
"""
import io


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from Word document."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text based on file extension."""
    fname = filename.lower()
    if fname.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif fname.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    elif fname.endswith(('.txt', '.md', '.csv')):
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            return file_bytes.decode('latin-1')
    elif fname.endswith('.xlsx') or fname.endswith('.xls'):
        # Basic Excel reading - just extract as text
        try:
            return file_bytes.decode('utf-8')
        except:
            return "[Excel file detected — paste text content for analysis]"
    else:
        try:
            return file_bytes.decode('utf-8')
        except:
            return file_bytes.decode('latin-1')
