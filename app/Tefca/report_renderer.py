"""
DocuAction TEFCA — Report Renderers (PDF + DOCX)
ONC TEFCA Review Protocol — Contract No. 7571MN26F80064 (HHS/ONC)

Renders the JSON reports produced by reporting.py (and the QA scorecard) into
formatted, AGT-branded PDF and DOCX deliverables. Rendering is generic and
schema-driven: it reads the ACTUAL fields the report carries (report_type, task,
overall_category_counts, per-QHIN breakdowns, executive_summary, suggested
methodology changes, contract info) and includes each section only when present.
This intentionally does NOT assume a fixed per-type field set — the reporting
engine emits different keys per report type, and a generic renderer stays correct
as those evolve.

Honest labeling: any report whose data_source is MOCK gets a prominent red
"MOCK DATA" banner so a demonstration report can never be mistaken for a
production deliverable.

Uses python-docx and reportlab (both pinned platform deps). No new dependencies.
"""
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, List, Tuple

# ── Branding ─────────────────────────────────────────────────────────────────
AGT_NAVY_HEX = "#003366"
AGT_GRAY_HEX = "#666666"
MOCK_RED_HEX = "#DC2626"

_CATEGORY_ORDER = ["no_discrepancy", "minor_administrative", "inexplicable", "non_compliant"]
_CATEGORY_LABELS = {
    "no_discrepancy": "No Discrepancy",
    "minor_administrative": "Minor / Administrative",
    "inexplicable": "Inexplicable",
    "non_compliant": "Non-Compliant",
}


def _is_mock(report_data: dict) -> bool:
    return str(report_data.get("data_source", "")).upper().startswith("MOCK")


def _title(report_data: dict) -> str:
    return report_data.get("task") or f"TEFCA {report_data.get('report_type', 'Report').title()} Report"


def _category_rows(report_data: dict) -> List[Tuple[str, int, str]]:
    """[(label, count, 'NN.N%')] from overall_category_counts, ordered by severity."""
    counts = report_data.get("overall_category_counts") or {}
    total = sum(v for v in counts.values() if isinstance(v, (int, float))) or 0
    rows = []
    for cat in _CATEGORY_ORDER:
        if cat in counts:
            n = counts.get(cat, 0)
            pct = f"{(100.0 * n / total):.1f}%" if total else "0.0%"
            rows.append((_CATEGORY_LABELS[cat], int(n), pct))
    return rows


def _qhin_rows(report_data: dict) -> List[Tuple[str, str, str]]:
    """[(qhin, total, compliance_score)] from whichever per-QHIN block exists."""
    block = (report_data.get("stratified_by_qhin")
             or report_data.get("per_qhin_breakdown")
             or report_data.get("per_qhin_scorecard") or {})
    rows = []
    if isinstance(block, dict):
        for qhin, d in sorted(block.items()):
            d = d or {}
            total = d.get("total", d.get("reviews", ""))
            score = d.get("compliance_score", d.get("pass_rate", ""))
            rows.append((str(qhin), str(total), str(score)))
    return rows


def _meta_lines(report_data: dict) -> List[str]:
    return [
        f"Report Type: {report_data.get('report_type', 'N/A')}",
        f"Period: {report_data.get('period_start', 'N/A')} to {report_data.get('period_end', 'N/A')}",
        f"Total Reviews: {report_data.get('total_reviews', 'N/A')}",
        f"Generated: {report_data.get('generated_at') or datetime.utcnow().isoformat()}",
        f"Data Source: {report_data.get('data_source', 'UNKNOWN')}",
    ]


def _suggestions(report_data: dict) -> List[str]:
    out = []
    for s in report_data.get("suggested_methodology_changes", []) or []:
        if isinstance(s, dict):
            out.append(f"{s.get('finding_type', '')} ({s.get('occurrences', 0)}x): {s.get('suggested_change', '')}")
        else:
            out.append(str(s))
    return out


# ══════════════════════════════════════════════════════════════════════════
# DOCX
# ══════════════════════════════════════════════════════════════════════════
def render_report_docx(report_data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    navy = RGBColor(0x00, 0x33, 0x66)
    red = RGBColor(0xDC, 0x26, 0x26)
    doc = Document()

    # Header / footer (add runs rather than mutating shared styles).
    hp = doc.sections[0].header.paragraphs[0]
    hr = hp.add_run("Alliance Global Tech, Inc. — HHS/ONC TEFCA ARC")
    hr.font.size = Pt(9)
    hr.font.color.rgb = navy
    fp = doc.sections[0].footer.paragraphs[0]
    fr = fp.add_run("CONFIDENTIAL — Alliance Global Tech, Inc. | Contract 7571MN26F80064")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    title = doc.add_heading(_title(report_data), level=1)
    if title.runs:
        title.runs[0].font.color.rgb = navy

    # MOCK banner up top (impossible to miss).
    if _is_mock(report_data):
        b = doc.add_paragraph()
        run = b.add_run("⚠ MOCK DATA — synthetic demonstration data only. "
                        "Do not use for operational decisions.")
        run.font.bold = True
        run.font.color.rgb = red

    for line in _meta_lines(report_data):
        doc.add_paragraph(line)

    # Executive summary (quarterly and others carry a prose summary).
    if report_data.get("executive_summary"):
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(str(report_data["executive_summary"]))

    # Discrepancy category counts.
    cat_rows = _category_rows(report_data)
    if cat_rows:
        doc.add_heading("Discrepancy Category Counts", level=2)
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        h = t.rows[0].cells
        h[0].text, h[1].text, h[2].text = "Category", "Count", "Percentage"
        for label, count, pct in cat_rows:
            c = t.add_row().cells
            c[0].text, c[1].text, c[2].text = label, str(count), pct

    # Per-QHIN breakdown.
    qhin_rows = _qhin_rows(report_data)
    if qhin_rows:
        doc.add_heading("Per-QHIN Breakdown", level=2)
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        h = t.rows[0].cells
        h[0].text, h[1].text, h[2].text = "QHIN", "Total", "Compliance Score"
        for qhin, total, score in qhin_rows:
            c = t.add_row().cells
            c[0].text, c[1].text, c[2].text = qhin, total, score

    # Suggested methodology changes.
    sug = _suggestions(report_data)
    if sug:
        doc.add_heading("Suggested Methodology Changes", level=2)
        for s in sug:
            doc.add_paragraph(s, style="List Bullet")

    # AGT disclaimer.
    note = report_data.get("agt_does_not_adjudicate")
    if note:
        doc.add_heading("Note", level=2)
        doc.add_paragraph(str(note))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════
# PDF
# ══════════════════════════════════════════════════════════════════════════
def render_report_pdf(report_data: dict) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="AGTTitle", fontSize=17, leading=21,
                              textColor=colors.HexColor(AGT_NAVY_HEX),
                              fontName="Helvetica-Bold", spaceAfter=12))
    styles.add(ParagraphStyle(name="AGTMeta", fontSize=9, leading=13,
                              textColor=colors.HexColor(AGT_GRAY_HEX)))
    styles.add(ParagraphStyle(name="AGTWarn", fontSize=11, leading=15,
                              textColor=colors.HexColor(MOCK_RED_HEX), fontName="Helvetica-Bold"))
    el: List[Any] = []

    el.append(Paragraph(_title(report_data), styles["AGTTitle"]))
    if _is_mock(report_data):
        el.append(Paragraph("MOCK DATA — synthetic demonstration data only. "
                            "Do not use for operational decisions.", styles["AGTWarn"]))
        el.append(Spacer(1, 8))
    for line in _meta_lines(report_data):
        el.append(Paragraph(line, styles["AGTMeta"]))
    el.append(Spacer(1, 14))

    if report_data.get("executive_summary"):
        el.append(Paragraph("Executive Summary", styles["Heading2"]))
        el.append(Paragraph(str(report_data["executive_summary"]), styles["Normal"]))
        el.append(Spacer(1, 10))

    def _table(title: str, header: List[str], rows: List[Tuple], widths: List[int]):
        el.append(Paragraph(title, styles["Heading2"]))
        data = [header] + [list(r) for r in rows]
        t = Table(data, colWidths=widths)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(AGT_NAVY_HEX)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f0f0")]),
        ]))
        el.append(t)
        el.append(Spacer(1, 12))

    cat_rows = _category_rows(report_data)
    if cat_rows:
        _table("Discrepancy Category Counts", ["Category", "Count", "Percentage"], cat_rows, [260, 90, 90])

    qhin_rows = _qhin_rows(report_data)
    if qhin_rows:
        _table("Per-QHIN Breakdown", ["QHIN", "Total", "Compliance Score"], qhin_rows, [240, 100, 120])

    sug = _suggestions(report_data)
    if sug:
        el.append(Paragraph("Suggested Methodology Changes", styles["Heading2"]))
        for s in sug:
            el.append(Paragraph("• " + s, styles["Normal"]))
        el.append(Spacer(1, 10))

    note = report_data.get("agt_does_not_adjudicate")
    if note:
        el.append(Paragraph("Note", styles["Heading2"]))
        el.append(Paragraph(str(note), styles["AGTMeta"]))

    doc.build(el)
    return buf.getvalue()
